"""
Import PAX-Zahlen aus Stärkemeldung Word-Dokumenten
Liest Word-Dokumente aus den Ordnern 01_Januar, 02_Februar, 03_März
und extrahiert PAX-Zahlen, um sie in die Datenbank einzutragen.
"""
import os
import re
from datetime import datetime
from docx import Document
from database.pax_db import speichere_tages_pax, speichere_tages_einsaetze

# Pfade zu den drei Ordnern
ORDNER = [
    r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26\06_Stärkemeldung\01_Januar",
    r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26\06_Stärkemeldung\02_Februar",
    r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26\06_Stärkemeldung\03_März",
]


def extrahiere_datum_aus_dateiname(dateiname: str) -> str:
    """
    Extrahiert das Von-Datum aus dem Dateinamen.
    Z.B. "Stärkemeldung 01.03.2026 - 02.03.2026.docx" -> "2026-03-01"
    """
    # Muster: DD.MM.YYYY
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', dateiname)
    if match:
        tag, monat, jahr = match.groups()
        return f"{jahr}-{monat}-{tag}"
    return None


def extrahiere_pax_aus_dokument(doc_pfad: str) -> tuple[int, int]:
    """
    Extrahiert PAX-Zahl und SL-Einsätze aus einem Word-Dokument.
    Sucht nach Zahlen im unteren Bereich des Dokuments.
    
    Returns: (pax_zahl, sl_einsaetze)
    """
    try:
        doc = Document(doc_pfad)
        
        # Suche in den letzten Absätzen nach PAX-Zahlen
        # Typisches Muster: "- 191 -" (Zahl zwischen Bindestrichen)
        pax_zahl = None
        sl_einsaetze = None
        
        # Durchsuche die letzten 30 Absätze von unten nach oben
        for para in reversed(doc.paragraphs[-30:]):
            text = para.text.strip()
            
            # Suche nach PAX-Zahl im Format "- 191 -"
            if pax_zahl is None:
                match = re.search(r'-\s*(\d+)\s*-', text)
                if match:
                    pax_zahl = int(match.group(1))
                    continue  # PAX-Zahl gefunden, weiter suchen nach SL-Einsätzen
            
            # Suche nach SL-Einsätze (optional, falls vorhanden)
            if sl_einsaetze is None:
                # Muster: "SL-Einsätze: 6", "Einsätze: 6"
                match = re.search(r'(?:SL[\s-]*)?Einsätze[\s:]+(\d+)', text, re.IGNORECASE)
                if match:
                    sl_einsaetze = int(match.group(1))
        
        # Wenn PAX-Zahl nicht gefunden, versuche auch in Tabellen
        if pax_zahl is None:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        
                        if pax_zahl is None:
                            match = re.search(r'-\s*(\d+)\s*-', text)
                            if match:
                                pax_zahl = int(match.group(1))
                        
                        if sl_einsaetze is None:
                            match = re.search(r'(?:SL[\s-]*)?Einsätze[\s:]+(\d+)', text, re.IGNORECASE)
                            if match:
                                sl_einsaetze = int(match.group(1))
        
        return pax_zahl or 0, sl_einsaetze or 0
        
    except Exception as e:
        print(f"  ⚠️  Fehler beim Lesen von {os.path.basename(doc_pfad)}: {e}")
        return 0, 0


def verarbeite_ordner(ordner_pfad: str):
    """Verarbeitet alle Word-Dokumente in einem Ordner."""
    if not os.path.exists(ordner_pfad):
        print(f"⚠️  Ordner existiert nicht: {ordner_pfad}")
        return
    
    ordner_name = os.path.basename(ordner_pfad)
    print(f"\n{'='*80}")
    print(f"Verarbeite Ordner: {ordner_name}")
    print(f"{'='*80}")
    
    dateien = [f for f in os.listdir(ordner_pfad) if f.endswith('.docx')]
    dateien.sort()
    
    erfolge = 0
    fehler = 0
    
    for datei in dateien:
        doc_pfad = os.path.join(ordner_pfad, datei)
        
        # Datum aus Dateiname extrahieren
        datum = extrahiere_datum_aus_dateiname(datei)
        if not datum:
            print(f"⚠️  Konnte Datum nicht extrahieren aus: {datei}")
            fehler += 1
            continue
        
        # PAX-Zahl und SL-Einsätze extrahieren
        pax_zahl, sl_einsaetze = extrahiere_pax_aus_dokument(doc_pfad)
        
        if pax_zahl > 0 or sl_einsaetze > 0:
            # In Datenbank speichern
            try:
                if pax_zahl > 0:
                    speichere_tages_pax(datum, pax_zahl)
                if sl_einsaetze > 0:
                    speichere_tages_einsaetze(datum, sl_einsaetze)
                
                print(f"✓ {datum}: PAX={pax_zahl:4d}, SL-Einsätze={sl_einsaetze:2d}  ({datei})")
                erfolge += 1
            except Exception as e:
                print(f"⚠️  Fehler beim Speichern für {datum}: {e}")
                fehler += 1
        else:
            print(f"⚠️  Keine Daten gefunden in: {datei}")
            fehler += 1
    
    print(f"\n{ordner_name}: {erfolge} erfolgreich, {fehler} Fehler")


def main():
    """Hauptfunktion: Verarbeitet alle drei Ordner."""
    print("\n" + "="*80)
    print("PAX-Import aus Word-Dokumenten")
    print("="*80)
    
    gesamt_erfolge = 0
    gesamt_fehler = 0
    
    for ordner in ORDNER:
        verarbeite_ordner(ordner)
    
    print("\n" + "="*80)
    print("Import abgeschlossen!")
    print("="*80)


if __name__ == "__main__":
    main()
