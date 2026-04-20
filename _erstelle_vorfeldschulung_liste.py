# -*- coding: utf-8 -*-
"""
Erstellt eine ausdruckbare Vorfeldschulung-Liste als Word-Dokument.
Tabelle: Nr. | Name | Vorname | Vorfeldschulung gültig bis | Unterschrift
Alle aktiven Mitarbeiter alphabetisch nach Nachname.

Aufruf: python _erstelle_vorfeldschulung_liste.py
Ausgabe: Daten/Vordrucke/Vorfeldschulung_Liste_JJJJMMTT.docx
         (+ Desktop-Kopie falls OneDrive-Desktop vorhanden)
"""
import os, sys
from datetime import date
from pathlib import Path

# Projektpfad in sys.path aufnehmen, damit functions/ importierbar ist
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# ── Pfade ─────────────────────────────────────────────────────────────────────
_BASE  = Path(os.path.dirname(os.path.abspath(__file__)))
LOGO   = _BASE / "Daten" / "Email" / "Logo.jpg"
ZIEL   = _BASE / "Daten" / "Vordrucke"
ZIEL.mkdir(parents=True, exist_ok=True)

_OD_DESKTOP = Path(r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V\Desktop")

# ── Mitarbeiter laden ─────────────────────────────────────────────────────────
def _lade_mitarbeiter() -> list[dict]:
    try:
        from functions.schulungen_db import lade_alle_mitarbeiter
        return lade_alle_mitarbeiter(aktiv_only=True)
    except Exception as e:
        print(f"[WARNUNG] Datenbankzugriff fehlgeschlagen: {e}")
        return []

# ── Hilfsfunktion: Zelle formatieren ─────────────────────────────────────────
def _zelle(cell, text: str, bold: bool = False, size_pt: int = 11,
           bg_hex: str = "", center: bool = False):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    if bg_hex:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex.lstrip("#"))
        tcPr.append(shd)

    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)

# ── Dokument erstellen ────────────────────────────────────────────────────────
def erstelle_liste(mitarbeiter: list[dict]) -> Path:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Seitenlayout: A4 Hochformat ───────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.5)

    # ── Logo + Kopfzeile ──────────────────────────────────────────────────
    if LOGO.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = logo_p.add_run()
        run.add_picture(str(LOGO), width=Cm(4.5))
        logo_p.paragraph_format.space_after = Pt(4)
    else:
        doc.add_paragraph("DRK – Erste-Hilfe-Station Flughafen Köln/Bonn")

    # ── Titel ─────────────────────────────────────────────────────────────
    titel_p = doc.add_paragraph()
    titel_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    titel_r = titel_p.add_run("Vorfeldschulung – Gültigkeitsliste")
    titel_r.bold = True
    titel_r.font.size = Pt(16)
    titel_r.font.color.rgb = RGBColor(0x15, 0x65, 0xA8)  # #1565a8
    titel_p.paragraph_format.space_before = Pt(4)
    titel_p.paragraph_format.space_after  = Pt(2)

    # ── Datum ─────────────────────────────────────────────────────────────
    datum_p = doc.add_paragraph()
    datum_r = datum_p.add_run(f"Stand: {date.today().strftime('%d.%m.%Y')}   |   "
                               f"Anzahl Mitarbeiter: {len(mitarbeiter)}")
    datum_r.font.size = Pt(9)
    datum_r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    datum_p.paragraph_format.space_after = Pt(6)

    # ── Tabelle ───────────────────────────────────────────────────────────
    # Spalten: Nr. | Nachname | Vorname | Vorfeldschulung gültig bis | Unterschrift
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"

    # Spaltenbreiten
    spalten_cm = [1.0, 5.5, 4.5, 4.5, 2.0]
    for i, breite in enumerate(spalten_cm):
        for cell in tbl.columns[i].cells:
            cell.width = Cm(breite)

    # Header-Zeile
    hdr_farbe = "1565A8"
    hdr_texte = ["Nr.", "Nachname", "Vorname", "Vorfeldschulung\ngültig bis", "Unterschrift"]
    hdr_cells = tbl.rows[0].cells
    for i, (zelle, text) in enumerate(zip(hdr_cells, hdr_texte)):
        _zelle(zelle, text, bold=True, size_pt=10, bg_hex=hdr_farbe, center=True)
        # Schriftfarbe weiß in der Überschrift
        p = zelle.paragraphs[0]
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Mitarbeiter-Zeilen
    zeile_farbe_hell  = "EBF3FC"  # leicht blau für jede 2. Zeile
    zeile_farbe_weiss = "FFFFFF"

    for nr, ma in enumerate(mitarbeiter, start=1):
        row = tbl.add_row()
        cells = row.cells
        bg = zeile_farbe_hell if nr % 2 == 0 else zeile_farbe_weiss

        _zelle(cells[0], str(nr),                        size_pt=10, bg_hex=bg, center=True)
        _zelle(cells[1], ma.get("nachname", ""),          size_pt=10, bg_hex=bg)
        _zelle(cells[2], ma.get("vorname",  ""),          size_pt=10, bg_hex=bg)
        _zelle(cells[3], "",                              size_pt=10, bg_hex=bg)  # leer zum Eintragen
        _zelle(cells[4], "",                              size_pt=10, bg_hex=bg)  # leer zum Unterschreiben

    # ── Hinweis-Fußzeile im Dokument ──────────────────────────────────────
    hinweis = doc.add_paragraph()
    hinweis.paragraph_format.space_before = Pt(8)
    hinweis_r = hinweis.add_run(
        "Bitte das Ablaufdatum der Vorfeldschulung in die Spalte"
        ' \u201eVorfeldschulung g\u00fcltig bis\u201c eintragen'
        " und mit Unterschrift best\u00e4tigen."
    )
    hinweis_r.font.size = Pt(8)
    hinweis_r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    hinweis_r.italic = True

    # ── Speichern ─────────────────────────────────────────────────────────
    dateiname = f"Vorfeldschulung_Liste_{date.today().strftime('%Y%m%d')}.docx"
    ziel_pfad = ZIEL / dateiname
    doc.save(str(ziel_pfad))
    print(f"[OK] Gespeichert: {ziel_pfad}")

    # Optionale Desktop-Kopie
    if _OD_DESKTOP.exists():
        desktop_pfad = _OD_DESKTOP / dateiname
        import shutil
        shutil.copy2(str(ziel_pfad), str(desktop_pfad))
        print(f"[OK] Desktop-Kopie: {desktop_pfad}")

    return ziel_pfad


# ── Hauptprogramm ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Vorfeldschulung-Liste wird erstellt …")

    mitarbeiter = _lade_mitarbeiter()
    if not mitarbeiter:
        print("[FEHLER] Keine Mitarbeiter gefunden. Bitte zuerst Mitarbeiter in Nesk3 anlegen.")
        sys.exit(1)

    print(f"  {len(mitarbeiter)} aktive Mitarbeiter geladen.")

    try:
        pfad = erstelle_liste(mitarbeiter)
        print(f"\nFertig! Datei öffnen mit:")
        print(f"  start \"\" \"{pfad}\"")
        # Datei direkt öffnen
        import subprocess
        subprocess.Popen(["start", "", str(pfad)], shell=True)
    except ImportError:
        print("[FEHLER] python-docx ist nicht installiert. Bitte ausführen:")
        print("  pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"[FEHLER] {e}")
        sys.exit(1)
