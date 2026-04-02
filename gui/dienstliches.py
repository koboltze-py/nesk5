"""
Dienstliches Widget
Dienstliche Protokolle: Einsätze nach Einsatzstatistik-Vorlage FKB
"""
import os
import sys
import sqlite3
from contextlib import contextmanager
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QFrame, QTabWidget, QTableWidget, QTableWidgetItem, QHeaderView,
    QDialog, QDialogButtonBox, QFormLayout, QLineEdit, QTextEdit,
    QComboBox, QDateEdit, QTimeEdit, QSpinBox, QRadioButton,
    QButtonGroup, QMessageBox, QSizePolicy, QScrollArea, QGroupBox,
    QCheckBox, QGridLayout
)
from PySide6.QtCore import Qt, QDate, QTime
from PySide6.QtGui import QFont, QColor

from config import FIORI_BLUE, FIORI_TEXT, BASE_DIR, FIORI_BORDER

try:
    from database.sanmat_db import SanmatDB as _SanmatDB
except ImportError:
    _SanmatDB = None


# ──────────────────────────────────────────────────────────────────────────────
#  Datenbank
# ──────────────────────────────────────────────────────────────────────────────

_EINSATZ_DB_DIR  = os.path.join(BASE_DIR, "database SQL")
_EINSATZ_DB_PFAD = os.path.join(_EINSATZ_DB_DIR, "einsaetze.db")
_EINSATZ_PROTO_DIR = os.path.join(BASE_DIR, "Daten", "Einsatz")

# ══════════════════════════════════════════════════════════════════════════════
#  Patienten DRK Station - Datenbank
# ══════════════════════════════════════════════════════════════════════════════

_PATIENTEN_DB_PFAD = os.path.join(_EINSATZ_DB_DIR, "patienten_station.db")

_CREATE_PATIENTEN_SQL = """
CREATE TABLE IF NOT EXISTS patienten (
    id                      INTEGER PRIMARY KEY AUTOINCREMENT,
    erstellt_am             TEXT    NOT NULL,
    datum                   TEXT    NOT NULL,
    uhrzeit                 TEXT    NOT NULL,
    behandlungsdauer        INTEGER DEFAULT 0,
    -- Patient
    patient_typ             TEXT    DEFAULT '',
    patient_abteilung       TEXT    DEFAULT '',
    patient_name            TEXT    DEFAULT '',
    patient_alter           INTEGER DEFAULT 0,
    geschlecht              TEXT    DEFAULT '',
    -- Ereignis
    hergang_was             TEXT    DEFAULT '',
    hergang_wie             TEXT    DEFAULT '',
    unfall_ort              TEXT    DEFAULT '',
    -- Beschwerdebild
    beschwerde_art          TEXT    DEFAULT '',
    symptome                TEXT    DEFAULT '',
    -- ABCDE
    abcde_a                 TEXT    DEFAULT '',
    abcde_b                 TEXT    DEFAULT '',
    abcde_c                 TEXT    DEFAULT '',
    abcde_d                 TEXT    DEFAULT '',
    abcde_e                 TEXT    DEFAULT '',
    -- Monitoring
    monitoring_bz           TEXT    DEFAULT '',
    monitoring_rr           TEXT    DEFAULT '',
    monitoring_spo2         TEXT    DEFAULT '',
    monitoring_hf           TEXT    DEFAULT '',
    -- Vorerkrankungen & Medikamente des Patienten
    vorerkrankungen         TEXT    DEFAULT '',
    medikamente_patient     TEXT    DEFAULT '',
    -- Behandlung
    diagnose                TEXT    DEFAULT '',
    massnahmen              TEXT    DEFAULT '',
    medikamente_gegeben     INTEGER DEFAULT 0,
    medikamente_gegeben_was TEXT    DEFAULT '',
    -- Arbeitsunfall
    arbeitsunfall           INTEGER DEFAULT 0,
    arbeitsunfall_details   TEXT    DEFAULT '',
    -- Personal & Abschluss
    drk_ma1                 TEXT    DEFAULT '',
    drk_ma2                 TEXT    DEFAULT '',
    weitergeleitet          TEXT    DEFAULT '',
    bemerkung               TEXT    DEFAULT ''
);

CREATE TABLE IF NOT EXISTS verbrauchsmaterial (
    id               INTEGER PRIMARY KEY AUTOINCREMENT,
    patienten_id     INTEGER NOT NULL,
    material         TEXT    NOT NULL,
    menge            INTEGER DEFAULT 1,
    einheit          TEXT    DEFAULT 'Stk',
    FOREIGN KEY (patienten_id) REFERENCES patienten(id) ON DELETE CASCADE
);

CREATE TABLE IF NOT EXISTS medikamente (
    id               INTEGER PRIMARY KEY AUTOINCREMENT,
    patienten_id     INTEGER NOT NULL,
    medikament       TEXT    NOT NULL,
    dosis            TEXT    DEFAULT '',
    applikation      TEXT    DEFAULT '',
    FOREIGN KEY (patienten_id) REFERENCES patienten(id) ON DELETE CASCADE
);
"""


def _ensured_patienten_db() -> str:
    _NEW_COLS = [
        ("patient_typ",             "TEXT DEFAULT ''"),
        ("patient_abteilung",       "TEXT DEFAULT ''"),
        ("hergang_was",             "TEXT DEFAULT ''"),
        ("hergang_wie",             "TEXT DEFAULT ''"),
        ("unfall_ort",              "TEXT DEFAULT ''"),
        ("beschwerde_art",          "TEXT DEFAULT ''"),
        ("abcde_a",                 "TEXT DEFAULT ''"),
        ("abcde_b",                 "TEXT DEFAULT ''"),
        ("abcde_c",                 "TEXT DEFAULT ''"),
        ("abcde_d",                 "TEXT DEFAULT ''"),
        ("abcde_e",                 "TEXT DEFAULT ''"),
        ("monitoring_bz",           "TEXT DEFAULT ''"),
        ("monitoring_rr",           "TEXT DEFAULT ''"),
        ("monitoring_spo2",         "TEXT DEFAULT ''"),
        ("monitoring_hf",           "TEXT DEFAULT ''"),
        ("vorerkrankungen",         "TEXT DEFAULT ''"),
        ("medikamente_patient",     "TEXT DEFAULT ''"),
        ("medikamente_gegeben",     "INTEGER DEFAULT 0"),
        ("medikamente_gegeben_was", "TEXT DEFAULT ''"),
        ("arbeitsunfall",           "INTEGER DEFAULT 0"),
        ("arbeitsunfall_details",   "TEXT DEFAULT ''"),
    ]
    os.makedirs(_EINSATZ_DB_DIR, exist_ok=True)
    con = sqlite3.connect(_PATIENTEN_DB_PFAD, timeout=5)
    con.execute("PRAGMA journal_mode = WAL")
    con.execute("PRAGMA synchronous  = NORMAL")
    con.execute("PRAGMA busy_timeout  = 5000")
    con.execute("PRAGMA foreign_keys = ON")
    con.executescript(_CREATE_PATIENTEN_SQL)
    for col, typ in _NEW_COLS:
        try:
            con.execute(f"ALTER TABLE patienten ADD COLUMN {col} {typ}")
        except Exception:
            pass  # Spalte existiert bereits
    con.commit()
    con.close()
    return _PATIENTEN_DB_PFAD


@contextmanager
def _patienten_db():
    _ensured_patienten_db()
    con = sqlite3.connect(_PATIENTEN_DB_PFAD, timeout=5)
    con.execute("PRAGMA journal_mode = WAL")
    con.execute("PRAGMA synchronous  = NORMAL")
    con.execute("PRAGMA busy_timeout  = 5000")
    con.execute("PRAGMA foreign_keys = ON")
    con.row_factory = sqlite3.Row
    try:
        yield con
        con.commit()
    finally:
        con.close()


def _push_pat(table: str, row_id: int) -> None:
    try:
        from database.turso_sync import push_row
        conn = sqlite3.connect(_PATIENTEN_DB_PFAD, timeout=5)
        conn.row_factory = sqlite3.Row
        row = conn.execute(f"SELECT * FROM {table} WHERE id = ?", (row_id,)).fetchone()
        conn.close()
        if row:
            push_row(_PATIENTEN_DB_PFAD, table, dict(row))
    except Exception:
        pass


def patient_speichern(daten: dict, verbrauchsmaterial: list[dict]) -> int:
    with _patienten_db() as con:
        cur = con.execute(
            """
            INSERT INTO patienten (
                erstellt_am, datum, uhrzeit, behandlungsdauer,
                patient_typ, patient_abteilung, patient_name, patient_alter, geschlecht,
                hergang_was, hergang_wie, unfall_ort,
                beschwerde_art, symptome,
                abcde_a, abcde_b, abcde_c, abcde_d, abcde_e,
                monitoring_bz, monitoring_rr, monitoring_spo2, monitoring_hf,
                vorerkrankungen, medikamente_patient,
                diagnose, massnahmen, medikamente_gegeben, medikamente_gegeben_was,
                arbeitsunfall, arbeitsunfall_details,
                drk_ma1, drk_ma2, weitergeleitet, bemerkung
            ) VALUES (
                :erstellt_am, :datum, :uhrzeit, :behandlungsdauer,
                :patient_typ, :patient_abteilung, :patient_name, :patient_alter, :geschlecht,
                :hergang_was, :hergang_wie, :unfall_ort,
                :beschwerde_art, :symptome,
                :abcde_a, :abcde_b, :abcde_c, :abcde_d, :abcde_e,
                :monitoring_bz, :monitoring_rr, :monitoring_spo2, :monitoring_hf,
                :vorerkrankungen, :medikamente_patient,
                :diagnose, :massnahmen, :medikamente_gegeben, :medikamente_gegeben_was,
                :arbeitsunfall, :arbeitsunfall_details,
                :drk_ma1, :drk_ma2, :weitergeleitet, :bemerkung
            )
            """,
            {
                "erstellt_am":             datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "datum":                   daten.get("datum", ""),
                "uhrzeit":                 daten.get("uhrzeit", ""),
                "behandlungsdauer":        daten.get("behandlungsdauer", 0),
                "patient_typ":             daten.get("patient_typ", ""),
                "patient_abteilung":       daten.get("patient_abteilung", ""),
                "patient_name":            daten.get("patient_name", ""),
                "patient_alter":           daten.get("patient_alter", 0),
                "geschlecht":              daten.get("geschlecht", ""),
                "hergang_was":             daten.get("hergang_was", ""),
                "hergang_wie":             daten.get("hergang_wie", ""),
                "unfall_ort":              daten.get("unfall_ort", ""),
                "beschwerde_art":          daten.get("beschwerde_art", ""),
                "symptome":                daten.get("symptome", ""),
                "abcde_a":                 daten.get("abcde_a", ""),
                "abcde_b":                 daten.get("abcde_b", ""),
                "abcde_c":                 daten.get("abcde_c", ""),
                "abcde_d":                 daten.get("abcde_d", ""),
                "abcde_e":                 daten.get("abcde_e", ""),
                "monitoring_bz":           daten.get("monitoring_bz", ""),
                "monitoring_rr":           daten.get("monitoring_rr", ""),
                "monitoring_spo2":         daten.get("monitoring_spo2", ""),
                "monitoring_hf":           daten.get("monitoring_hf", ""),
                "vorerkrankungen":         daten.get("vorerkrankungen", ""),
                "medikamente_patient":     daten.get("medikamente_patient", ""),
                "diagnose":                daten.get("diagnose", ""),
                "massnahmen":              daten.get("massnahmen", ""),
                "medikamente_gegeben":     daten.get("medikamente_gegeben", 0),
                "medikamente_gegeben_was": daten.get("medikamente_gegeben_was", ""),
                "arbeitsunfall":           daten.get("arbeitsunfall", 0),
                "arbeitsunfall_details":   daten.get("arbeitsunfall_details", ""),
                "drk_ma1":                 daten.get("drk_ma1", ""),
                "drk_ma2":                 daten.get("drk_ma2", ""),
                "weitergeleitet":          daten.get("weitergeleitet", ""),
                "bemerkung":               daten.get("bemerkung", ""),
            },
        )
        patienten_id = cur.lastrowid
        for vm in verbrauchsmaterial:
            con.execute(
                "INSERT INTO verbrauchsmaterial (patienten_id, material, menge, einheit) VALUES (?, ?, ?, ?)",
                (patienten_id, vm.get("material", ""), vm.get("menge", 1), vm.get("einheit", "Stk"))
            )
        for med in daten.get("_medikamente", []):
            con.execute(
                "INSERT INTO medikamente (patienten_id, medikament, dosis, applikation) VALUES (?, ?, ?, ?)",
                (patienten_id, med.get("medikament", ""), med.get("dosis", ""), med.get("applikation", ""))
            )
    _push_pat("patienten", patienten_id)
    try:
        from database.turso_sync import push_replace_by_fk
        push_replace_by_fk(_PATIENTEN_DB_PFAD, "verbrauchsmaterial", "patienten_id", patienten_id)
        push_replace_by_fk(_PATIENTEN_DB_PFAD, "medikamente", "patienten_id", patienten_id)
    except Exception:
        pass
    return patienten_id


def patient_aktualisieren(row_id: int, daten: dict, verbrauchsmaterial: list[dict]) -> None:
    with _patienten_db() as con:
        con.execute(
            """
            UPDATE patienten SET
                datum=:datum, uhrzeit=:uhrzeit, behandlungsdauer=:behandlungsdauer,
                patient_typ=:patient_typ, patient_abteilung=:patient_abteilung,
                patient_name=:patient_name, patient_alter=:patient_alter, geschlecht=:geschlecht,
                hergang_was=:hergang_was, hergang_wie=:hergang_wie, unfall_ort=:unfall_ort,
                beschwerde_art=:beschwerde_art, symptome=:symptome,
                abcde_a=:abcde_a, abcde_b=:abcde_b, abcde_c=:abcde_c,
                abcde_d=:abcde_d, abcde_e=:abcde_e,
                monitoring_bz=:monitoring_bz, monitoring_rr=:monitoring_rr,
                monitoring_spo2=:monitoring_spo2, monitoring_hf=:monitoring_hf,
                vorerkrankungen=:vorerkrankungen, medikamente_patient=:medikamente_patient,
                diagnose=:diagnose, massnahmen=:massnahmen,
                medikamente_gegeben=:medikamente_gegeben,
                medikamente_gegeben_was=:medikamente_gegeben_was,
                arbeitsunfall=:arbeitsunfall, arbeitsunfall_details=:arbeitsunfall_details,
                drk_ma1=:drk_ma1, drk_ma2=:drk_ma2,
                weitergeleitet=:weitergeleitet, bemerkung=:bemerkung
            WHERE id=:id
            """,
            {
                "id":                      row_id,
                "datum":                   daten.get("datum", ""),
                "uhrzeit":                 daten.get("uhrzeit", ""),
                "behandlungsdauer":        daten.get("behandlungsdauer", 0),
                "patient_typ":             daten.get("patient_typ", ""),
                "patient_abteilung":       daten.get("patient_abteilung", ""),
                "patient_name":            daten.get("patient_name", ""),
                "patient_alter":           daten.get("patient_alter", 0),
                "geschlecht":              daten.get("geschlecht", ""),
                "hergang_was":             daten.get("hergang_was", ""),
                "hergang_wie":             daten.get("hergang_wie", ""),
                "unfall_ort":              daten.get("unfall_ort", ""),
                "beschwerde_art":          daten.get("beschwerde_art", ""),
                "symptome":                daten.get("symptome", ""),
                "abcde_a":                 daten.get("abcde_a", ""),
                "abcde_b":                 daten.get("abcde_b", ""),
                "abcde_c":                 daten.get("abcde_c", ""),
                "abcde_d":                 daten.get("abcde_d", ""),
                "abcde_e":                 daten.get("abcde_e", ""),
                "monitoring_bz":           daten.get("monitoring_bz", ""),
                "monitoring_rr":           daten.get("monitoring_rr", ""),
                "monitoring_spo2":         daten.get("monitoring_spo2", ""),
                "monitoring_hf":           daten.get("monitoring_hf", ""),
                "vorerkrankungen":         daten.get("vorerkrankungen", ""),
                "medikamente_patient":     daten.get("medikamente_patient", ""),
                "diagnose":                daten.get("diagnose", ""),
                "massnahmen":              daten.get("massnahmen", ""),
                "medikamente_gegeben":     daten.get("medikamente_gegeben", 0),
                "medikamente_gegeben_was": daten.get("medikamente_gegeben_was", ""),
                "arbeitsunfall":           daten.get("arbeitsunfall", 0),
                "arbeitsunfall_details":   daten.get("arbeitsunfall_details", ""),
                "drk_ma1":                 daten.get("drk_ma1", ""),
                "drk_ma2":                 daten.get("drk_ma2", ""),
                "weitergeleitet":          daten.get("weitergeleitet", ""),
                "bemerkung":               daten.get("bemerkung", ""),
            },
        )
        con.execute("DELETE FROM verbrauchsmaterial WHERE patienten_id=?", (row_id,))
        for vm in verbrauchsmaterial:
            con.execute(
                "INSERT INTO verbrauchsmaterial (patienten_id, material, menge, einheit) VALUES (?, ?, ?, ?)",
                (row_id, vm.get("material", ""), vm.get("menge", 1), vm.get("einheit", "Stk"))
            )
        con.execute("DELETE FROM medikamente WHERE patienten_id=?", (row_id,))
        for med in daten.get("_medikamente", []):
            con.execute(
                "INSERT INTO medikamente (patienten_id, medikament, dosis, applikation) VALUES (?, ?, ?, ?)",
                (row_id, med.get("medikament", ""), med.get("dosis", ""), med.get("applikation", ""))
            )
    try:
        from database.turso_sync import push_replace_by_fk
        push_replace_by_fk(_PATIENTEN_DB_PFAD, "verbrauchsmaterial", "patienten_id", row_id)
        push_replace_by_fk(_PATIENTEN_DB_PFAD, "medikamente", "patienten_id", row_id)
    except Exception:
        pass
    _push_pat("patienten", row_id)


def patient_loeschen(row_id: int) -> None:
    with _patienten_db() as con:
        con.execute("DELETE FROM patienten WHERE id=?", (row_id,))
    try:
        from database.turso_sync import push_delete
        push_delete(_PATIENTEN_DB_PFAD, "patienten", row_id)
    except Exception:
        pass


def lade_patienten(
    *,
    monat: int | None = None,
    jahr: int | None = None,
    suchtext: str | None = None,
) -> list[dict]:
    where_parts: list[str] = []
    params: list = []
    if monat is not None:
        where_parts.append("substr(datum,4,2)=?")
        params.append(f"{monat:02d}")
    if jahr is not None:
        where_parts.append("substr(datum,7,4)=?")
        params.append(str(jahr))
    if suchtext:
        t = f"%{suchtext}%"
        where_parts.append(
            "(patient_name LIKE ? OR symptome LIKE ? OR diagnose LIKE ?"
            " OR hergang_was LIKE ? OR beschwerde_art LIKE ?"
            " OR massnahmen LIKE ? OR unfall_ort LIKE ?"
            " OR drk_ma1 LIKE ? OR drk_ma2 LIKE ?)"
        )
        params.extend([t, t, t, t, t, t, t, t, t])
    sql = "SELECT * FROM patienten"
    if where_parts:
        sql += " WHERE " + " AND ".join(where_parts)
    sql += " ORDER BY substr(datum,7,4)||substr(datum,4,2)||substr(datum,1,2) ASC, id ASC"
    with _patienten_db() as con:
        rows = con.execute(sql, params).fetchall()
    return [dict(r) for r in rows]


def lade_verbrauchsmaterial(patienten_id: int) -> list[dict]:
    with _patienten_db() as con:
        rows = con.execute(
            "SELECT * FROM verbrauchsmaterial WHERE patienten_id=? ORDER BY id",
            (patienten_id,)
        ).fetchall()
    return [dict(r) for r in rows]


def lade_medikamente(patienten_id: int) -> list[dict]:
    with _patienten_db() as con:
        rows = con.execute(
            "SELECT * FROM medikamente WHERE patienten_id=? ORDER BY id",
            (patienten_id,)
        ).fetchall()
    return [dict(r) for r in rows]


def verfuegbare_jahre_patienten() -> list[int]:
    with _patienten_db() as con:
        rows = con.execute(
            "SELECT DISTINCT substr(datum,7,4) AS j FROM patienten"
            " WHERE length(datum)=10 ORDER BY j DESC"
        ).fetchall()
    result = []
    for r in rows:
        try:
            result.append(int(r["j"]))
        except (ValueError, TypeError):
            pass
    return result

_CREATE_SQL = """
CREATE TABLE IF NOT EXISTS einsaetze (
    id               INTEGER PRIMARY KEY AUTOINCREMENT,
    erstellt_am      TEXT    NOT NULL,
    datum            TEXT    NOT NULL,          -- dd.MM.yyyy
    uhrzeit          TEXT    NOT NULL,          -- HH:mm (Alarmierungszeit)
    einsatzdauer     INTEGER DEFAULT 0,         -- Minuten
    einsatzstichwort TEXT    DEFAULT '',
    einsatzort       TEXT    DEFAULT '',
    einsatznr_drk    TEXT    DEFAULT '',
    drk_ma1          TEXT    DEFAULT '',
    drk_ma2          TEXT    DEFAULT '',
    angenommen       INTEGER DEFAULT 1,         -- 1=Ja, 0=Nein
    grund_abgelehnt  TEXT    DEFAULT '',        -- warum nicht angenommen
    bemerkung        TEXT    DEFAULT '',
    gesendet         INTEGER DEFAULT 0          -- 1 = per E-Mail versendet
);
"""


def _ensured_db() -> str:
    os.makedirs(_EINSATZ_DB_DIR, exist_ok=True)
    con = sqlite3.connect(_EINSATZ_DB_PFAD, timeout=5)
    con.execute("PRAGMA journal_mode = WAL")
    con.execute("PRAGMA synchronous  = NORMAL")
    con.execute("PRAGMA busy_timeout  = 5000")
    con.execute(_CREATE_SQL)
    # Migration: gesendet-Spalte für bestehende Datenbanken ergänzen
    try:
        con.execute("ALTER TABLE einsaetze ADD COLUMN gesendet INTEGER DEFAULT 0")
    except Exception:
        pass  # Spalte existiert bereits
    con.commit()
    con.close()
    return _EINSATZ_DB_PFAD


@contextmanager
def _db():
    _ensured_db()
    con = sqlite3.connect(_EINSATZ_DB_PFAD, timeout=5)
    con.execute("PRAGMA journal_mode = WAL")
    con.execute("PRAGMA synchronous  = NORMAL")
    con.execute("PRAGMA busy_timeout  = 5000")
    con.row_factory = sqlite3.Row
    try:
        yield con
        con.commit()
    finally:
        con.close()


def _push_ein(table: str, row_id: int) -> None:
    try:
        from database.turso_sync import push_row
        conn = sqlite3.connect(_EINSATZ_DB_PFAD, timeout=5)
        conn.row_factory = sqlite3.Row
        row = conn.execute(f"SELECT * FROM {table} WHERE id = ?", (row_id,)).fetchone()
        conn.close()
        if row:
            push_row(_EINSATZ_DB_PFAD, table, dict(row))
    except Exception:
        pass


def einsatz_speichern(daten: dict) -> int:
    with _db() as con:
        cur = con.execute(
            """
            INSERT INTO einsaetze (
                erstellt_am, datum, uhrzeit, einsatzdauer,
                einsatzstichwort, einsatzort, einsatznr_drk,
                drk_ma1, drk_ma2, angenommen, grund_abgelehnt, bemerkung
            ) VALUES (
                :erstellt_am, :datum, :uhrzeit, :einsatzdauer,
                :einsatzstichwort, :einsatzort, :einsatznr_drk,
                :drk_ma1, :drk_ma2, :angenommen, :grund_abgelehnt, :bemerkung
            )
            """,
            {
                "erstellt_am":      datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "datum":            daten.get("datum", ""),
                "uhrzeit":          daten.get("uhrzeit", ""),
                "einsatzdauer":     daten.get("einsatzdauer", 0),
                "einsatzstichwort": daten.get("einsatzstichwort", ""),
                "einsatzort":       daten.get("einsatzort", ""),
                "einsatznr_drk":    daten.get("einsatznr_drk", ""),
                "drk_ma1":          daten.get("drk_ma1", ""),
                "drk_ma2":          daten.get("drk_ma2", ""),
                "angenommen":       1 if daten.get("angenommen", True) else 0,
                "grund_abgelehnt":  daten.get("grund_abgelehnt", ""),
                "bemerkung":        daten.get("bemerkung", ""),
            },
        )
        new_id = cur.lastrowid
    _push_ein("einsaetze", new_id)
    return new_id


def einsatz_aktualisieren(row_id: int, daten: dict) -> None:
    with _db() as con:
        con.execute(
            """
            UPDATE einsaetze SET
                datum=:datum, uhrzeit=:uhrzeit, einsatzdauer=:einsatzdauer,
                einsatzstichwort=:einsatzstichwort, einsatzort=:einsatzort,
                einsatznr_drk=:einsatznr_drk, drk_ma1=:drk_ma1, drk_ma2=:drk_ma2,
                angenommen=:angenommen, grund_abgelehnt=:grund_abgelehnt, bemerkung=:bemerkung
            WHERE id=:id
            """,
            {
                "id":               row_id,
                "datum":            daten.get("datum", ""),
                "uhrzeit":          daten.get("uhrzeit", ""),
                "einsatzdauer":     daten.get("einsatzdauer", 0),
                "einsatzstichwort": daten.get("einsatzstichwort", ""),
                "einsatzort":       daten.get("einsatzort", ""),
                "einsatznr_drk":    daten.get("einsatznr_drk", ""),
                "drk_ma1":          daten.get("drk_ma1", ""),
                "drk_ma2":          daten.get("drk_ma2", ""),
                "angenommen":       1 if daten.get("angenommen", True) else 0,
                "grund_abgelehnt":  daten.get("grund_abgelehnt", ""),
                "bemerkung":        daten.get("bemerkung", ""),
            },
        )
    _push_ein("einsaetze", row_id)


def einsatz_loeschen(row_id: int) -> None:
    with _db() as con:
        con.execute("DELETE FROM einsaetze WHERE id=?", (row_id,))
    try:
        from database.turso_sync import push_delete
        push_delete(_EINSATZ_DB_PFAD, "einsaetze", row_id)
    except Exception:
        pass


def lade_einsaetze(
    *,
    monat: int | None = None,
    jahr: int | None = None,
    suchtext: str | None = None,
) -> list[dict]:
    where_parts: list[str] = []
    params: list = []
    if monat is not None:
        where_parts.append("substr(datum,4,2)=?")
        params.append(f"{monat:02d}")
    if jahr is not None:
        where_parts.append("substr(datum,7,4)=?")
        params.append(str(jahr))
    if suchtext:
        t = f"%{suchtext}%"
        where_parts.append(
            "(einsatzstichwort LIKE ? OR einsatzort LIKE ? OR drk_ma1 LIKE ?"
            " OR drk_ma2 LIKE ? OR einsatznr_drk LIKE ?)"
        )
        params.extend([t, t, t, t, t])

    sql = "SELECT * FROM einsaetze"
    if where_parts:
        sql += " WHERE " + " AND ".join(where_parts)
    sql += " ORDER BY substr(datum,7,4)||substr(datum,4,2)||substr(datum,1,2) ASC, id ASC"

    with _db() as con:
        rows = con.execute(sql, params).fetchall()
    return [dict(r) for r in rows]


def verfuegbare_jahre_einsaetze() -> list[int]:
    with _db() as con:
        rows = con.execute(
            "SELECT DISTINCT substr(datum,7,4) AS j FROM einsaetze"
            " WHERE length(datum)=10 ORDER BY j DESC"
        ).fetchall()
    result = []
    for r in rows:
        try:
            result.append(int(r["j"]))
        except (ValueError, TypeError):
            pass
    return result


def markiere_einsatz_gesendet(row_id: int) -> None:
    """Markiert einen Einsatz als per E-Mail versendet."""
    with _db() as con:
        con.execute("UPDATE einsaetze SET gesendet=1 WHERE id=?", (row_id,))
    _push_ein("einsaetze", row_id)


# ──────────────────────────────────────────────────────────────────────────────
#  Excel-Export
# ──────────────────────────────────────────────────────────────────────────────

_PROTOKOLL_DIR = os.path.join(_EINSATZ_PROTO_DIR, "Protokolle")
_PATIENTEN_PROTO_DIR = os.path.join(BASE_DIR, "Daten", "Patienten Station", "Protokolle")


def export_einsaetze_excel(
    eintraege: list[dict],
    ziel_pfad: str | None = None,
    titel_zeitraum: str = "",
) -> str:
    """
    Exportiert eine Liste von Einsätzen als Excel-Datei.
    Gibt den tatsächlichen Speicherpfad zurück.
    ziel_pfad: vollständiger Pfad inkl. Dateiname; None = Standardordner.
    """
    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side
        )
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        raise ImportError(
            "openpyxl ist nicht installiert. Bitte 'pip install openpyxl' ausführen."
        ) from e

    os.makedirs(_PROTOKOLL_DIR, exist_ok=True)

    if not ziel_pfad:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        dateiname = f"Einsatzprotokoll_{ts}.xlsx"
        ziel_pfad = os.path.join(_PROTOKOLL_DIR, dateiname)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Einsätze"

    # ── Farben / Stile ────────────────────────────────────────────────────
    rot_drk   = "C8102E"   # DRK-Rot
    hell_grau = "F2F2F2"
    gruen_bg  = "E8F5E9"
    rot_bg    = "FFE0E0"
    thin = Side(style="thin", color="AAAAAA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    hdr_font  = Font(bold=True, color="FFFFFF", size=11)
    hdr_fill  = PatternFill("solid", fgColor=rot_drk)
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # ── Titel-Zeile ────────────────────────────────────────────────────────
    ws.merge_cells("A1:K1")
    titel_cell = ws["A1"]
    titel_text = "Einsatzstatistik – Notfalleinsätze FKB (DRK)"
    if titel_zeitraum:
        titel_text += f"  |  {titel_zeitraum}"
    titel_cell.value = titel_text
    titel_cell.font  = Font(bold=True, size=13, color="FFFFFF")
    titel_cell.fill  = PatternFill("solid", fgColor=rot_drk)
    titel_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # ── Untertitel ─────────────────────────────────────────────────────────
    ws.merge_cells("A2:K2")
    sub = ws["A2"]
    sub.value = (
        f"Erstellt am {datetime.now().strftime('%d.%m.%Y %H:%M')}  –  "
        f"{len(eintraege)} Einsatz/ätze"
    )
    sub.font      = Font(italic=True, size=9, color="666666")
    sub.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 16

    # ── Spaltenüberschriften ───────────────────────────────────────────────
    headers = [
        "Lfd. Nr.", "Datum", "Uhrzeit", "Dauer\n(Min.)",
        "Einsatzstichwort", "Einsatzort", "Einsatznr. DRK",
        "DRK MA 1", "DRK MA 2", "Angenommen", "Grund (bei Nein) / Bemerkung"
    ]
    col_widths = [9, 13, 10, 10, 36, 22, 16, 20, 20, 13, 36]

    for col, (h, w) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font      = hdr_font
        cell.fill      = hdr_fill
        cell.alignment = hdr_align
        cell.border    = border
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.row_dimensions[3].height = 32

    # ── Datenzeilen ────────────────────────────────────────────────────────
    ang_count = 0
    for row_idx, e in enumerate(eintraege, start=1):
        ang = bool(e.get("angenommen", 1))
        if ang:
            ang_count += 1
        row_num = row_idx + 3  # Zeile 4+
        bg_color = gruen_bg if ang else rot_bg
        row_fill = PatternFill("solid", fgColor=bg_color)

        dauer = e.get("einsatzdauer", 0) or 0
        grund_bem = e.get("grund_abgelehnt", "") or ""
        if e.get("bemerkung"):
            if grund_bem:
                grund_bem += "  |  " + e["bemerkung"]
            else:
                grund_bem = e["bemerkung"]

        values = [
            row_idx,
            e.get("datum", ""),
            e.get("uhrzeit", ""),
            dauer if dauer else "",
            e.get("einsatzstichwort", ""),
            e.get("einsatzort", ""),
            e.get("einsatznr_drk", "") or "",
            e.get("drk_ma1", ""),
            e.get("drk_ma2", "") or "",
            "Ja" if ang else "Nein",
            grund_bem,
        ]
        for col, val in enumerate(values, start=1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.fill   = row_fill
            cell.border = border
            cell.alignment = Alignment(vertical="center", wrap_text=(col in (5, 11)))
            if col in (1, 3, 4, 10):
                cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[row_num].height = 16

    # ── Statistik-Zeile ────────────────────────────────────────────────────
    stat_row = len(eintraege) + 5
    ws.merge_cells(f"A{stat_row}:K{stat_row}")
    stat_cell = ws.cell(row=stat_row, column=1)
    abgelehnt = len(eintraege) - ang_count
    stat_cell.value = (
        f"Gesamt: {len(eintraege)}   ✅ Angenommen: {ang_count}   "
        f"❌ Abgelehnt / nicht angenommen: {abgelehnt}"
    )
    stat_cell.font  = Font(bold=True, size=10)
    stat_cell.fill  = PatternFill("solid", fgColor="E3F2FD")
    stat_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[stat_row].height = 20

    # ── Zeile einfrieren ────────────────────────────────────────────────────
    ws.freeze_panes = "A4"

    wb.save(ziel_pfad)
    try:
        from functions.dokument_archiv import kopiere_ins_archiv
        kopiere_ins_archiv(ziel_pfad, "einsatzprotokoll")
    except Exception:
        pass
    return ziel_pfad


_PATIENTEN_EXPORT_DIR = os.path.join(BASE_DIR, "Daten", "Pat. Station")


def export_patienten_excel(
    eintraege: list[dict],
    ziel_pfad: str | None = None,
    titel_zeitraum: str = "",
) -> str:
    """
    Exportiert Patienten als DIN-A4-Protokollformulare (ein Blatt pro Patient).
    Gibt den tatsächlichen Speicherpfad zurück.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.worksheet.page import PageMargins
    except ImportError as e:
        raise ImportError(
            "openpyxl ist nicht installiert. Bitte 'pip install openpyxl' ausführen."
        ) from e

    os.makedirs(_PATIENTEN_EXPORT_DIR, exist_ok=True)
    if not ziel_pfad:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        ziel_pfad = os.path.join(_PATIENTEN_EXPORT_DIR, f"Patienten_Export_{ts}.xlsx")

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # Standard-Sheet entfernen

    # ── gemeinsame Stile ───────────────────────────────────────────────────
    ROT     = "C8102E"
    HELLROT = "FAD4DA"
    GRAU    = "F5F5F5"
    DUNKEL  = "333333"
    thin    = Side(style="thin",   color="AAAAAA")
    thick   = Side(style="medium", color="C8102E")
    b_thin  = Border(left=thin, right=thin, top=thin, bottom=thin)
    b_bot   = Border(bottom=Border(bottom=Side(style="thin", color="999999")).bottom)

    def _hdr(ws, row, col, text, span=1, fill=ROT, fsize=10, bold=True, color="FFFFFF"):
        c = ws.cell(row=row, column=col, value=text)
        c.font = Font(bold=bold, size=fsize, color=color)
        c.fill = PatternFill("solid", fgColor=fill)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        if span > 1:
            ws.merge_cells(start_row=row, start_column=col,
                           end_row=row, end_column=col + span - 1)
        return c

    def _lbl(ws, row, col, text, span=1, fsize=9, bold=True):
        c = ws.cell(row=row, column=col, value=text)
        c.font = Font(bold=bold, size=fsize, color=DUNKEL)
        c.fill = PatternFill("solid", fgColor=GRAU)
        c.alignment = Alignment(vertical="center", wrap_text=True)
        if span > 1:
            ws.merge_cells(start_row=row, start_column=col,
                           end_row=row, end_column=col + span - 1)
        return c

    def _val(ws, row, col, text="", span=1, fsize=9, bold=False, wrap=False):
        c = ws.cell(row=row, column=col, value=text or "")
        c.font = Font(bold=bold, size=fsize, color=DUNKEL)
        c.alignment = Alignment(vertical="center", wrap_text=wrap)
        c.border = b_thin
        if span > 1:
            ws.merge_cells(start_row=row, start_column=col,
                           end_row=row, end_column=col + span - 1)
        return c

    def _handschrift(ws, row, col, span=1):
        """Leeres Feld mit Hinweis 'handschriftlich'."""
        c = ws.cell(row=row, column=col, value="____________________")
        c.font = Font(italic=True, size=9, color="AAAAAA")
        c.alignment = Alignment(vertical="bottom")
        c.border = b_thin
        if span > 1:
            ws.merge_cells(start_row=row, start_column=col,
                           end_row=row, end_column=col + span - 1)
        return c

    # ── Pro Patient ein Sheet ──────────────────────────────────────────────
    for idx, p in enumerate(eintraege, start=1):
        sheet_name = f"Patient {idx}"
        ws = wb.create_sheet(title=sheet_name)

        # A4 Portrait, schmale Ränder
        ws.page_setup.paperSize  = ws.PAPERSIZE_A4
        ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
        ws.page_setup.fitToPage  = True
        ws.page_setup.fitToHeight = 1
        ws.page_setup.fitToWidth  = 1
        ws.page_margins = PageMargins(left=0.4, right=0.4, top=0.5, bottom=0.5,
                                      header=0.2, footer=0.2)

        # Spaltenbreiten (6 Spalten = bequem A4)
        COLS = [3, 12, 12, 12, 12, 14]
        for i, w in enumerate(COLS, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        r = 1  # aktueller Row-Zähler

        # ── Kopfzeile ──────────────────────────────────────────────────────
        ws.merge_cells(f"A{r}:F{r}")
        c = ws.cell(row=r, column=1,
                    value=" 🏥  Patientenprotokoll – DRK Erste-Hilfe-Station FKB (DRK)")
        c.font = Font(bold=True, size=13, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=ROT)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[r].height = 26
        r += 1

        ws.merge_cells(f"A{r}:F{r}")
        sub = ws.cell(row=r, column=1,
                      value=f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}   "
                            f"|   Patient Nr. {idx} von {len(eintraege)}")
        sub.font = Font(italic=True, size=8, color="777777")
        sub.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[r].height = 12
        r += 1

        # ── Abschnitt 1: Basisdaten ────────────────────────────────────────
        _hdr(ws, r, 1, "BASISDATEN", span=6, fsize=9)
        ws.row_dimensions[r].height = 15
        r += 1

        # Datum | Uhrzeit | Behandlungsdauer
        _lbl(ws, r, 1, "Datum")
        _val(ws, r, 2, p.get("datum", ""))
        _lbl(ws, r, 3, "Uhrzeit")
        _val(ws, r, 4, p.get("uhrzeit", ""))
        _lbl(ws, r, 5, "Dauer (Min.)")
        _val(ws, r, 6, p.get("behandlungsdauer", "") or "")
        ws.row_dimensions[r].height = 16
        r += 1

        # Typ | BG-Fall
        _lbl(ws, r, 1, "Personentyp")
        _val(ws, r, 2, p.get("patient_typ", "") or "", span=3)
        _lbl(ws, r, 5, "BG-Fall")
        _val(ws, r, 6, "Ja" if p.get("arbeitsunfall") else "Nein")
        ws.row_dimensions[r].height = 16
        r += 1

        # Abteilung
        _lbl(ws, r, 1, "Abteilung")
        _val(ws, r, 2, p.get("patient_abteilung", "") or "", span=5)
        ws.row_dimensions[r].height = 16
        r += 1

        # ── Abschnitt 2: Patient (handschriftlich) ─────────────────────────
        _hdr(ws, r, 1, "PATIENTENDATEN  (handschriftlich – nicht digital gespeichert)",
             span=6, fsize=9, fill="8B1A1A")
        ws.row_dimensions[r].height = 15
        r += 1

        _lbl(ws, r, 1, "Name")
        _handschrift(ws, r, 2, span=2)
        _lbl(ws, r, 4, "Geburtsdatum")
        _handschrift(ws, r, 5, span=2)
        ws.row_dimensions[r].height = 18
        r += 1

        _lbl(ws, r, 1, "Geschlecht")
        _val(ws, r, 2, p.get("geschlecht", "") or "", span=2)
        _lbl(ws, r, 4, "Alter (ca.)")
        _handschrift(ws, r, 5, span=2)
        ws.row_dimensions[r].height = 18
        r += 1

        # ── Abschnitt 3: Ereignis ──────────────────────────────────────────
        _hdr(ws, r, 1, "EREIGNIS", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        _lbl(ws, r, 1, "Was")
        _val(ws, r, 2, p.get("hergang_was", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 32
        r += 1

        _lbl(ws, r, 1, "Wie")
        _val(ws, r, 2, p.get("hergang_wie", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 26
        r += 1

        _lbl(ws, r, 1, "Wo")
        _val(ws, r, 2, p.get("unfall_ort", "") or "", span=5)
        ws.row_dimensions[r].height = 16
        r += 1

        # ── Abschnitt 4: Beschwerdebild ────────────────────────────────────
        _hdr(ws, r, 1, "BESCHWERDEBILD / SYMPTOME", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        _lbl(ws, r, 1, "Art")
        _val(ws, r, 2, p.get("beschwerde_art", "") or "", span=5)
        ws.row_dimensions[r].height = 16
        r += 1

        _lbl(ws, r, 1, "Symptome")
        _val(ws, r, 2, p.get("symptome", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 28
        r += 1

        # ── Abschnitt 5: ABCDE ────────────────────────────────────────────
        _hdr(ws, r, 1, "ABCDE-SCHEMA", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        for letter, key in [("A – Atemweg", "abcde_a"), ("B – Atmung", "abcde_b"),
                              ("C – Kreislauf", "abcde_c"), ("D – Neurologie", "abcde_d"),
                              ("E – Exposition", "abcde_e")]:
            _lbl(ws, r, 1, letter)
            _val(ws, r, 2, p.get(key, "") or "", span=5)
            ws.row_dimensions[r].height = 15
            r += 1

        # ── Abschnitt 6: Monitoring ────────────────────────────────────────
        _hdr(ws, r, 1, "MONITORING", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        _lbl(ws, r, 1, "BZ")
        _val(ws, r, 2, p.get("monitoring_bz", "") or "")
        _lbl(ws, r, 3, "RR")
        _val(ws, r, 4, p.get("monitoring_rr", "") or "")
        _lbl(ws, r, 5, "SpO₂")
        _val(ws, r, 6, p.get("monitoring_spo2", "") or "")
        ws.row_dimensions[r].height = 16
        r += 1

        _lbl(ws, r, 1, "HF")
        _val(ws, r, 2, p.get("monitoring_hf", "") or "", span=5)
        ws.row_dimensions[r].height = 15
        r += 1

        # ── Abschnitt 7: Diagnostik / Therapie ────────────────────────────
        _hdr(ws, r, 1, "DIAGNOSE / MAßNAHMEN", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        _lbl(ws, r, 1, "Diagnose")
        _val(ws, r, 2, p.get("diagnose", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 26
        r += 1

        _lbl(ws, r, 1, "Maßnahmen")
        _val(ws, r, 2, p.get("massnahmen", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 32
        r += 1

        # ── Abschnitt 8: Vorerkrankungen / Medikamente ────────────────────
        _hdr(ws, r, 1, "ANAMNESE", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        _lbl(ws, r, 1, "Vorerkrank.")
        _val(ws, r, 2, p.get("vorerkrankungen", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 22
        r += 1

        _lbl(ws, r, 1, "Med. Patient")
        _val(ws, r, 2, p.get("medikamente_patient", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 22
        r += 1

        # ── Abschnitt 9: BG-Fall ───────────────────────────────────────────
        if p.get("arbeitsunfall"):
            _hdr(ws, r, 1, "⚠  ARBEITSUNFALL / BG-FALL", span=6, fsize=9, fill="8B4513")
            ws.row_dimensions[r].height = 14
            r += 1
            _lbl(ws, r, 1, "Details")
            _val(ws, r, 2, p.get("arbeitsunfall_details", "") or "", span=5, wrap=True)
            ws.row_dimensions[r].height = 22
            r += 1

        # ── Abschnitt 10: Abschluss ────────────────────────────────────────
        _hdr(ws, r, 1, "WEITERGABE / ABSCHLUSS", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        _lbl(ws, r, 1, "Werterleitet")
        _val(ws, r, 2, p.get("weitergeleitet", "") or "", span=5)
        ws.row_dimensions[r].height = 16
        r += 1

        _lbl(ws, r, 1, "Bemerkung")
        _val(ws, r, 2, p.get("bemerkung", "") or "", span=5, wrap=True)
        ws.row_dimensions[r].height = 26
        r += 1

        # ── Unterschriftszeile ─────────────────────────────────────────────
        _hdr(ws, r, 1, "PERSONAL", span=6, fsize=9)
        ws.row_dimensions[r].height = 14
        r += 1

        _lbl(ws, r, 1, "DRK MA 1")
        _val(ws, r, 2, p.get("drk_ma1", "") or "", span=2)
        _lbl(ws, r, 4, "DRK MA 2")
        _val(ws, r, 5, p.get("drk_ma2", "") or "", span=2)
        ws.row_dimensions[r].height = 16
        r += 1

        _lbl(ws, r, 1, "Unterschrift 1")
        _handschrift(ws, r, 2, span=2)
        _lbl(ws, r, 4, "Unterschrift 2")
        _handschrift(ws, r, 5, span=2)
        ws.row_dimensions[r].height = 22
        r += 1

        # ── Druckbereich festlegen ─────────────────────────────────────────
        ws.print_area = f"A1:{get_column_letter(6)}{r-1}"

    wb.save(ziel_pfad)
    try:
        from functions.dokument_archiv import kopiere_ins_archiv
        kopiere_ins_archiv(ziel_pfad, "patienten_excel")
    except Exception:
        pass
    return ziel_pfad


# ──────────────────────────────────────────────────────────────────────────────
#  Hilfsstile
# ──────────────────────────────────────────────────────────────────────────────

def _btn(text: str, color: str = FIORI_BLUE, hover: str = "#0057b8") -> QPushButton:
    btn = QPushButton(text)
    btn.setFixedHeight(32)
    btn.setCursor(Qt.CursorShape.PointingHandCursor)
    btn.setStyleSheet(f"""
        QPushButton {{
            background: {color}; color: white; border: none;
            border-radius: 4px; padding: 4px 14px; font-size: 12px;
        }}
        QPushButton:hover {{ background: {hover}; }}
        QPushButton:disabled {{ background: #bbb; color: #888; }}
    """)
    return btn


def _btn_light(text: str) -> QPushButton:
    btn = QPushButton(text)
    btn.setFixedHeight(32)
    btn.setCursor(Qt.CursorShape.PointingHandCursor)
    btn.setStyleSheet("""
        QPushButton { background:#eee; color:#333; border:none;
            border-radius:4px; padding:4px 14px; font-size:12px; }
        QPushButton:hover { background:#ddd; }
        QPushButton:disabled { background:#f5f5f5; color:#bbb; }
    """)
    return btn


# ──────────────────────────────────────────────────────────────────────────────
#  Dialog: Einsatz erfassen / bearbeiten
# ──────────────────────────────────────────────────────────────────────────────

class _ArtikelPickerDialog(QDialog):
    """Artikel aus Sanmat-Datenbank auswählen und Menge eingeben."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("🧰 Artikel hinzufügen")
        self.setMinimumWidth(520)
        self.setMinimumHeight(420)
        self._artikel: list[dict] = []
        self._table_data: list[dict] = []
        self._build_ui()
        self._laden()

    def _build_ui(self):
        lay = QVBoxLayout(self)
        lay.setSpacing(8)
        lay.setContentsMargins(12, 12, 12, 8)

        f_lay = QHBoxLayout()
        f_lay.addWidget(QLabel("Suche:"))
        self._filter = QLineEdit()
        self._filter.setPlaceholderText("Bezeichnung oder Artikelnr …")
        self._filter.textChanged.connect(self._filtern)
        f_lay.addWidget(self._filter)
        lay.addLayout(f_lay)

        self._table = QTableWidget()
        self._table.setColumnCount(4)
        self._table.setHorizontalHeaderLabels(["Artikelnr.", "Bezeichnung", "Kategorie", "Bestand"])
        self._table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self._table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self._table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self._table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        self._table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self._table.verticalHeader().setVisible(False)
        self._table.setAlternatingRowColors(True)
        self._table.doubleClicked.connect(self._on_double_click)
        lay.addWidget(self._table, 1)

        menge_lay = QHBoxLayout()
        menge_lay.addWidget(QLabel("Menge:"))
        self._menge = QSpinBox()
        self._menge.setRange(1, 9999)
        self._menge.setValue(1)
        self._menge.setFixedWidth(80)
        menge_lay.addWidget(self._menge)
        menge_lay.addStretch()
        lay.addLayout(menge_lay)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("Hinzufügen")
        btns.accepted.connect(self._on_accept)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)

    def _laden(self):
        if _SanmatDB is None:
            return
        try:
            db = _SanmatDB()
            db.initialize()
            self._artikel = db.get_artikel(nur_aktiv=True)
        except Exception:
            self._artikel = []
        self._filtern()

    def _filtern(self):
        text = self._filter.text().strip().lower()
        if text:
            self._table_data = [
                a for a in self._artikel
                if text in (a.get("bezeichnung", "") + " " + a.get("artikelnr", "")).lower()
            ]
        else:
            self._table_data = list(self._artikel)

        self._table.setRowCount(len(self._table_data))
        for r, a in enumerate(self._table_data):
            menge = int(a.get("menge", 0))
            vals = [
                a.get("artikelnr", ""),
                a.get("bezeichnung", ""),
                a.get("kategorie", ""),
                str(menge),
            ]
            for c, v in enumerate(vals):
                item = QTableWidgetItem(v)
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                if c == 3 and menge <= 0:
                    item.setForeground(QColor("#c0392b"))
                self._table.setItem(r, c, item)

    def _on_double_click(self):
        row = self._table.currentRow()
        if row >= 0:
            self.accept()

    def _on_accept(self):
        row = self._table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Auswahl", "Bitte einen Artikel auswählen.")
            return
        self.accept()

    def get_auswahl(self) -> dict | None:
        row = self._table.currentRow()
        if row < 0 or row >= len(self._table_data):
            return None
        a = self._table_data[row]
        return {
            "artikel_id":  a["id"],
            "bezeichnung": a.get("bezeichnung", ""),
            "menge":       self._menge.value(),
        }


class _EinsatzDialog(QDialog):
    """Formular zum Anlegen oder Bearbeiten eines Notfalleinsatzes."""

    _FIELD_STYLE = (
        "QLineEdit, QTextEdit, QTimeEdit, QDateEdit, QSpinBox, QComboBox {"
        "border:1px solid #ccc; border-radius:4px; padding:4px;"
        "font-size:12px; background:white;}"
    )
    _GROUP_STYLE = (
        "QGroupBox { font-weight:bold; font-size:12px; color:#1565a8;"
        "border:1px solid #c5d8f0; border-radius:6px;"
        "margin-top:6px; padding-top:10px; background:#f8fbff;}"
        "QGroupBox::title { subcontrol-origin:margin; left:12px;"
        "padding:0 4px; background:#f8fbff;}"
    )

    # Einsatzstichwörter
    _STICHWOERTER = [
        "",
        "Intern 1",
        "Intern 2",
        "Chirurgisch 1",
        "Chirurgisch 2",
        "Sandienst",
        "Pat. Station",
    ]

    def __init__(self, daten: dict | None = None, parent=None):
        super().__init__(parent)
        self._edit_daten = daten  # None = Neu, dict = Bearbeiten
        self._verbrauch_positionen: list[dict] = []
        self.setWindowTitle(
            "✏️  Einsatz bearbeiten" if daten else "🚑  Neuen Einsatz erfassen"
        )
        self.setMinimumWidth(560)
        self.setMinimumHeight(620)
        self.resize(580, 720)
        self._build_ui()
        if daten:
            self._befuellen(daten)
        self._update_visibility()

    # ── UI ──────────────────────────────────────────────────────────────────

    def _build_ui(self):
        main = QVBoxLayout(self)
        main.setContentsMargins(12, 12, 12, 8)
        main.setSpacing(8)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea{border:none;}")
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setSpacing(10)
        layout.setContentsMargins(4, 4, 4, 4)
        scroll.setWidget(container)
        main.addWidget(scroll, 1)

        # ── Zeitangaben ────────────────────────────────────────────────────
        grp_zeit = QGroupBox("⏱️  Zeitangaben")
        grp_zeit.setStyleSheet(self._GROUP_STYLE)
        zeit_fl = QFormLayout(grp_zeit)
        zeit_fl.setSpacing(8)

        self._datum = QDateEdit()
        self._datum.setCalendarPopup(True)
        self._datum.setDisplayFormat("dd.MM.yyyy")
        self._datum.setDate(QDate.currentDate())
        self._datum.setStyleSheet(self._FIELD_STYLE)
        zeit_fl.addRow("Datum:", self._datum)

        self._uhrzeit = QTimeEdit()
        self._uhrzeit.setDisplayFormat("HH:mm")
        self._uhrzeit.setTime(QTime.currentTime())
        self._uhrzeit.setStyleSheet(self._FIELD_STYLE)
        zeit_fl.addRow("Alarmierungszeit:", self._uhrzeit)

        self._einsatzdauer = QSpinBox()
        self._einsatzdauer.setRange(0, 999)
        self._einsatzdauer.setSuffix("  min")
        self._einsatzdauer.setValue(0)
        self._einsatzdauer.setStyleSheet(self._FIELD_STYLE)
        zeit_fl.addRow("Einsatzdauer:", self._einsatzdauer)

        layout.addWidget(grp_zeit)

        # ── Einsatz-Details ────────────────────────────────────────────────
        grp_detail = QGroupBox("🚨  Einsatz-Details")
        grp_detail.setStyleSheet(self._GROUP_STYLE)
        det_fl = QFormLayout(grp_detail)
        det_fl.setSpacing(8)

        self._stichwort = QComboBox()
        self._stichwort.setEditable(True)
        self._stichwort.addItems(self._STICHWOERTER)
        self._stichwort.setStyleSheet(self._FIELD_STYLE)
        det_fl.addRow("Einsatzstichwort:", self._stichwort)

        self._einsatzort = QLineEdit()
        self._einsatzort.setPlaceholderText("z.B. T1 Gate B11, T2 Ebene 3 ...")
        self._einsatzort.setStyleSheet(self._FIELD_STYLE)
        det_fl.addRow("Einsatzort:", self._einsatzort)

        self._einsatznr_drk = QLineEdit()
        self._einsatznr_drk.setPlaceholderText("DRK interne Einsatznummer")
        self._einsatznr_drk.setStyleSheet(self._FIELD_STYLE)
        det_fl.addRow("Einsatznr. DRK:", self._einsatznr_drk)

        layout.addWidget(grp_detail)

        # ── Alarmierung ────────────────────────────────────────────────────
        grp_alarm = QGroupBox("📻  Alarmierung – Einsatz angenommen?")
        grp_alarm.setStyleSheet(self._GROUP_STYLE)
        alarm_layout = QVBoxLayout(grp_alarm)
        alarm_layout.setSpacing(6)

        self._alarm_group = QButtonGroup(self)
        self._radio_ja   = QRadioButton("✅  Ja – Einsatz angenommen")
        self._radio_nein = QRadioButton("❌  Nein – Einsatz abgelehnt / nicht angenommen")
        self._radio_ja.setChecked(True)
        for rb in (self._radio_ja, self._radio_nein):
            self._alarm_group.addButton(rb)
            alarm_layout.addWidget(rb)

        self._radio_ja.toggled.connect(self._update_visibility)
        self._radio_nein.toggled.connect(self._update_visibility)
        layout.addWidget(grp_alarm)

        # ── Grund (wenn nicht angenommen) ─────────────────────────────────
        self._grp_grund = QGroupBox("❓  Grund der Ablehnung")
        self._grp_grund.setStyleSheet(self._GROUP_STYLE)
        grund_layout = QVBoxLayout(self._grp_grund)
        self._grund_abgelehnt = QTextEdit()
        self._grund_abgelehnt.setPlaceholderText("Warum wurde der Einsatz nicht angenommen?")
        self._grund_abgelehnt.setMaximumHeight(80)
        self._grund_abgelehnt.setStyleSheet(self._FIELD_STYLE)
        grund_layout.addWidget(self._grund_abgelehnt)
        layout.addWidget(self._grp_grund)

        # ── Einsatzkräfte ──────────────────────────────────────────────────
        grp_ma = QGroupBox("👥  Eingesetzte DRK Mitarbeiter")
        grp_ma.setStyleSheet(self._GROUP_STYLE)
        ma_fl = QFormLayout(grp_ma)
        ma_fl.setSpacing(8)

        self._drk_ma1 = QLineEdit()
        self._drk_ma1.setPlaceholderText("DRK Mitarbeiter 1")
        self._drk_ma1.setStyleSheet(self._FIELD_STYLE)
        ma_fl.addRow("DRK MA 1:", self._drk_ma1)

        self._drk_ma2 = QLineEdit()
        self._drk_ma2.setPlaceholderText("DRK Mitarbeiter 2 (optional)")
        self._drk_ma2.setStyleSheet(self._FIELD_STYLE)
        ma_fl.addRow("DRK MA 2:", self._drk_ma2)

        layout.addWidget(grp_ma)

        # ── Verbrauchsmaterial ─────────────────────────────────────────────
        grp_mat = QGroupBox("🧰  Verbrauchsmaterial (Sanmat)")
        grp_mat.setStyleSheet(self._GROUP_STYLE)
        mat_layout = QVBoxLayout(grp_mat)
        mat_layout.setSpacing(6)

        entnehmer_fl = QFormLayout()
        entnehmer_fl.setSpacing(6)
        self._entnehmer = QLineEdit()
        self._entnehmer.setPlaceholderText("Name des Entnehmers (Pflicht bei Materialverbrauch)")
        self._entnehmer.setStyleSheet(self._FIELD_STYLE)
        entnehmer_fl.addRow("Entnehmer:", self._entnehmer)
        mat_layout.addLayout(entnehmer_fl)

        self._mat_table = QTableWidget()
        self._mat_table.setColumnCount(2)
        self._mat_table.setHorizontalHeaderLabels(["Artikel", "Menge"])
        self._mat_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self._mat_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self._mat_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._mat_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._mat_table.verticalHeader().setVisible(False)
        self._mat_table.setMaximumHeight(150)
        self._mat_table.setAlternatingRowColors(True)
        mat_layout.addWidget(self._mat_table)

        mat_btn_lay = QHBoxLayout()
        btn_add_mat = QPushButton("➕ Artikel hinzufügen")
        btn_add_mat.clicked.connect(self._artikel_hinzufuegen)
        mat_btn_lay.addWidget(btn_add_mat)
        btn_rm_mat = QPushButton("🗑 Entfernen")
        btn_rm_mat.clicked.connect(self._artikel_entfernen)
        mat_btn_lay.addWidget(btn_rm_mat)
        mat_btn_lay.addStretch()
        mat_layout.addLayout(mat_btn_lay)

        layout.addWidget(grp_mat)

        # ── Bemerkungen ────────────────────────────────────────────────────
        grp_bem = QGroupBox("📝  Bemerkungen (optional)")
        grp_bem.setStyleSheet(self._GROUP_STYLE)
        bem_layout = QVBoxLayout(grp_bem)
        self._bemerkung = QTextEdit()
        self._bemerkung.setPlaceholderText("Weitere Anmerkungen zum Einsatz ...")
        self._bemerkung.setMaximumHeight(80)
        self._bemerkung.setStyleSheet(self._FIELD_STYLE)
        bem_layout.addWidget(self._bemerkung)
        layout.addWidget(grp_bem)
        layout.addStretch()

        # ── Buttons ────────────────────────────────────────────────────────
        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        label = "✏️  Speichern" if self._edit_daten else "🚑  Einsatz erfassen"
        btns.button(QDialogButtonBox.StandardButton.Ok).setText(label)
        btns.accepted.connect(self._on_accept)
        btns.rejected.connect(self.reject)
        main.addWidget(btns)

    def _update_visibility(self):
        self._grp_grund.setVisible(self._radio_nein.isChecked())

    def _artikel_hinzufuegen(self):
        if _SanmatDB is None:
            QMessageBox.warning(self, "Sanmat", "Sanmat-Modul nicht verfügbar.")
            return
        dlg = _ArtikelPickerDialog(parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            auswahl = dlg.get_auswahl()
            if auswahl:
                self._verbrauch_positionen.append(auswahl)
                self._mat_tabelle_aktualisieren()

    def _artikel_entfernen(self):
        row = self._mat_table.currentRow()
        if row >= 0:
            del self._verbrauch_positionen[row]
            self._mat_tabelle_aktualisieren()

    def _mat_tabelle_aktualisieren(self):
        self._mat_table.setRowCount(len(self._verbrauch_positionen))
        for r, p in enumerate(self._verbrauch_positionen):
            item_bez = QTableWidgetItem(p["bezeichnung"])
            item_bez.setFlags(item_bez.flags() & ~Qt.ItemFlag.ItemIsEditable)
            item_menge = QTableWidgetItem(str(p["menge"]))
            item_menge.setFlags(item_menge.flags() & ~Qt.ItemFlag.ItemIsEditable)
            item_menge.setTextAlignment(Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter)
            self._mat_table.setItem(r, 0, item_bez)
            self._mat_table.setItem(r, 1, item_menge)

    def _befuellen(self, d: dict):
        """Felder mit bestehenden Daten füllen."""
        try:
            parts = d.get("datum", "").split(".")
            if len(parts) == 3:
                self._datum.setDate(QDate(int(parts[2]), int(parts[1]), int(parts[0])))
        except Exception:
            pass
        try:
            h, m = d.get("uhrzeit", "00:00").split(":")
            self._uhrzeit.setTime(QTime(int(h), int(m)))
        except Exception:
            pass
        self._einsatzdauer.setValue(d.get("einsatzdauer", 0) or 0)
        # Stichwort
        stw = d.get("einsatzstichwort", "")
        idx = self._stichwort.findText(stw)
        if idx >= 0:
            self._stichwort.setCurrentIndex(idx)
        else:
            self._stichwort.setCurrentText(stw)
        self._einsatzort.setText(d.get("einsatzort", ""))
        self._einsatznr_drk.setText(d.get("einsatznr_drk", ""))
        self._drk_ma1.setText(d.get("drk_ma1", ""))
        self._drk_ma2.setText(d.get("drk_ma2", ""))
        angenommen = bool(d.get("angenommen", 1))
        if angenommen:
            self._radio_ja.setChecked(True)
        else:
            self._radio_nein.setChecked(True)
        self._grund_abgelehnt.setPlainText(d.get("grund_abgelehnt", ""))
        self._bemerkung.setPlainText(d.get("bemerkung", ""))

    def _on_accept(self):
        if self._verbrauch_positionen and not self._entnehmer.text().strip():
            QMessageBox.warning(
                self, "Pflichtfeld",
                "Bitte einen Entnehmer angeben.\nDieses Feld ist Pflicht wenn Verbrauchsmaterial eingetragen wurde."
            )
            self._entnehmer.setFocus()
            return
        self.accept()

    def get_daten(self) -> dict:
        return dict(
            datum            = self._datum.date().toString("dd.MM.yyyy"),
            uhrzeit          = self._uhrzeit.time().toString("HH:mm"),
            einsatzdauer     = self._einsatzdauer.value(),
            einsatzstichwort = self._stichwort.currentText().strip(),
            einsatzort       = self._einsatzort.text().strip(),
            einsatznr_drk    = self._einsatznr_drk.text().strip(),
            drk_ma1          = self._drk_ma1.text().strip(),
            drk_ma2          = self._drk_ma2.text().strip(),
            angenommen       = self._radio_ja.isChecked(),
            grund_abgelehnt  = self._grund_abgelehnt.toPlainText().strip(),
            bemerkung        = self._bemerkung.toPlainText().strip(),
            entnehmer        = self._entnehmer.text().strip(),
            verbrauch_liste  = list(self._verbrauch_positionen),
        )


# ──────────────────────────────────────────────────────────────────────────────
#  Word-Protokoll für Patienten-Station
# ──────────────────────────────────────────────────────────────────────────────

def export_patient_word(
    patient: dict,
    verbrauchsmaterial: list[dict],
    medikamente: list[dict] | None = None,
    ziel_pfad: str | None = None,
) -> str:
    """
    Erstellt ein Word-Protokoll (.docx) für einen einzelnen Stationspatienten.
    Gibt den gespeicherten Dateipfad zurück.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise ImportError(
            "python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen."
        ) from e

    os.makedirs(_PATIENTEN_PROTO_DIR, exist_ok=True)

    if not ziel_pfad:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        datum_str = patient.get("datum", "").replace(".", "")
        dateiname = f"Patientenprotokoll_{datum_str}_{ts}.docx"
        ziel_pfad = os.path.join(_PATIENTEN_PROTO_DIR, dateiname)

    doc = Document()

    # ── Seitenränder ────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(1.0)
        sec.bottom_margin = Cm(1.0)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)

    DRK_ROT  = RGBColor(0xC8, 0x10, 0x2E)
    DRK_BLAU = RGBColor(0x15, 0x65, 0xA8)

    def _set_cell_bg(cell, rgb_hex: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), rgb_hex)
        tcPr.append(shd)

    def _row(label: str, wert):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        lc = tbl.rows[0].cells[0]
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(0)
        run_l = lp.add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(8)
        _set_cell_bg(lc, "E8EFF8")
        vc = tbl.rows[0].cells[1]
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(0)
        vp.paragraph_format.space_after  = Pt(0)
        wert_str = str(wert).strip() if (wert is not None and str(wert).strip()) else "—"
        vp.add_run(wert_str).font.size = Pt(8)

    def _abschnitt(titel: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(f"  {titel}")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1565A8")
        pPr.append(shd)

    # ── Logo ─────────────────────────────────────────────────────────────────
    logo_pfad = os.path.join(BASE_DIR, "Daten", "Email", "Logo.jpg")
    if os.path.isfile(logo_pfad):
        doc.add_picture(logo_pfad, width=Cm(3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.paragraphs[-1].paragraph_format.space_before = Pt(0)
        doc.paragraphs[-1].paragraph_format.space_after  = Pt(0)

    # ── Titel ─────────────────────────────────────────────────────────────────
    titel_p = doc.add_heading("Patientenprotokoll", level=1)
    titel_p.runs[0].font.color.rgb = DRK_ROT
    titel_p.runs[0].font.size = Pt(13)
    titel_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titel_p.paragraph_format.space_before = Pt(2)
    titel_p.paragraph_format.space_after  = Pt(1)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(0)
    sub_run = sub_p.add_run("DRK Erste-Hilfe-Station – Flughafen Köln/Bonn (FKB)")
    sub_run.font.size = Pt(8)
    sub_run.font.color.rgb = DRK_BLAU

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(0)
    info_p.paragraph_format.space_after  = Pt(0)
    info_p.add_run(
        f"Datum: {patient.get('datum', '')}  |  Uhrzeit: {patient.get('uhrzeit', '')}  |  "
        f"Behandlungsdauer: {patient.get('behandlungsdauer', 0) or 0} Min."
    ).font.size = Pt(8)

    bg_jn = "Ja ⚠" if patient.get("arbeitsunfall") else "Nein"
    bg_p = doc.add_paragraph()
    bg_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bg_p.paragraph_format.space_before = Pt(0)
    bg_p.paragraph_format.space_after  = Pt(0)
    bg_run = bg_p.add_run(f"BG-Fall / Arbeitsunfall: {bg_jn}")
    bg_run.font.size = Pt(8)
    bg_run.bold = True
    if patient.get("arbeitsunfall"):
        bg_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    # ── 1. Patient ────────────────────────────────────────────────────────────
    _abschnitt("1 │ Patient")
    _row("Typ",                    patient.get("patient_typ", ""))
    _row("Abteilung",              patient.get("patient_abteilung", ""))
    _row("Name / Anonymisierung",  patient.get("patient_name", ""))
    _row("Alter",                  patient.get("patient_alter", "") or "")
    _row("Geschlecht",             patient.get("geschlecht", ""))

    # ── 2. Ereignis ───────────────────────────────────────────────────────────
    _abschnitt("2 │ Ereignis")
    _row("Was ist passiert?",    patient.get("hergang_was", ""))
    _row("Wie ist es passiert?", patient.get("hergang_wie", ""))
    _row("Ort / Bereich",        patient.get("unfall_ort", ""))

    # ── 3. Beschwerdebild ─────────────────────────────────────────────────────
    _abschnitt("3 │ Beschwerdebild")
    _row("Beschwerdeart", patient.get("beschwerde_art", ""))
    _row("Symptome",      patient.get("symptome", ""))

    # ── 4. ABCDE-Schema ───────────────────────────────────────────────────────
    _abschnitt("4 │ ABCDE-Schema")
    _row("A – Airway (Atemweg)",           patient.get("abcde_a", ""))
    _row("B – Breathing (Atmung)",         patient.get("abcde_b", ""))
    _row("C – Circulation (Kreislauf)",    patient.get("abcde_c", ""))
    _row("D – Disability (Neurologie)",    patient.get("abcde_d", ""))
    _row("E – Exposure (Umgebung/Körper)", patient.get("abcde_e", ""))

    # ── 5. Monitoring ─────────────────────────────────────────────────────────
    _abschnitt("5 │ Monitoring")
    mon_tbl = doc.add_table(rows=2, cols=4)
    mon_tbl.style = "Table Grid"
    for _i, _h in enumerate(["BZ (mg/dl)", "RR (mmHg)", "SpO2 (%)", "HF (bpm)"]):
        _c = mon_tbl.rows[0].cells[_i]
        _cp = _c.paragraphs[0]
        _cp.paragraph_format.space_before = Pt(0)
        _cp.paragraph_format.space_after  = Pt(0)
        _rc = _cp.add_run(_h)
        _rc.bold = True
        _rc.font.size = Pt(8)
        _set_cell_bg(_c, "E8EFF8")
    for _i, _v in enumerate([
        patient.get("monitoring_bz",   "") or "—",
        patient.get("monitoring_rr",   "") or "—",
        patient.get("monitoring_spo2", "") or "—",
        patient.get("monitoring_hf",   "") or "—",
    ]):
        _vp = mon_tbl.rows[1].cells[_i].paragraphs[0]
        _vp.paragraph_format.space_before = Pt(0)
        _vp.paragraph_format.space_after  = Pt(0)
        _vp.add_run(_v).font.size = Pt(8)

    # ── 6. Vorerkrankungen & Medikamente ──────────────────────────────────────
    _abschnitt("6 │ Vorerkrankungen & Medikamente des Patienten")
    _row("Vorerkrankungen",       patient.get("vorerkrankungen", ""))
    _row("Medikamente (Patient)", patient.get("medikamente_patient", ""))

    # ── 7. Behandlung ─────────────────────────────────────────────────────────
    _abschnitt("7 │ Behandlung durch DRK")
    _row("Diagnose / Einweisung", patient.get("diagnose", ""))
    _row("Maßnahmen",             patient.get("massnahmen", ""))
    # Medikamentengabe als Tabelle
    _med_liste = medikamente or []
    if _med_liste:
        _abschnitt_med = doc.add_paragraph()
        _abschnitt_med.paragraph_format.space_before = Pt(2)
        _abschnitt_med.paragraph_format.space_after  = Pt(0)
        _abschnitt_med.add_run("  Medikamente gegeben:").bold = True
        _abschnitt_med.runs[0].font.size = Pt(8)
        med_word_tbl = doc.add_table(rows=1 + len(_med_liste), cols=3)
        med_word_tbl.style = "Table Grid"
        for _i, _h in enumerate(["Medikament", "Dosis", "Applikation"]):
            _cw = med_word_tbl.rows[0].cells[_i]
            _cwp = _cw.paragraphs[0]
            _cwp.paragraph_format.space_before = Pt(0)
            _cwp.paragraph_format.space_after  = Pt(0)
            _rcw = _cwp.add_run(_h)
            _rcw.bold = True
            _rcw.font.size = Pt(8)
            _set_cell_bg(_cw, "E8EFF8")
        for _ri, _m in enumerate(_med_liste, start=1):
            for _ci, _val in enumerate([_m.get("medikament", ""), _m.get("dosis", "") or "—", _m.get("applikation", "") or "—"]):
                _rp = med_word_tbl.rows[_ri].cells[_ci].paragraphs[0]
                _rp.paragraph_format.space_before = Pt(0)
                _rp.paragraph_format.space_after  = Pt(0)
                _rp.add_run(_val).font.size = Pt(8)
    else:
        _row("Medikamentengabe", "Nein")

    # ── 8. Verbrauchsmaterial ─────────────────────────────────────────────────
    _abschnitt("8 │ Verbrauchsmaterial")
    if verbrauchsmaterial:
        vm_tbl = doc.add_table(rows=1 + len(verbrauchsmaterial), cols=3)
        vm_tbl.style = "Table Grid"
        for _i, _h in enumerate(["Material", "Menge", "Einheit"]):
            _c2 = vm_tbl.rows[0].cells[_i]
            _c2p = _c2.paragraphs[0]
            _c2p.paragraph_format.space_before = Pt(0)
            _c2p.paragraph_format.space_after  = Pt(0)
            _rc2 = _c2p.add_run(_h)
            _rc2.bold = True
            _rc2.font.size = Pt(8)
            _set_cell_bg(_c2, "E8EFF8")
        for _ri, _m in enumerate(verbrauchsmaterial, start=1):
            for _ci, _val in enumerate([_m.get("material", ""), str(_m.get("menge", "")), _m.get("einheit", "")]):
                _rp2 = vm_tbl.rows[_ri].cells[_ci].paragraphs[0]
                _rp2.paragraph_format.space_before = Pt(0)
                _rp2.paragraph_format.space_after  = Pt(0)
                _rp2.add_run(_val).font.size = Pt(8)
    else:
        p_vm = doc.add_paragraph("Kein Verbrauchsmaterial erfasst.")
        p_vm.paragraph_format.space_before = Pt(0)
        p_vm.paragraph_format.space_after  = Pt(0)
        p_vm.runs[0].font.size = Pt(8)

    # ── 9. Arbeitsunfall / BG-Fall ────────────────────────────────────────────
    if patient.get("arbeitsunfall"):
        _abschnitt("9 │ Arbeitsunfall / BG-Fall")
        _row("Details / Hergang", patient.get("arbeitsunfall_details", ""))

    # ── 10. Personal & Abschluss ──────────────────────────────────────────────
    _abschnitt("10 │ Personal & Abschluss")
    _row("DRK MA 1",          patient.get("drk_ma1", ""))
    _row("DRK MA 2",          patient.get("drk_ma2", ""))
    _row("Weitergeleitet an", patient.get("weitergeleitet", ""))

    # ── 11. Bemerkung ─────────────────────────────────────────────────────────
    if patient.get("bemerkung", "").strip():
        _abschnitt("11 │ Bemerkung")
        _row("Bemerkung", patient.get("bemerkung", ""))

    # ── Fußzeile ──────────────────────────────────────────────────────────────
    fuss_p = doc.add_paragraph()
    fuss_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fuss_p.paragraph_format.space_before = Pt(3)
    fuss_p.paragraph_format.space_after  = Pt(0)
    fuss_r = fuss_p.add_run(
        f"Erstellt am {datetime.now().strftime('%d.%m.%Y %H:%M')}  –  "
        "DRK-Kreisverband Köln e.V. | Erste-Hilfe-Station FKB"
    )
    fuss_r.font.size = Pt(7)
    fuss_r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(ziel_pfad)
    try:
        from functions.dokument_archiv import kopiere_ins_archiv
        kopiere_ins_archiv(ziel_pfad, "patienten_word")
    except Exception:
        pass
    return ziel_pfad


# ──────────────────────────────────────────────────────────────────────────────
#  Dialog: Patient erfassen / bearbeiten
# ──────────────────────────────────────────────────────────────────────────────

class _PatientenDialog(QDialog):
    """Erweitertes Formular zum Erfassen / Bearbeiten eines Stationspatienten."""

    _FIELD_STYLE = (
        "QLineEdit, QTextEdit, QTimeEdit, QDateEdit, QSpinBox, QComboBox {"
        "border:1px solid #ccc; border-radius:4px; padding:4px;"
        "font-size:12px; background:white;}"
    )
    _GROUP_STYLE = (
        "QGroupBox { font-weight:bold; font-size:12px; color:#1565a8;"
        "border:1px solid #c5d8f0; border-radius:6px;"
        "margin-top:6px; padding-top:10px; background:#f8fbff;}"
        "QGroupBox::title { subcontrol-origin:margin; left:12px;"
        "padding:0 4px; background:#f8fbff;}"
    )
    _DISABLED_STYLE = (
        "QLineEdit{border:1px solid #ddd;border-radius:4px;padding:4px;"
        "font-size:12px;background:#f5f5f5;color:#aaa;}"
    )
    _PATIENT_TYPEN = [
        "Fluggast", "Mitarbeiter (intern)", "Besucher",
        "Handwerker / Dienstleister", "Sonstiges",
    ]
    _BESCHWERDE_ARTEN = [
        "", "Verletzung / Trauma", "Internistisch / Erkrankung",
        "Neurologisch", "Allergisch / Anaphylaxie",
        "Bewusstlosigkeit / Synkope", "Kreislaufprobleme",
        "Atemnot / Respiratorisch", "Psychisch / Psychiatrisch", "Sonstiges",
    ]

    def __init__(self, daten: dict | None = None, parent=None):
        super().__init__(parent)
        self._edit_daten = daten
        self._verbrauchsmaterial_liste: list[dict] = []
        self._medikamente_liste: list[dict] = []
        self.setWindowTitle(
            "✏️  Patient bearbeiten" if daten
            else "🏥  Patienten erfassen – DRK Station"
        )
        self.setMinimumWidth(680)
        self.setMinimumHeight(750)
        self.resize(740, 880)
        self._build_ui()
        if daten:
            self._befuellen(daten)

    # ── UI-Aufbau ──────────────────────────────────────────────────────────────────────
    def _build_ui(self):
        main = QVBoxLayout(self)
        main.setContentsMargins(12, 12, 12, 8)
        main.setSpacing(8)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea{border:none;}")
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setSpacing(10)
        layout.setContentsMargins(4, 4, 4, 4)
        scroll.setWidget(container)
        main.addWidget(scroll, 1)

        layout.addWidget(self._build_grp_zeit())
        layout.addWidget(self._build_grp_patient())
        layout.addWidget(self._build_grp_ereignis())
        layout.addWidget(self._build_grp_beschwerde())
        layout.addWidget(self._build_grp_abcde())
        layout.addWidget(self._build_grp_monitoring())
        layout.addWidget(self._build_grp_vorerkrankungen())
        layout.addWidget(self._build_grp_behandlung())
        layout.addWidget(self._build_grp_medikamente())
        layout.addWidget(self._build_grp_material())
        layout.addWidget(self._build_grp_arbeitsunfall())
        layout.addWidget(self._build_grp_personal())
        layout.addWidget(self._build_grp_bemerkung())

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        btns.accepted.connect(self._validate)
        btns.rejected.connect(self.reject)
        main.addWidget(btns)

    # ── Hilfsmethoden ────────────────────────────────────────────────────────────────────
    def _grp(self, titel: str) -> tuple[QGroupBox, QFormLayout]:
        g = QGroupBox(titel)
        g.setStyleSheet(self._GROUP_STYLE)
        f = QFormLayout(g)
        f.setSpacing(8)
        return g, f

    def _le(self, placeholder: str = "") -> QLineEdit:
        w = QLineEdit()
        w.setPlaceholderText(placeholder)
        w.setStyleSheet(self._FIELD_STYLE)
        return w

    def _te(self, placeholder: str = "", h: int = 65) -> QTextEdit:
        w = QTextEdit()
        w.setPlaceholderText(placeholder)
        w.setFixedHeight(h)
        w.setStyleSheet(self._FIELD_STYLE)
        return w

    # ── Abschnitte ───────────────────────────────────────────────────────────────────
    def _build_grp_zeit(self) -> QGroupBox:
        g, f = self._grp("⏱️  Zeitangaben")
        self._datum = QDateEdit()
        self._datum.setCalendarPopup(True)
        self._datum.setDisplayFormat("dd.MM.yyyy")
        self._datum.setDate(QDate.currentDate())
        self._datum.setStyleSheet(self._FIELD_STYLE)
        f.addRow("Datum:", self._datum)
        self._uhrzeit = QTimeEdit()
        self._uhrzeit.setDisplayFormat("HH:mm")
        self._uhrzeit.setTime(QTime.currentTime())
        self._uhrzeit.setStyleSheet(self._FIELD_STYLE)
        f.addRow("Uhrzeit:", self._uhrzeit)
        self._behandlungsdauer = QSpinBox()
        self._behandlungsdauer.setRange(0, 999)
        self._behandlungsdauer.setSuffix("  min")
        self._behandlungsdauer.setValue(0)
        self._behandlungsdauer.setStyleSheet(self._FIELD_STYLE)
        f.addRow("Behandlungsdauer:", self._behandlungsdauer)
        return g

    def _build_grp_patient(self) -> QGroupBox:
        g, f = self._grp("👤  Patienteninformationen")
        self._patient_typ = QComboBox()
        self._patient_typ.addItems(self._PATIENT_TYPEN)
        self._patient_typ.setEditable(True)
        self._patient_typ.setStyleSheet(self._FIELD_STYLE)
        self._patient_typ.currentTextChanged.connect(self._on_typ_changed)
        f.addRow("Personentyp *:", self._patient_typ)
        self._patient_abteilung = self._le("Abteilung / Firma / Bereich")
        self._patient_abteilung.setEnabled(False)
        self._patient_abteilung.setStyleSheet(self._DISABLED_STYLE)
        f.addRow("Abteilung:", self._patient_abteilung)
        self._patient_name = self._le("")
        self._patient_name.setReadOnly(True)
        self._patient_name.setEnabled(False)
        self._patient_name.setPlaceholderText("✏ handschriftlich eintragen")
        self._patient_name.setToolTip("Aus Datenschutzgründen nicht digital erfasst – bitte im Ausdruck handschriftlich eintragen")
        self._patient_name.setStyleSheet(self._DISABLED_STYLE)
        f.addRow("Name / Geburtsdatum:", self._patient_name)
        self._alter = QSpinBox()
        self._alter.setRange(0, 150)
        self._alter.setSuffix("  Jahre")
        self._alter.setValue(0)
        self._alter.setEnabled(False)
        self._alter.setStyleSheet(self._DISABLED_STYLE)
        self._alter.setToolTip("Aus Datenschutzgründen nicht digital erfasst – bitte im Ausdruck handschriftlich eintragen")
        f.addRow("Alter:", self._alter)
        _hinweis = QLabel("ℹ  Name und Alter werden nicht digital gespeichert – bitte handschriftlich im Protokollausdruck eintragen.")
        _hinweis.setWordWrap(True)
        _hinweis.setStyleSheet("color:#b06000;font-size:10px;font-style:italic;border:none;padding:2px 0;")
        f.addRow("", _hinweis)
        self._geschlecht = QComboBox()
        self._geschlecht.addItems(["", "männlich", "weiblich", "divers"])
        self._geschlecht.setStyleSheet(self._FIELD_STYLE)
        f.addRow("Geschlecht:", self._geschlecht)
        return g

    def _on_typ_changed(self, text: str):
        is_intern = "mitarbeiter" in text.lower() or "intern" in text.lower()
        self._patient_abteilung.setEnabled(is_intern)
        self._patient_abteilung.setStyleSheet(
            self._FIELD_STYLE if is_intern else self._DISABLED_STYLE
        )

    def _build_grp_ereignis(self) -> QGroupBox:
        g, f = self._grp("📋  Ereignis")
        self._hergang_was = self._te(
            "Was ist passiert? (Beschreibung des Ereignisses)", 70
        )
        f.addRow("Was ist passiert *:", self._hergang_was)
        self._hergang_wie = self._te(
            "Wie ist es passiert? (Unfallhergang, Entstehungsweise)", 60
        )
        f.addRow("Wie ist es passiert:", self._hergang_wie)
        self._unfall_ort = self._le(
            "z.B. Terminal 1 / Gate B42 / Gepäckband 7 / Außenbereich …"
        )
        f.addRow("Wo ist es passiert:", self._unfall_ort)
        return g

    def _build_grp_beschwerde(self) -> QGroupBox:
        g, f = self._grp("🩺  Beschwerdebild / Verletzung")
        self._beschwerde_art = QComboBox()
        self._beschwerde_art.addItems(self._BESCHWERDE_ARTEN)
        self._beschwerde_art.setEditable(True)
        self._beschwerde_art.setStyleSheet(self._FIELD_STYLE)
        f.addRow("Art der Beschwerde:", self._beschwerde_art)
        self._symptome = self._te("Symptome, Beschwerden und Befunde …")
        f.addRow("Symptome / Befunde:", self._symptome)
        return g

    def _build_grp_abcde(self) -> QGroupBox:
        g = QGroupBox("🔬  ABCDE-Schema (Notfallsystematik Rettungsdienst)")
        g.setStyleSheet(self._GROUP_STYLE)
        grid = QGridLayout(g)
        grid.setSpacing(6)
        grid.setContentsMargins(10, 18, 10, 10)
        for col, txt in [(1, "Kriterium"), (2, "Befund / Beurteilung")]:
            lbl = QLabel(txt)
            lbl.setStyleSheet("font-weight:bold; font-size:11px; color:#555;")
            grid.addWidget(lbl, 0, col)
        abcde_defs = [
            ("A", "Airway – Atemweg",        "frei / verlegt / Stridor / Fremdkörper …",           "_abcde_a"),
            ("B", "Breathing – Atmung",       "unauffällig / Tachypnoe / Dyspnoe / SpO₂ …",    "_abcde_b"),
            ("C", "Circulation – Kreislauf",  "Puls tastbar / RR … / Tachykardie / Schock …",     "_abcde_c"),
            ("D", "Disability – Neurologie",  "wach / GCS … / orientiert / Pupillen … / AVPU …", "_abcde_d"),
            ("E", "Exposure – Untersuchung",  "Wunden / Traumazeichen / Hypothermie / Hautturgur …",   "_abcde_e"),
        ]
        for r, (letter, label_txt, placeholder, attr) in enumerate(abcde_defs, start=1):
            badge = QLabel(letter)
            badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
            badge.setFixedSize(26, 26)
            badge.setStyleSheet(
                "background:#1565a8;color:white;border-radius:13px;"
                "font-weight:bold;font-size:13px;"
            )
            lbl = QLabel(label_txt)
            lbl.setStyleSheet("font-size:11px;")
            edit = self._le(placeholder)
            setattr(self, attr, edit)
            grid.addWidget(badge, r, 0)
            grid.addWidget(lbl,   r, 1)
            grid.addWidget(edit,  r, 2)
        grid.setColumnStretch(2, 1)
        return g

    def _build_grp_monitoring(self) -> QGroupBox:
        g = QGroupBox("📊  Monitoring-Werte")
        g.setStyleSheet(self._GROUP_STYLE)
        grid = QGridLayout(g)
        grid.setSpacing(8)
        grid.setContentsMargins(10, 18, 10, 10)
        mon_defs = [
            ("BZ (mg/dl):", "_mon_bz",   "z.B. 95"),
            ("RR (mmHg):",  "_mon_rr",   "z.B. 120/80"),
            ("SpO₂ (%):",  "_mon_spo2", "z.B. 98"),
            ("HF (bpm):",   "_mon_hf",   "z.B. 78"),
        ]
        for i, (label, attr, placeholder) in enumerate(mon_defs):
            r, c = divmod(i, 2)
            lbl = QLabel(label)
            lbl.setStyleSheet("font-size:12px;")
            edit = self._le(placeholder)
            edit.setFixedWidth(140)
            setattr(self, attr, edit)
            grid.addWidget(lbl,  r, c * 2)
            grid.addWidget(edit, r, c * 2 + 1)
        return g

    def _build_grp_vorerkrankungen(self) -> QGroupBox:
        g, f = self._grp("📜  Vorerkrankungen & Medikamente des Patienten")
        self._vorerkrankungen = self._te(
            "Bekannte Vorerkrankungen, Allergien …", 60
        )
        f.addRow("Vorerkrankungen:", self._vorerkrankungen)
        self._medikamente_patient = self._te(
            "Dauermedikamente des Patienten …", 55
        )
        f.addRow("Medikamente (Patient):", self._medikamente_patient)
        return g

    def _build_grp_behandlung(self) -> QGroupBox:
        g, f = self._grp("🏥  Art und Umfang der Behandlung")
        self._massnahmen = self._te(
            "Art und Umfang der durchgeführten Maßnahmen …", 80
        )
        f.addRow("Maßnahmen / Behandlung *:", self._massnahmen)
        self._diagnose = self._te("Diagnose / Verdachtsdiagnose …", 55)
        f.addRow("Diagnose:", self._diagnose)
        return g

    def _build_grp_medikamente(self) -> QGroupBox:
        g = QGroupBox("💊  Medikamentengabe durch DRK")
        g.setStyleSheet(self._GROUP_STYLE)
        med_layout = QVBoxLayout(g)
        med_layout.setSpacing(8)
        btn_add = QPushButton("➕  Medikament hinzufügen")
        btn_add.setFixedHeight(32)
        btn_add.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_add.setStyleSheet(
            "QPushButton{background:#1565a8;color:white;border:none;"
            "border-radius:4px;padding:4px 14px;font-size:12px;}"
            "QPushButton:hover{background:#0d47a1;}"
        )
        btn_add.clicked.connect(self._medikament_hinzufuegen)
        med_layout.addWidget(btn_add, 0, Qt.AlignmentFlag.AlignLeft)
        self._medikament_table = QTableWidget()
        self._medikament_table.setColumnCount(4)
        self._medikament_table.setHorizontalHeaderLabels(["Medikament", "Dosis", "Applikation", ""])
        mh = self._medikament_table.horizontalHeader()
        mh.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        mh.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        mh.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        mh.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        self._medikament_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._medikament_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._medikament_table.setMaximumHeight(150)
        self._medikament_table.verticalHeader().setVisible(False)
        self._medikament_table.setStyleSheet("font-size:12px;")
        med_layout.addWidget(self._medikament_table)
        return g

    def _build_grp_material(self) -> QGroupBox:
        g = QGroupBox("📦  Verbrauchsmaterial")
        g.setStyleSheet(self._GROUP_STYLE)
        mat_layout = QVBoxLayout(g)
        mat_layout.setSpacing(8)
        btn_add = QPushButton("➕  Material hinzufügen")
        btn_add.setFixedHeight(32)
        btn_add.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_add.setStyleSheet(
            "QPushButton{background:#107e3e;color:white;border:none;"
            "border-radius:4px;padding:4px 14px;font-size:12px;}"
            "QPushButton:hover{background:#0a5c2e;}"
        )
        btn_add.clicked.connect(self._material_hinzufuegen)
        mat_layout.addWidget(btn_add, 0, Qt.AlignmentFlag.AlignLeft)
        self._material_table = QTableWidget()
        self._material_table.setColumnCount(4)
        self._material_table.setHorizontalHeaderLabels(["Material", "Menge", "Einheit", ""])
        hh = self._material_table.horizontalHeader()
        hh.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        hh.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        hh.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        hh.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        self._material_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._material_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._material_table.setMaximumHeight(150)
        self._material_table.verticalHeader().setVisible(False)
        self._material_table.setStyleSheet("font-size:12px;")
        mat_layout.addWidget(self._material_table)
        return g

    def _build_grp_arbeitsunfall(self) -> QGroupBox:
        g = QGroupBox("⚠️  Arbeitsunfall / Berufsgenossenschaft")
        g.setStyleSheet(self._GROUP_STYLE)
        lay = QVBoxLayout(g)
        lay.setSpacing(8)
        self._arbeitsunfall_cb = QCheckBox(
            "Es handelt sich um einen Arbeitsunfall im Sinne der BG"
        )
        self._arbeitsunfall_cb.setStyleSheet(
            "QCheckBox{font-size:12px;font-weight:bold;color:#b00000;}"
        )
        lay.addWidget(self._arbeitsunfall_cb)
        fl = QFormLayout()
        self._arbeitsunfall_details = self._le(
            "Arbeitgeber / BG / Detailangaben zum Arbeitsunfall …"
        )
        self._arbeitsunfall_details.setEnabled(False)
        self._arbeitsunfall_cb.toggled.connect(self._arbeitsunfall_details.setEnabled)
        fl.addRow("Details (BG / Arbeitgeber):", self._arbeitsunfall_details)
        lay.addLayout(fl)
        return g

    def _build_grp_personal(self) -> QGroupBox:
        g, f = self._grp("👥  Personal & Weiterleitung")
        self._drk_ma1 = self._le("Name Mitarbeiter 1")
        f.addRow("DRK MA 1 *:", self._drk_ma1)
        self._drk_ma2 = self._le("Name Mitarbeiter 2 (optional)")
        f.addRow("DRK MA 2:", self._drk_ma2)
        self._weitergeleitet = self._le(
            "z.B. 'Krankenhaus XY', 'Hausarzt', '112 gerufen', 'keine' …"
        )
        f.addRow("Weitergeleitet an:", self._weitergeleitet)
        return g

    def _build_grp_bemerkung(self) -> QGroupBox:
        g = QGroupBox("📝  Bemerkung")
        g.setStyleSheet(self._GROUP_STYLE)
        lay = QVBoxLayout(g)
        self._bemerkung = self._te("Zusätzliche Hinweise …", 55)
        lay.addWidget(self._bemerkung)
        return g

    # ── Medikamente ───────────────────────────────────────────────────────────────────────
    def _medikament_hinzufuegen(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("Medikament hinzufügen")
        dlg.setMinimumWidth(400)
        layout = QVBoxLayout(dlg)
        fl = QFormLayout()
        med_edit = QLineEdit()
        med_edit.setPlaceholderText("z.B. Aspirin, Glucose, Sauerstoff …")
        dosis_edit = QLineEdit()
        dosis_edit.setPlaceholderText("z.B. 500 mg, 4 l/min, 40%")
        applikation_combo = QComboBox()
        applikation_combo.addItems(["", "oral", "sublingual", "inhalativ", "intravenös (IV)", "subkutan", "intramuskulär", "intranasal", "topisch"])
        applikation_combo.setEditable(True)
        fl.addRow("Medikament:", med_edit)
        fl.addRow("Dosis:", dosis_edit)
        fl.addRow("Applikation:", applikation_combo)
        layout.addLayout(fl)
        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        btns.accepted.connect(lambda: dlg.accept() if med_edit.text().strip() else None)
        btns.rejected.connect(dlg.reject)
        layout.addWidget(btns)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            medikament = med_edit.text().strip()
            if medikament:
                self._medikamente_liste.append({
                    "medikament":   medikament,
                    "dosis":        dosis_edit.text().strip(),
                    "applikation":  applikation_combo.currentText().strip(),
                })
                self._aktualisiere_medikament_tabelle()

    def _aktualisiere_medikament_tabelle(self):
        self._medikament_table.setRowCount(len(self._medikamente_liste))
        for row, med in enumerate(self._medikamente_liste):
            self._medikament_table.setItem(row, 0, QTableWidgetItem(med.get("medikament", "")))
            self._medikament_table.setItem(row, 1, QTableWidgetItem(med.get("dosis", "")))
            self._medikament_table.setItem(row, 2, QTableWidgetItem(med.get("applikation", "")))
            btn_del = QPushButton("🗑")
            btn_del.setFixedSize(28, 28)
            btn_del.setCursor(Qt.CursorShape.PointingHandCursor)
            btn_del.setStyleSheet(
                "QPushButton{background:#ffcccc;border:none;border-radius:3px;}"
                "QPushButton:hover{background:#ff9999;}"
            )
            btn_del.clicked.connect(lambda checked, r=row: self._medikament_entfernen(r))
            self._medikament_table.setCellWidget(row, 3, btn_del)

    def _medikament_entfernen(self, row: int):
        if 0 <= row < len(self._medikamente_liste):
            self._medikamente_liste.pop(row)
            self._aktualisiere_medikament_tabelle()

    # ── Verbrauchsmaterial ──────────────────────────────────────────────────────────────────
    def _material_hinzufuegen(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("Material hinzufügen")
        dlg.setMinimumWidth(400)
        layout = QVBoxLayout(dlg)
        fl = QFormLayout()
        material_edit = QLineEdit()
        material_edit.setPlaceholderText("z.B. Pflaster, Kompresse, Handschuhe …")
        menge_spin = QSpinBox()
        menge_spin.setRange(1, 999)
        menge_spin.setValue(1)
        einheit_combo = QComboBox()
        einheit_combo.addItems(["Stk", "Paar", "Pkg", "ml"])
        einheit_combo.setEditable(True)
        fl.addRow("Material:", material_edit)
        fl.addRow("Menge:", menge_spin)
        fl.addRow("Einheit:", einheit_combo)
        layout.addLayout(fl)
        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        btns.accepted.connect(lambda: dlg.accept() if material_edit.text().strip() else None)
        btns.rejected.connect(dlg.reject)
        layout.addWidget(btns)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            material = material_edit.text().strip()
            if material:
                self._verbrauchsmaterial_liste.append({
                    "material": material,
                    "menge": menge_spin.value(),
                    "einheit": einheit_combo.currentText()
                })
                self._aktualisiere_material_tabelle()

    def _aktualisiere_material_tabelle(self):
        self._material_table.setRowCount(len(self._verbrauchsmaterial_liste))
        for row, mat in enumerate(self._verbrauchsmaterial_liste):
            self._material_table.setItem(row, 0, QTableWidgetItem(mat["material"]))
            self._material_table.setItem(row, 1, QTableWidgetItem(str(mat["menge"])))
            self._material_table.setItem(row, 2, QTableWidgetItem(mat["einheit"]))
            btn_del = QPushButton("🗑")
            btn_del.setFixedSize(28, 28)
            btn_del.setCursor(Qt.CursorShape.PointingHandCursor)
            btn_del.setStyleSheet(
                "QPushButton{background:#ffcccc;border:none;border-radius:3px;}"
                "QPushButton:hover{background:#ff9999;}"
            )
            btn_del.clicked.connect(lambda checked, r=row: self._material_entfernen(r))
            self._material_table.setCellWidget(row, 3, btn_del)

    def _material_entfernen(self, row: int):
        if 0 <= row < len(self._verbrauchsmaterial_liste):
            self._verbrauchsmaterial_liste.pop(row)
            self._aktualisiere_material_tabelle()

    # ── Daten befüllen ───────────────────────────────────────────────────────────────────────
    def _befuellen(self, d: dict):
        try:
            datum_obj = QDate.fromString(d.get("datum", ""), "dd.MM.yyyy")
            if datum_obj.isValid():
                self._datum.setDate(datum_obj)
        except Exception:
            pass
        try:
            zeit_obj = QTime.fromString(d.get("uhrzeit", ""), "HH:mm")
            if zeit_obj.isValid():
                self._uhrzeit.setTime(zeit_obj)
        except Exception:
            pass
        self._behandlungsdauer.setValue(d.get("behandlungsdauer", 0) or 0)
        # Patient
        typ = d.get("patient_typ", "")
        if typ:
            idx = self._patient_typ.findText(typ)
            if idx >= 0:
                self._patient_typ.setCurrentIndex(idx)
            else:
                self._patient_typ.setCurrentText(typ)
        self._patient_abteilung.setText(d.get("patient_abteilung", ""))
        self._on_typ_changed(self._patient_typ.currentText())
        self._patient_name.setText(d.get("patient_name", ""))
        self._alter.setValue(d.get("patient_alter", 0) or 0)
        idx = self._geschlecht.findText(d.get("geschlecht", ""))
        if idx >= 0:
            self._geschlecht.setCurrentIndex(idx)
        # Ereignis
        self._hergang_was.setPlainText(d.get("hergang_was", ""))
        self._hergang_wie.setPlainText(d.get("hergang_wie", ""))
        self._unfall_ort.setText(d.get("unfall_ort", ""))
        # Beschwerdebild
        ba = d.get("beschwerde_art", "")
        if ba:
            idx2 = self._beschwerde_art.findText(ba)
            if idx2 >= 0:
                self._beschwerde_art.setCurrentIndex(idx2)
            else:
                self._beschwerde_art.setCurrentText(ba)
        self._symptome.setPlainText(d.get("symptome", ""))
        # ABCDE
        self._abcde_a.setText(d.get("abcde_a", ""))
        self._abcde_b.setText(d.get("abcde_b", ""))
        self._abcde_c.setText(d.get("abcde_c", ""))
        self._abcde_d.setText(d.get("abcde_d", ""))
        self._abcde_e.setText(d.get("abcde_e", ""))
        # Monitoring
        self._mon_bz.setText(d.get("monitoring_bz", ""))
        self._mon_rr.setText(d.get("monitoring_rr", ""))
        self._mon_spo2.setText(d.get("monitoring_spo2", ""))
        self._mon_hf.setText(d.get("monitoring_hf", ""))
        # Vorerkrankungen
        self._vorerkrankungen.setPlainText(d.get("vorerkrankungen", ""))
        self._medikamente_patient.setPlainText(d.get("medikamente_patient", ""))
        # Behandlung
        self._massnahmen.setPlainText(d.get("massnahmen", ""))
        self._diagnose.setPlainText(d.get("diagnose", ""))
        # Medikamente
        if "id" in d:
            self._medikamente_liste = lade_medikamente(d["id"])
            self._aktualisiere_medikament_tabelle()
        # Arbeitsunfall
        self._arbeitsunfall_cb.setChecked(bool(d.get("arbeitsunfall", 0)))
        self._arbeitsunfall_details.setText(d.get("arbeitsunfall_details", ""))
        # Personal
        self._drk_ma1.setText(d.get("drk_ma1", ""))
        self._drk_ma2.setText(d.get("drk_ma2", ""))
        self._weitergeleitet.setText(d.get("weitergeleitet", ""))
        self._bemerkung.setPlainText(d.get("bemerkung", ""))
        # Verbrauchsmaterial
        if "id" in d:
            self._verbrauchsmaterial_liste = lade_verbrauchsmaterial(d["id"])
            self._aktualisiere_material_tabelle()

    # ── Validierung ─────────────────────────────────────────────────────────────────────────
    def _validate(self):
        self.accept()

    # ── Daten auslesen ──────────────────────────────────────────────────────────────────────
    def get_data(self) -> tuple[dict, list[dict]]:
        daten = {
            "datum":                   self._datum.date().toString("dd.MM.yyyy"),
            "uhrzeit":                 self._uhrzeit.time().toString("HH:mm"),
            "behandlungsdauer":        self._behandlungsdauer.value(),
            "patient_typ":             self._patient_typ.currentText(),
            "patient_abteilung":       self._patient_abteilung.text().strip(),
            "patient_name":            self._patient_name.text().strip(),
            "patient_alter":           self._alter.value(),
            "geschlecht":              self._geschlecht.currentText(),
            "hergang_was":             self._hergang_was.toPlainText().strip(),
            "hergang_wie":             self._hergang_wie.toPlainText().strip(),
            "unfall_ort":              self._unfall_ort.text().strip(),
            "beschwerde_art":          self._beschwerde_art.currentText(),
            "symptome":                self._symptome.toPlainText().strip(),
            "abcde_a":                 self._abcde_a.text().strip(),
            "abcde_b":                 self._abcde_b.text().strip(),
            "abcde_c":                 self._abcde_c.text().strip(),
            "abcde_d":                 self._abcde_d.text().strip(),
            "abcde_e":                 self._abcde_e.text().strip(),
            "monitoring_bz":           self._mon_bz.text().strip(),
            "monitoring_rr":           self._mon_rr.text().strip(),
            "monitoring_spo2":         self._mon_spo2.text().strip(),
            "monitoring_hf":           self._mon_hf.text().strip(),
            "vorerkrankungen":         self._vorerkrankungen.toPlainText().strip(),
            "medikamente_patient":     self._medikamente_patient.toPlainText().strip(),
            "massnahmen":              self._massnahmen.toPlainText().strip(),
            "diagnose":                self._diagnose.toPlainText().strip(),
            "medikamente_gegeben":     1 if self._medikamente_liste else 0,
            "medikamente_gegeben_was": ", ".join(m["medikament"] for m in self._medikamente_liste),
            "_medikamente":            self._medikamente_liste,
            "arbeitsunfall":           1 if self._arbeitsunfall_cb.isChecked() else 0,
            "arbeitsunfall_details":   self._arbeitsunfall_details.text().strip(),
            "drk_ma1":                 self._drk_ma1.text().strip(),
            "drk_ma2":                 self._drk_ma2.text().strip(),
            "weitergeleitet":          self._weitergeleitet.text().strip(),
            "bemerkung":               self._bemerkung.toPlainText().strip(),
        }
        return daten, self._verbrauchsmaterial_liste


# ──────────────────────────────────────────────────────────────────────────────
#  Mail-Dialog für Patienten-Protokoll
# ──────────────────────────────────────────────────────────────────────────────

class _PatientenMailDialog(QDialog):
    """Mail-Dialog für ein einzelnes Patienten-Protokoll (.docx)."""

    _FIELD_STYLE = (
        "QLineEdit, QTextEdit { border:1px solid #ccc; border-radius:4px;"
        "padding:4px; font-size:12px; background:white;}"
    )

    def __init__(
        self,
        word_pfad: str,
        patient_name: str,
        datum: str,
        parent=None,
    ):
        super().__init__(parent)
        self.setWindowTitle("📧  Patientenprotokoll per E-Mail senden")
        self.setMinimumWidth(500)
        self.resize(540, 420)
        layout = QVBoxLayout(self)
        fl = QFormLayout()
        fl.setSpacing(8)

        self._empfaenger = QLineEdit()
        self._empfaenger.setPlaceholderText("empfaenger@drk-koeln.de")
        self._empfaenger.setStyleSheet(self._FIELD_STYLE)
        fl.addRow("Empfänger:", self._empfaenger)

        self._betreff = QLineEdit()
        pat_lbl = patient_name.strip() if patient_name.strip() else "Patient"
        self._betreff.setText(f"Patientenprotokoll FKB – {pat_lbl} ({datum})")
        self._betreff.setStyleSheet(self._FIELD_STYLE)
        fl.addRow("Betreff:", self._betreff)

        layout.addLayout(fl)
        layout.addWidget(QLabel("Nachrichtentext:"))

        self._body = QTextEdit()
        self._body.setStyleSheet(self._FIELD_STYLE)
        self._body.setPlainText(
            f"Hallo,\n\n"
            f"anbei das Patientenprotokoll vom {datum}."
            f"{(' (Patient: ' + patient_name.strip() + ')') if patient_name.strip() else ''}\n\n"
            f"Das Protokoll ist als Word-Dokument angehängt.\n\n"
            f"Mit freundlichen Grüßen\n"
            f"DRK-Kreisverband Köln e.V."
        )
        self._body.setMinimumHeight(140)
        layout.addWidget(self._body, 1)

        anhang_lbl = QLabel(f"📎  Anhang: {os.path.basename(word_pfad)}")
        anhang_lbl.setStyleSheet(
            "background:#e3f2fd; color:#154360; border-radius:4px;"
            "padding:6px 10px; font-size:11px;"
        )
        layout.addWidget(anhang_lbl)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("📧  Outlook-Entwurf öffnen")
        btns.accepted.connect(self._on_accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def _on_accept(self):
        if not self._empfaenger.text().strip():
            QMessageBox.warning(self, "Pflichtfeld", "Bitte einen Empfänger angeben.")
            return
        self.accept()

    def get_daten(self) -> tuple[str, str, str]:
        return (
            self._empfaenger.text().strip(),
            self._betreff.text().strip(),
            self._body.toPlainText().strip(),
        )


# ──────────────────────────────────────────────────────────────────────────────
#  Patienten DRK Station Tab
# ──────────────────────────────────────────────────────────────────────────────

class _PatientenTab(QWidget):
    """Tab 'Patienten DRK Station' – Behandlung von Patienten auf der Station."""

    _COLS = [
        "Nr.", "Datum", "Uhrzeit", "Dauer\n(min)",
        "Typ", "Patient", "Alter", "Beschwerde",
        "Ort", "Maßnahmen", "DRK MA 1", "Weitergeleitet", "BG-Fall",
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self._eintraege: list[dict] = []
        self._build_ui()
        self.refresh()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        titel_lbl = QLabel("🏥  Patienten DRK Station")
        titel_lbl.setFont(QFont("Arial", 15, QFont.Weight.Bold))
        titel_lbl.setStyleSheet(f"color:{FIORI_TEXT}; padding: 4px 0;")
        layout.addWidget(titel_lbl)

        hinweis_lbl = QLabel("Behandlung von Patienten direkt auf der DRK-Station am Flughafen")
        hinweis_lbl.setStyleSheet("color:#666; font-size:11px; font-style:italic;")
        layout.addWidget(hinweis_lbl)

        # ── Filter-Leiste ─────────────────────────────────────────────────
        filter_frame = QFrame()
        filter_frame.setStyleSheet(
            "QFrame{background:#f8f9fa;border:1px solid #ddd;border-radius:4px;padding:4px;}"
        )
        fl = QHBoxLayout(filter_frame)
        fl.setContentsMargins(8, 6, 8, 6)
        fl.setSpacing(10)

        fl.addWidget(QLabel("Jahr:"))
        self._combo_jahr = QComboBox()
        self._combo_jahr.setFixedWidth(80)
        self._combo_jahr.addItem("Alle", None)
        self._combo_jahr.currentIndexChanged.connect(self._filter_changed)
        fl.addWidget(self._combo_jahr)

        fl.addWidget(QLabel("Monat:"))
        self._combo_monat = QComboBox()
        self._combo_monat.setFixedWidth(110)
        for i, m in enumerate([
            "Alle", "Januar", "Februar", "März", "April", "Mai", "Juni",
            "Juli", "August", "September", "Oktober", "November", "Dezember"
        ]):
            self._combo_monat.addItem(m, None if i == 0 else i)
        self._combo_monat.currentIndexChanged.connect(self._filter_changed)
        fl.addWidget(self._combo_monat)

        fl.addWidget(QLabel("Suche:"))
        self._suche = QLineEdit()
        self._suche.setPlaceholderText("Patient, Symptome, Mitarbeiter …")
        self._suche.setMinimumWidth(200)
        self._suche.textChanged.connect(self._filter_changed)
        fl.addWidget(self._suche, 1)

        btn_reset = _btn_light("✖ Zurücksetzen")
        btn_reset.setFixedHeight(28)
        btn_reset.clicked.connect(self._filter_reset)
        fl.addWidget(btn_reset)
        layout.addWidget(filter_frame)

        # ── Aktions-Buttons ────────────────────────────────────────────────
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        self._btn_neu = _btn("＋  Neuer Patient", "#107e3e", "#0a5c2e")
        self._btn_neu.setToolTip("Neuen Patienten erfassen")
        self._btn_neu.clicked.connect(self._neu)
        btn_row.addWidget(self._btn_neu)

        self._btn_bearbeiten = _btn_light("✏  Bearbeiten")
        self._btn_bearbeiten.setEnabled(False)
        self._btn_bearbeiten.setToolTip("Ausgewählten Patienten bearbeiten")
        self._btn_bearbeiten.clicked.connect(self._bearbeiten)
        btn_row.addWidget(self._btn_bearbeiten)

        self._btn_loeschen = _btn_light("🗑  Löschen")
        self._btn_loeschen.setEnabled(False)
        self._btn_loeschen.setStyleSheet(
            "QPushButton{background:#eee;color:#333;border:none;"
            "border-radius:4px;padding:4px 14px;font-size:12px;}"
            "QPushButton:hover{background:#ffcccc;color:#a00;}"
            "QPushButton:disabled{background:#f5f5f5;color:#bbb;}"
        )
        self._btn_loeschen.clicked.connect(self._loeschen)
        btn_row.addWidget(self._btn_loeschen)

        self._btn_material = _btn_light("📦  Verbrauchsmaterial anzeigen")
        self._btn_material.setEnabled(False)
        self._btn_material.setToolTip("Verbrauchsmaterial für ausgewählten Patienten anzeigen")
        self._btn_material.clicked.connect(self._material_anzeigen)
        btn_row.addWidget(self._btn_material)

        self._btn_word = _btn_light("📄  Word-Protokoll")
        self._btn_word.setEnabled(False)
        self._btn_word.setToolTip("Protokoll des ausgewählten Patienten als Word-Datei erstellen und öffnen")
        self._btn_word.clicked.connect(self._word_protokoll)
        btn_row.addWidget(self._btn_word)

        self._btn_mail = _btn("📧  Per E-Mail senden", "#6a1b9a", "#4a148c")
        self._btn_mail.setEnabled(False)
        self._btn_mail.setToolTip("Word-Protokoll erstellen und als Outlook-Entwurf versenden")
        self._btn_mail.clicked.connect(self._mail_protokoll)
        btn_row.addWidget(self._btn_mail)

        # Trennstrich
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.VLine)
        sep.setStyleSheet("color:#ccc;")
        btn_row.addWidget(sep)

        self._btn_excel = _btn("📈  Excel exportieren", "#1565a8", "#0d47a1")
        self._btn_excel.setToolTip(
            f"Aktuelle Ansicht als Excel speichern\n"
            f"Standard-Ordner: Daten/Pat. Station"
        )
        self._btn_excel.clicked.connect(self._excel_exportieren)
        btn_row.addWidget(self._btn_excel)

        self._btn_excel_wo = _btn_light("📂  Speicherort wählen")
        self._btn_excel_wo.setToolTip("Excel exportieren und Speicherort manuell wählen")
        self._btn_excel_wo.clicked.connect(lambda: self._excel_exportieren(ask_path=True))
        btn_row.addWidget(self._btn_excel_wo)

        btn_row.addStretch()
        self._treffer_lbl = QLabel()
        self._treffer_lbl.setStyleSheet("color:#666; font-size:11px;")
        btn_row.addWidget(self._treffer_lbl)
        layout.addLayout(btn_row)

        # ── Tabelle ────────────────────────────────────────────────────────
        self._table = QTableWidget()
        self._table.setColumnCount(len(self._COLS))
        self._table.setHorizontalHeaderLabels(self._COLS)
        hh = self._table.horizontalHeader()
        for _c in (0, 1, 2, 3, 4, 5, 6, 8, 10, 11, 12):
            hh.setSectionResizeMode(_c, QHeaderView.ResizeMode.ResizeToContents)
        hh.setSectionResizeMode(7,  QHeaderView.ResizeMode.Stretch)
        hh.setSectionResizeMode(9,  QHeaderView.ResizeMode.Stretch)
        self._table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._table.setAlternatingRowColors(True)
        self._table.setStyleSheet(
            "QTableWidget{border:1px solid #ddd; font-size:12px;}"
            "QTableWidget::item:selected{background:#d0e4f8; color:#000;}"
        )
        self._table.verticalHeader().setVisible(False)
        self._table.itemSelectionChanged.connect(self._auswahl_geaendert)
        self._table.itemDoubleClicked.connect(lambda _: self._bearbeiten())
        layout.addWidget(self._table, 1)

        self._stat_lbl = QLabel()
        self._stat_lbl.setStyleSheet(
            "background:#e8f5e9; color:#256029; border-radius:4px;"
            "padding:6px 12px; font-size:11px;"
        )
        layout.addWidget(self._stat_lbl)

    def refresh(self):
        current_j = self._combo_jahr.currentData()
        self._combo_jahr.blockSignals(True)
        self._combo_jahr.clear()
        self._combo_jahr.addItem("Alle", None)
        for j in verfuegbare_jahre_patienten():
            self._combo_jahr.addItem(str(j), j)
        for i in range(self._combo_jahr.count()):
            if self._combo_jahr.itemData(i) == current_j:
                self._combo_jahr.setCurrentIndex(i)
                break
        self._combo_jahr.blockSignals(False)
        self._lade()

    def _filter_reset(self):
        self._combo_jahr.blockSignals(True)
        self._combo_monat.blockSignals(True)
        self._suche.blockSignals(True)
        self._combo_jahr.setCurrentIndex(0)
        self._combo_monat.setCurrentIndex(0)
        self._suche.clear()
        self._combo_jahr.blockSignals(False)
        self._combo_monat.blockSignals(False)
        self._suche.blockSignals(False)
        self._lade()

    def _filter_changed(self):
        self._lade()

    def _lade(self):
        jahr  = self._combo_jahr.currentData()
        monat = self._combo_monat.currentData()
        suche = self._suche.text().strip() or None
        try:
            eintraege = lade_patienten(monat=monat, jahr=jahr, suchtext=suche)
        except Exception as exc:
            QMessageBox.critical(self, "Datenbankfehler", str(exc))
            return
        self._eintraege = eintraege
        self._table.setRowCount(len(eintraege))
        for row, p in enumerate(eintraege):
            def _item(text: str, center: bool = False) -> QTableWidgetItem:
                item = QTableWidgetItem(str(text) if text else "")
                if center:
                    item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                return item
            name    = p.get("patient_name", "")
            abtlg   = p.get("patient_abteilung", "")
            display = f"{name} ({abtlg})" if name and abtlg else (name or "—")
            beschwerde = p.get("beschwerde_art", "") or p.get("symptome", "")
            massnahmen = p.get("massnahmen", "")
            self._table.setItem(row, 0,  _item(str(row + 1), center=True))
            self._table.setItem(row, 1,  _item(p.get("datum", "")))
            self._table.setItem(row, 2,  _item(p.get("uhrzeit", ""), center=True))
            dauer = p.get("behandlungsdauer", 0) or 0
            self._table.setItem(row, 3,  _item(str(dauer) if dauer else "—", center=True))
            self._table.setItem(row, 4,  _item(p.get("patient_typ", "") or "—"))
            self._table.setItem(row, 5,  _item(display))
            alter = p.get("patient_alter", 0)
            self._table.setItem(row, 6,  _item(str(alter) if alter else "—", center=True))
            self._table.setItem(row, 7,  _item(
                beschwerde[:47] + "..." if len(beschwerde) > 50 else beschwerde
            ))
            self._table.setItem(row, 8,  _item(p.get("unfall_ort", "") or "—"))
            self._table.setItem(row, 9,  _item(
                massnahmen[:47] + "..." if len(massnahmen) > 50 else massnahmen
            ))
            self._table.setItem(row, 10, _item(p.get("drk_ma1", "")))
            self._table.setItem(row, 11, _item(p.get("weitergeleitet", "") or "—"))
            bg = p.get("arbeitsunfall", 0)
            bg_item = _item("✓ BG" if bg else "—", center=True)
            if bg:
                bg_item.setForeground(QColor("#b00000"))
            self._table.setItem(row, 12, bg_item)
        n = len(eintraege)
        self._treffer_lbl.setText(f"{n} Patient{'en' if n != 1 else ''} gefunden")
        self._stat_lbl.setText(f"📊  Gesamt: {n} Patient{'en' if n != 1 else ''}")
        self._auswahl_geaendert()

    def _auswahl_geaendert(self):
        has = len(self._table.selectedItems()) > 0
        self._btn_bearbeiten.setEnabled(has)
        self._btn_loeschen.setEnabled(has)
        self._btn_material.setEnabled(has)
        self._btn_word.setEnabled(has)
        self._btn_mail.setEnabled(has)

    def _neu(self):
        dlg = _PatientenDialog(parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            daten, material = dlg.get_data()
            try:
                patient_speichern(daten, material)
                self.refresh()
                QMessageBox.information(self, "Erfolg", "Patient erfolgreich erfasst.")
            except Exception as exc:
                QMessageBox.critical(self, "Fehler beim Speichern", str(exc))

    def _bearbeiten(self):
        row = self._table.currentRow()
        if row < 0 or row >= len(self._eintraege):
            return
        eintrag = self._eintraege[row]
        dlg = _PatientenDialog(daten=eintrag, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            daten, material = dlg.get_data()
            try:
                patient_aktualisieren(eintrag["id"], daten, material)
                self.refresh()
                QMessageBox.information(self, "Erfolg", "Patient aktualisiert.")
            except Exception as exc:
                QMessageBox.critical(self, "Fehler beim Aktualisieren", str(exc))

    def _loeschen(self):
        row = self._table.currentRow()
        if row < 0 or row >= len(self._eintraege):
            return
        eintrag = self._eintraege[row]
        antwort = QMessageBox.question(
            self, "Löschen bestätigen",
            f"Patienten-Datensatz vom {eintrag.get('datum','')} wirklich löschen?\n"
            f"Diese Aktion kann nicht rückgängig gemacht werden.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )
        if antwort == QMessageBox.StandardButton.Yes:
            try:
                patient_loeschen(eintrag["id"])
                self.refresh()
                QMessageBox.information(self, "Gelöscht", "Patient wurde gelöscht.")
            except Exception as exc:
                QMessageBox.critical(self, "Fehler beim Löschen", str(exc))

    def _material_anzeigen(self):
        row = self._table.currentRow()
        if row < 0 or row >= len(self._eintraege):
            return
        eintrag = self._eintraege[row]
        try:
            material = lade_verbrauchsmaterial(eintrag["id"])
        except Exception as exc:
            QMessageBox.critical(self, "Fehler", str(exc))
            return
        dlg = QDialog(self)
        dlg.setWindowTitle(
            f"📦  Verbrauchsmaterial – "
            f"{eintrag.get('patient_name', 'Patient')} ({eintrag.get('datum', '')})"
        )
        dlg.setMinimumWidth(500)
        layout = QVBoxLayout(dlg)
        if not material:
            layout.addWidget(QLabel("Kein Verbrauchsmaterial erfasst."))
        else:
            table = QTableWidget()
            table.setColumnCount(3)
            table.setHorizontalHeaderLabels(["Material", "Menge", "Einheit"])
            table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
            table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
            table.setRowCount(len(material))
            for row_m, m in enumerate(material):
                table.setItem(row_m, 0, QTableWidgetItem(m.get("material", "")))
                table.setItem(row_m, 1, QTableWidgetItem(str(m.get("menge", ""))))
                table.setItem(row_m, 2, QTableWidgetItem(m.get("einheit", "")))
            layout.addWidget(table)
        btn_close = QPushButton("Schließen")
        btn_close.clicked.connect(dlg.accept)
        layout.addWidget(btn_close)
        dlg.exec()

    def _word_protokoll(self):
        row = self._table.currentRow()
        if row < 0 or row >= len(self._eintraege):
            return
        eintrag = self._eintraege[row]
        try:
            material = lade_verbrauchsmaterial(eintrag["id"])
            medikamente = lade_medikamente(eintrag["id"])
            pfad = export_patient_word(eintrag, material, medikamente)
        except Exception as exc:
            QMessageBox.critical(self, "Fehler beim Word-Export", str(exc))
            return
        antwort = QMessageBox.question(
            self, "Word-Protokoll erstellt",
            f"Das Protokoll wurde gespeichert:\n{pfad}\n\nJetzt in Word öffnen?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.Yes,
        )
        if antwort == QMessageBox.StandardButton.Yes:
            os.startfile(pfad)

    def _mail_protokoll(self):
        row = self._table.currentRow()
        if row < 0 or row >= len(self._eintraege):
            return
        eintrag = self._eintraege[row]

        # Word-Datei erstellen
        try:
            material = lade_verbrauchsmaterial(eintrag["id"])
            medikamente = lade_medikamente(eintrag["id"])
            word_pfad = export_patient_word(eintrag, material, medikamente)
        except Exception as exc:
            QMessageBox.critical(self, "Fehler beim Word-Export", str(exc))
            return

        # Mail-Dialog
        dlg = _PatientenMailDialog(
            word_pfad=word_pfad,
            patient_name=eintrag.get("patient_name", ""),
            datum=eintrag.get("datum", ""),
            parent=self,
        )
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        empfaenger, betreff, body = dlg.get_daten()

        try:
            from functions.mail_functions import create_outlook_draft
            logo_pfad = os.path.join(BASE_DIR, "Daten", "Email", "Logo.jpg")
            create_outlook_draft(
                to=empfaenger,
                subject=betreff,
                body_text=body,
                attachment_path=word_pfad,
                logo_path=logo_pfad if os.path.isfile(logo_pfad) else None,
            )
            QMessageBox.information(
                self, "Outlook geöffnet",
                "Der Outlook-Entwurf wurde geöffnet.\n"
                "Bitte prüfen und manuell absenden."
            )
        except Exception as exc:
            QMessageBox.critical(self, "Fehler beim E-Mail-Versand", str(exc))

    def _excel_exportieren(self, ask_path: bool = False):
        if not self._eintraege:
            QMessageBox.information(self, "Kein Inhalt", "Keine Einträge zur Ansicht vorhanden.")
            return

        # Datumszeitraum wählen
        zeitraum_dlg = _DatumsbereichDialog(self._eintraege, self)
        if zeitraum_dlg.exec() != QDialog.DialogCode.Accepted:
            return
        export_eintraege = zeitraum_dlg.get_gefilterte()
        zeitraum_lbl = zeitraum_dlg.get_zeitraum_label()

        ziel = None
        if ask_path:
            from PySide6.QtWidgets import QFileDialog
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            vorschlag = os.path.join(_PATIENTEN_EXPORT_DIR, f"Patienten_Export_{ts}.xlsx")
            os.makedirs(_PATIENTEN_EXPORT_DIR, exist_ok=True)
            ziel, _ = QFileDialog.getSaveFileName(
                self,
                "Excel-Datei speichern",
                vorschlag,
                "Excel-Dateien (*.xlsx)",
            )
            if not ziel:
                return

        try:
            pfad = export_patienten_excel(
                export_eintraege,
                ziel_pfad=ziel,
                titel_zeitraum=zeitraum_lbl,
            )
        except ImportError as e:
            QMessageBox.critical(self, "Modul fehlt", str(e))
            return
        except Exception as exc:
            QMessageBox.critical(self, "Export-Fehler", f"Fehler beim Exportieren:\n{exc}")
            return

        antwort = QMessageBox.question(
            self,
            "Excel gespeichert",
            f"Excel wurde erfolgreich gespeichert:\n\n{pfad}\n\nDatei jetzt öffnen?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if antwort == QMessageBox.StandardButton.Yes:
            import subprocess
            subprocess.Popen(["explorer", pfad], shell=False)

        return pfad


# ──────────────────────────────────────────────────────────────────────────────
#  Einsätze-Tab-Widget
# ──────────────────────────────────────────────────────────────────────────────

class _EinsaetzeTab(QWidget):
    """Tab 'Einsätze' – Einsatzprotokoll nach Vorlage FKB."""

    # Tabellen-Spalten
    _COLS = [
        "Lfd. Nr.", "Datum", "Uhrzeit", "Dauer\n(Min.)",
        "Einsatzstichwort", "Einsatzort", "Einsatznr.\nDRK",
        "DRK MA 1", "DRK MA 2", "Angenommen", "Grund (bei Nein)", "Versendet"
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self._eintraege: list[dict] = []
        self._build_ui()
        self.refresh()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # ── Titelzeile ────────────────────────────────────────────────────
        titel_lbl = QLabel("🚑  Einsatzstatistik – Notfalleinsätze FKB")
        titel_lbl.setFont(QFont("Arial", 15, QFont.Weight.Bold))
        titel_lbl.setStyleSheet(f"color:{FIORI_TEXT}; padding: 4px 0;")
        layout.addWidget(titel_lbl)

        hinweis_lbl = QLabel(
            "Alarmierung über die Leitstelle per Telefon oder Digitalmelder"
        )
        hinweis_lbl.setStyleSheet("color:#666; font-size:11px; font-style:italic;")
        layout.addWidget(hinweis_lbl)

        # ── Filter-Leiste ─────────────────────────────────────────────────
        filter_frame = QFrame()
        filter_frame.setStyleSheet(
            "QFrame{background:#f8f9fa;border:1px solid #ddd;"
            "border-radius:4px;padding:4px;}"
        )
        fl = QHBoxLayout(filter_frame)
        fl.setContentsMargins(8, 6, 8, 6)
        fl.setSpacing(10)

        fl.addWidget(QLabel("Jahr:"))
        self._combo_jahr = QComboBox()
        self._combo_jahr.setFixedWidth(80)
        self._combo_jahr.addItem("Alle", None)
        self._combo_jahr.currentIndexChanged.connect(self._filter_changed)
        fl.addWidget(self._combo_jahr)

        fl.addWidget(QLabel("Monat:"))
        self._combo_monat = QComboBox()
        self._combo_monat.setFixedWidth(110)
        for i, m in enumerate([
            "Alle", "Januar", "Februar", "März", "April", "Mai", "Juni",
            "Juli", "August", "September", "Oktober", "November", "Dezember"
        ]):
            self._combo_monat.addItem(m, None if i == 0 else i)
        self._combo_monat.currentIndexChanged.connect(self._filter_changed)
        fl.addWidget(self._combo_monat)

        fl.addWidget(QLabel("Suche:"))
        self._suche = QLineEdit()
        self._suche.setPlaceholderText("Stichwort, Ort, Mitarbeiter …")
        self._suche.setMinimumWidth(200)
        self._suche.textChanged.connect(self._filter_changed)
        fl.addWidget(self._suche, 1)

        btn_reset = _btn_light("✖ Zurücksetzen")
        btn_reset.setFixedHeight(28)
        btn_reset.clicked.connect(self._filter_reset)
        fl.addWidget(btn_reset)

        layout.addWidget(filter_frame)

        # ── Aktions-Buttons ────────────────────────────────────────────────
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        self._btn_neu = _btn("＋  Neuer Einsatz", "#107e3e", "#0a5c2e")
        self._btn_neu.setToolTip("Neuen Notfalleinsatz erfassen")
        self._btn_neu.clicked.connect(self._neu)
        btn_row.addWidget(self._btn_neu)

        self._btn_bearbeiten = _btn_light("✏  Bearbeiten")
        self._btn_bearbeiten.setEnabled(False)
        self._btn_bearbeiten.setToolTip("Ausgewählten Einsatz bearbeiten")
        self._btn_bearbeiten.clicked.connect(self._bearbeiten)
        btn_row.addWidget(self._btn_bearbeiten)

        self._btn_loeschen = _btn_light("🗑  Löschen")
        self._btn_loeschen.setEnabled(False)
        self._btn_loeschen.setStyleSheet(
            "QPushButton{background:#eee;color:#333;border:none;"
            "border-radius:4px;padding:4px 14px;font-size:12px;}"
            "QPushButton:hover{background:#ffcccc;color:#a00;}"
            "QPushButton:disabled{background:#f5f5f5;color:#bbb;}"
        )
        self._btn_loeschen.clicked.connect(self._loeschen)
        btn_row.addWidget(self._btn_loeschen)

        # Trennstrich
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.VLine)
        sep.setStyleSheet("color:#ccc;")
        btn_row.addWidget(sep)

        self._btn_excel = _btn("📈  Excel exportieren", "#1565a8", "#0d47a1")
        self._btn_excel.setToolTip(
            f"Aktuelle Ansicht als Excel speichern\n"
            f"Standard-Ordner: Daten/Einsatz/Protokolle"
        )
        self._btn_excel.clicked.connect(self._excel_exportieren)
        btn_row.addWidget(self._btn_excel)

        self._btn_excel_wo = _btn_light("📂  Speicherort wählen")
        self._btn_excel_wo.setToolTip("Excel exportieren und Speicherort manuell wählen")
        self._btn_excel_wo.clicked.connect(lambda: self._excel_exportieren(ask_path=True))
        btn_row.addWidget(self._btn_excel_wo)

        self._btn_mail = _btn("📧  Per E-Mail senden", "#6a1b9a", "#4a148c")
        self._btn_mail.setToolTip("Excel exportieren und als Outlook-Entwurf öffnen")
        self._btn_mail.clicked.connect(self._per_mail_senden)
        btn_row.addWidget(self._btn_mail)

        btn_row.addStretch()

        self._treffer_lbl = QLabel()
        self._treffer_lbl.setStyleSheet("color:#666; font-size:11px;")
        btn_row.addWidget(self._treffer_lbl)

        layout.addLayout(btn_row)

        # ── Tabelle ────────────────────────────────────────────────────────
        self._table = QTableWidget()
        self._table.setColumnCount(len(self._COLS))
        self._table.setHorizontalHeaderLabels(self._COLS)
        hh = self._table.horizontalHeader()
        hh.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)  # Lfd. Nr.
        hh.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)  # Datum
        hh.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)  # Uhrzeit
        hh.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)  # Dauer
        hh.setSectionResizeMode(4, QHeaderView.ResizeMode.Stretch)           # Stichwort
        hh.setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)  # Ort
        hh.setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents)  # Einsatznr.
        hh.setSectionResizeMode(7, QHeaderView.ResizeMode.ResizeToContents)  # MA1
        hh.setSectionResizeMode(8, QHeaderView.ResizeMode.ResizeToContents)  # MA2
        hh.setSectionResizeMode(9, QHeaderView.ResizeMode.ResizeToContents)  # Angenommen
        hh.setSectionResizeMode(10, QHeaderView.ResizeMode.ResizeToContents) # Grund
        hh.setSectionResizeMode(11, QHeaderView.ResizeMode.ResizeToContents) # Versendet
        self._table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._table.setAlternatingRowColors(True)
        self._table.setStyleSheet(
            "QTableWidget{border:1px solid #ddd; font-size:12px;}"
            "QTableWidget::item:selected{background:#d0e4f8; color:#000;}"
        )
        self._table.verticalHeader().setVisible(False)
        self._table.itemSelectionChanged.connect(self._auswahl_geaendert)
        self._table.itemDoubleClicked.connect(lambda _: self._bearbeiten())
        layout.addWidget(self._table, 1)

        # ── Statistik-Leiste ───────────────────────────────────────────────
        self._stat_lbl = QLabel()
        self._stat_lbl.setStyleSheet(
            "background:#e8f5e9; color:#256029; border-radius:4px;"
            "padding:6px 12px; font-size:11px;"
        )
        layout.addWidget(self._stat_lbl)

    # ── Refresh / Laden ───────────────────────────────────────────────────────

    def refresh(self):
        # Jahre aktualisieren
        current_j = self._combo_jahr.currentData()
        self._combo_jahr.blockSignals(True)
        self._combo_jahr.clear()
        self._combo_jahr.addItem("Alle", None)
        for j in verfuegbare_jahre_einsaetze():
            self._combo_jahr.addItem(str(j), j)
        # Aktuelles Jahr wiederherstellen
        for i in range(self._combo_jahr.count()):
            if self._combo_jahr.itemData(i) == current_j:
                self._combo_jahr.setCurrentIndex(i)
                break
        self._combo_jahr.blockSignals(False)
        self._lade()

    def _filter_reset(self):
        self._combo_jahr.blockSignals(True)
        self._combo_monat.blockSignals(True)
        self._suche.blockSignals(True)
        self._combo_jahr.setCurrentIndex(0)
        self._combo_monat.setCurrentIndex(0)
        self._suche.clear()
        self._combo_jahr.blockSignals(False)
        self._combo_monat.blockSignals(False)
        self._suche.blockSignals(False)
        self._lade()

    def _filter_changed(self):
        self._lade()

    def _lade(self):
        jahr   = self._combo_jahr.currentData()
        monat  = self._combo_monat.currentData()
        suche  = self._suche.text().strip() or None
        try:
            eintraege = lade_einsaetze(monat=monat, jahr=jahr, suchtext=suche)
        except Exception as exc:
            QMessageBox.critical(self, "Datenbankfehler", str(exc))
            return

        self._eintraege = eintraege
        lfd_nr = 1
        self._table.setRowCount(len(eintraege))
        angenommen_count = 0
        for row, e in enumerate(eintraege):
            ang = bool(e.get("angenommen", 1))
            if ang:
                angenommen_count += 1

            def _item(text: str, center: bool = False) -> QTableWidgetItem:
                item = QTableWidgetItem(str(text) if text else "")
                if center:
                    item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                return item

            self._table.setItem(row, 0, _item(str(row + 1), center=True))
            self._table.setItem(row, 1, _item(e.get("datum", "")))
            self._table.setItem(row, 2, _item(e.get("uhrzeit", ""), center=True))
            dauer = e.get("einsatzdauer", 0) or 0
            self._table.setItem(row, 3, _item(str(dauer) if dauer else "—", center=True))
            self._table.setItem(row, 4, _item(e.get("einsatzstichwort", "")))
            self._table.setItem(row, 5, _item(e.get("einsatzort", "")))
            self._table.setItem(row, 6, _item(e.get("einsatznr_drk", "") or "—"))
            self._table.setItem(row, 7, _item(e.get("drk_ma1", "")))
            self._table.setItem(row, 8, _item(e.get("drk_ma2", "") or "—"))

            ang_item = _item("✅ Ja" if ang else "❌ Nein", center=True)
            if not ang:
                ang_item.setBackground(QColor("#ffe0e0"))
            else:
                ang_item.setBackground(QColor("#e8f5e9"))
            self._table.setItem(row, 9, ang_item)

            self._table.setItem(row, 10, _item(e.get("grund_abgelehnt", "") or "—"))

            gesendet = bool(e.get("gesendet"))
            gs_item = _item("✅ Ja" if gesendet else "⬜ Nein", center=True)
            if gesendet:
                gs_item.setBackground(QColor("#e8f5e9"))
            self._table.setItem(row, 11, gs_item)

        n = len(eintraege)
        abgelehnt = n - angenommen_count
        self._treffer_lbl.setText(f"{n} Einsatz{'ätze' if n != 1 else ''} gefunden")
        self._stat_lbl.setText(
            f"📊  Gesamt: {n}  |  ✅ Angenommen: {angenommen_count}  |  ❌ Abgelehnt: {abgelehnt}"
        )
        self._auswahl_geaendert()

    # ── Aktionen ──────────────────────────────────────────────────────────────

    def _auswahl_geaendert(self):
        hat = bool(self._aktuell())
        self._btn_bearbeiten.setEnabled(hat)
        self._btn_loeschen.setEnabled(hat)

    def _aktuell(self) -> dict | None:
        row = self._table.currentRow()
        if 0 <= row < len(self._eintraege):
            return self._eintraege[row]
        return None

    def _neu(self):
        dlg = _EinsatzDialog(parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            daten = dlg.get_daten()
            try:
                einsatz_id = einsatz_speichern(daten)
                self.refresh()
            except Exception as exc:
                QMessageBox.critical(self, "Fehler", f"Fehler beim Speichern:\n{exc}")
                return
            # Verbrauchsmaterial in Sanmat buchen
            verbrauch = daten.get("verbrauch_liste", [])
            if verbrauch:
                try:
                    db = _SanmatDB()
                    db.initialize()
                    datum_raw = daten.get("datum", "")
                    try:
                        t = datum_raw.split(".")
                        datum_iso = f"{t[2]}-{t[1]}-{t[0]}" if len(t) == 3 else datum_raw
                    except Exception:
                        datum_iso = datum_raw
                    entnehmer = daten.get("entnehmer", "")
                    stichwort = daten.get("einsatzstichwort", "") or f"Einsatz-ID {einsatz_id}"
                    fehler = []
                    for pos in verbrauch:
                        ok, msg = db.entnehmen(
                            artikel_id=pos["artikel_id"],
                            artikel_name=pos["bezeichnung"],
                            menge=pos["menge"],
                            datum=datum_iso,
                            typ="verbrauch",
                            von=entnehmer,
                            bemerkung=f"Einsatz: {stichwort}  (ID {einsatz_id})",
                            negativ_erlaubt=True,
                        )
                        if not ok:
                            fehler.append(f"{pos['bezeichnung']}: {msg}")
                    if fehler:
                        QMessageBox.warning(
                            self, "Sanmat-Buchung",
                            "Einige Buchungen fehlgeschlagen:\n" + "\n".join(fehler)
                        )
                except Exception as exc:
                    QMessageBox.warning(
                        self, "Sanmat-Buchung",
                        f"Fehler beim Buchen des Verbrauchsmaterials:\n{exc}"
                    )

    def _bearbeiten(self):
        e = self._aktuell()
        if not e:
            return
        dlg = _EinsatzDialog(daten=e, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            daten = dlg.get_daten()
            try:
                einsatz_aktualisieren(e["id"], daten)
                self.refresh()
            except Exception as exc:
                QMessageBox.critical(self, "Fehler", f"Fehler beim Speichern:\n{exc}")

    def _loeschen(self):
        e = self._aktuell()
        if not e:
            return
        antwort = QMessageBox.question(
            self, "Einsatz löschen",
            f"Einsatz wirklich löschen?\n\n"
            f"Datum:       {e.get('datum', '')}\n"
            f"Uhrzeit:     {e.get('uhrzeit', '')}\n"
            f"Stichwort:   {e.get('einsatzstichwort', '')}\n"
            f"Einsatzort:  {e.get('einsatzort', '')}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if antwort == QMessageBox.StandardButton.Yes:
            try:
                einsatz_loeschen(e["id"])
                self.refresh()
            except Exception as exc:
                QMessageBox.critical(self, "Fehler", f"Fehler beim Löschen:\n{exc}")

    # ── Export-Helfer ─────────────────────────────────────────────────────────

    def _zeitraum_label(self) -> str:
        """Gibt einen lesbaren Zeitraum-Text aus den aktuellen Filtern zurück."""
        teile = []
        j = self._combo_jahr.currentData()
        m = self._combo_monat.currentData()
        _MONATE = ["","Januar","Februar","März","April","Mai","Juni",
                   "Juli","August","September","Oktober","November","Dezember"]
        if m:
            teile.append(_MONATE[m])
        if j:
            teile.append(str(j))
        return " ".join(teile) if teile else "Alle Einträge"

    def _excel_exportieren(self, ask_path: bool = False):
        if not self._eintraege:
            QMessageBox.information(self, "Kein Inhalt", "Keine Einträge zur Ansicht vorhanden.")
            return

        # Datumszeitraum wählen
        zeitraum_dlg = _DatumsbereichDialog(self._eintraege, self)
        if zeitraum_dlg.exec() != QDialog.DialogCode.Accepted:
            return
        export_eintraege = zeitraum_dlg.get_gefilterte()
        zeitraum_lbl = zeitraum_dlg.get_zeitraum_label()

        ziel = None
        if ask_path:
            from PySide6.QtWidgets import QFileDialog
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            vorschlag = os.path.join(_PROTOKOLL_DIR, f"Einsatzprotokoll_{ts}.xlsx")
            os.makedirs(_PROTOKOLL_DIR, exist_ok=True)
            ziel, _ = QFileDialog.getSaveFileName(
                self,
                "Excel-Datei speichern",
                vorschlag,
                "Excel-Dateien (*.xlsx)",
            )
            if not ziel:
                return

        try:
            pfad = export_einsaetze_excel(
                export_eintraege,
                ziel_pfad=ziel,
                titel_zeitraum=zeitraum_lbl,
            )
        except ImportError as e:
            QMessageBox.critical(self, "Modul fehlt", str(e))
            return
        except Exception as exc:
            QMessageBox.critical(self, "Export-Fehler", f"Fehler beim Exportieren:\n{exc}")
            return

        antwort = QMessageBox.question(
            self,
            "Excel gespeichert",
            f"Excel wurde erfolgreich gespeichert:\n\n{pfad}\n\nDatei jetzt öffnen?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if antwort == QMessageBox.StandardButton.Yes:
            import subprocess
            subprocess.Popen(["explorer", pfad], shell=False)

        return pfad

    def _per_mail_senden(self):
        if not self._eintraege:
            QMessageBox.information(self, "Kein Inhalt", "Keine Einträge zur Ansicht vorhanden.")
            return

        # Datumszeitraum wählen
        zeitraum_dlg = _DatumsbereichDialog(self._eintraege, self)
        if zeitraum_dlg.exec() != QDialog.DialogCode.Accepted:
            return
        export_eintraege = zeitraum_dlg.get_gefilterte()
        zeitraum_lbl = zeitraum_dlg.get_zeitraum_label()

        # Excel speichern
        try:
            excel_pfad = export_einsaetze_excel(
                export_eintraege,
                titel_zeitraum=zeitraum_lbl,
            )
        except ImportError as e:
            QMessageBox.critical(self, "Modul fehlt", str(e))
            return
        except Exception as exc:
            QMessageBox.critical(self, "Export-Fehler", f"Fehler beim Erstellen der Excel:\n{exc}")
            return

        # Mail-Dialog
        dlg = _MailDialog(
            excel_pfad=excel_pfad,
            zeitraum=zeitraum_lbl,
            anzahl=len(export_eintraege),
            parent=self,
        )
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        empfaenger, betreff, body = dlg.get_daten()

        try:
            from functions.mail_functions import create_outlook_draft
            logo_pfad = os.path.join(BASE_DIR, "Daten", "Email", "Logo.jpg")
            create_outlook_draft(
                to=empfaenger,
                subject=betreff,
                body_text=body,
                attachment_path=excel_pfad,
                logo_path=logo_pfad if os.path.isfile(logo_pfad) else None,
            )
            QMessageBox.information(
                self, "Outlook geöffnet",
                "Der Outlook-Entwurf wurde geöffnet.\n"
                "Bitte prüfen und manuell absenden."
            )
        except Exception as exc:
            QMessageBox.critical(
                self, "E-Mail-Fehler",
                f"Outlook-Entwurf konnte nicht erstellt werden:\n{exc}\n\n"
                f"Die Excel-Datei wurde trotzdem gespeichert:\n{excel_pfad}"
            )


# ──────────────────────────────────────────────────────────────────────────────
#  Dialog: E-Mail versenden
# ──────────────────────────────────────────────────────────────────────────────

class _DatumsbereichDialog(QDialog):
    """
    Wählt einen Von-Bis-Datumsbereich für den Export.
    Gibt die gefilterten Einträge zurück.
    """

    _FIELD_STYLE = (
        "QDateEdit { border:1px solid #ccc; border-radius:4px;"
        "padding:4px; font-size:12px; background:white;}"
    )

    def __init__(self, eintraege: list[dict], parent=None):
        super().__init__(parent)
        self.setWindowTitle("📅  Datumszeitraum wählen")
        self.setMinimumWidth(400)
        self.setFixedHeight(230)
        self._alle = eintraege

        layout = QVBoxLayout(self)
        layout.setSpacing(10)
        layout.setContentsMargins(16, 14, 16, 10)

        info = QLabel(
            "Über welchen Zeitraum soll die Excel-Datei erstellt werden?"
        )
        info.setStyleSheet("color:#444; font-size:12px;")
        info.setWordWrap(True)
        layout.addWidget(info)

        # Von / Bis
        fl = QFormLayout()
        fl.setSpacing(8)

        # Defaultwerte: frühestes / spätestes Datum aus den Einträgen
        def _to_qdate(s: str) -> QDate:
            try:
                d, m, y = s.split(".")
                return QDate(int(y), int(m), int(d))
            except Exception:
                return QDate.currentDate()

        daten = [e.get("datum", "") for e in eintraege if e.get("datum")]
        if daten:
            # sortierbar machen: yyyymmdd
            def _sort_key(s):
                try:
                    d, m, y = s.split(".")
                    return y + m + d
                except Exception:
                    return ""
            daten_sorted = sorted(daten, key=_sort_key)
            von_default = _to_qdate(daten_sorted[0])
            bis_default = _to_qdate(daten_sorted[-1])
        else:
            von_default = bis_default = QDate.currentDate()

        self._von = QDateEdit()
        self._von.setCalendarPopup(True)
        self._von.setDisplayFormat("dd.MM.yyyy")
        self._von.setDate(von_default)
        self._von.setStyleSheet(self._FIELD_STYLE)
        fl.addRow("Von:", self._von)

        self._bis = QDateEdit()
        self._bis.setCalendarPopup(True)
        self._bis.setDisplayFormat("dd.MM.yyyy")
        self._bis.setDate(bis_default)
        self._bis.setStyleSheet(self._FIELD_STYLE)
        fl.addRow("Bis:", self._bis)

        layout.addLayout(fl)

        self._treffer_lbl = QLabel()
        self._treffer_lbl.setStyleSheet(
            "background:#e3f2fd; color:#154360; border-radius:4px;"
            "padding:5px 10px; font-size:11px;"
        )
        layout.addWidget(self._treffer_lbl)
        self._update_treffer()
        self._von.dateChanged.connect(self._update_treffer)
        self._bis.dateChanged.connect(self._update_treffer)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("📈  Exportieren")
        btns.accepted.connect(self._on_accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def _filter(self) -> list[dict]:
        """Gibt alle Einträge zurück die im gewählten Zeitraum liegen."""
        von_str = self._von.date().toString("yyyyMMdd")
        bis_str = self._bis.date().toString("yyyyMMdd")

        def _key(e) -> str:
            try:
                d, m, y = e.get("datum", "").split(".")
                return y + m + d
            except Exception:
                return ""

        return [e for e in self._alle if von_str <= _key(e) <= bis_str]

    def _update_treffer(self):
        n = len(self._filter())
        von = self._von.date().toString("dd.MM.yyyy")
        bis = self._bis.date().toString("dd.MM.yyyy")
        self._treffer_lbl.setText(
            f"📊  {n} Einsatz/ätze im Zeitraum {von} – {bis}"
        )

    def _on_accept(self):
        if not self._filter():
            QMessageBox.warning(
                self, "Keine Einträge",
                "Im gewählten Zeitraum gibt es keine Einsätze."
            )
            return
        self.accept()

    def get_gefilterte(self) -> list[dict]:
        return self._filter()

    def get_zeitraum_label(self) -> str:
        von = self._von.date().toString("dd.MM.yyyy")
        bis = self._bis.date().toString("dd.MM.yyyy")
        return f"{von} – {bis}"


class _MailDialog(QDialog):
    """Kleiner Dialog: Empfänger, Betreff und Nachrichtentext für Einsatz-Mail."""

    _FIELD_STYLE = (
        "QLineEdit, QTextEdit { border:1px solid #ccc; border-radius:4px;"
        "padding:4px; font-size:12px; background:white;}"
    )

    def __init__(
        self,
        excel_pfad: str,
        zeitraum: str,
        anzahl: int,
        parent=None,
    ):
        super().__init__(parent)
        self.setWindowTitle("📧  Einsatzprotokoll per E-Mail senden")
        self.setMinimumWidth(500)
        self.resize(540, 420)
        self._excel_pfad = excel_pfad
        layout = QVBoxLayout(self)
        fl = QFormLayout()
        fl.setSpacing(8)

        self._empfaenger = QLineEdit()
        self._empfaenger.setPlaceholderText("empfaenger@drk-koeln.de")
        self._empfaenger.setStyleSheet(self._FIELD_STYLE)
        fl.addRow("Empfänger:", self._empfaenger)

        self._betreff = QLineEdit()
        self._betreff.setText(
            f"Einsatzprotokoll FKB – {zeitraum}  ({anzahl} Einsätze)"
        )
        self._betreff.setStyleSheet(self._FIELD_STYLE)
        fl.addRow("Betreff:", self._betreff)

        layout.addLayout(fl)
        layout.addWidget(QLabel("Nachrichtentext:"))

        self._body = QTextEdit()
        self._body.setStyleSheet(self._FIELD_STYLE)
        self._body.setPlainText(
            f"Hallo,\n\n"
            f"anbei das Einsatzprotokoll für den Zeitraum {zeitraum} "
            f"({anzahl} Einsätze).\n\n"
            f"Das Protokoll ist als Excel-Datei angehängt.\n\n"
            f"Mit freundlichen Grüßen\n"
            f"DRK-Kreisverband Köln e.V."
        )
        self._body.setMinimumHeight(140)
        layout.addWidget(self._body, 1)

        anhang_lbl = QLabel(f"📎  Anhang: {os.path.basename(excel_pfad)}")
        anhang_lbl.setStyleSheet(
            "background:#e3f2fd; color:#154360; border-radius:4px;"
            "padding:6px 10px; font-size:11px;"
        )
        layout.addWidget(anhang_lbl)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("📧  Outlook-Entwurf öffnen")
        btns.accepted.connect(self._on_accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def _on_accept(self):
        if not self._empfaenger.text().strip():
            QMessageBox.warning(self, "Pflichtfeld", "Bitte einen Empfänger angeben.")
            return
        self.accept()

    def get_daten(self) -> tuple[str, str, str]:
        return (
            self._empfaenger.text().strip(),
            self._betreff.text().strip(),
            self._body.toPlainText().strip(),
        )


# ──────────────────────────────────────────────────────────────────────────────
#  Übersicht-Tab
# ──────────────────────────────────────────────────────────────────────────────

class _UebersichtTab(QWidget):
    """Tab 'Übersicht' – Statistiken und Kennzahlen über alle Einsätze."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._build_ui()
        self.refresh()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(10)

        # ── Titel ──────────────────────────────────────────────────────────
        titel = QLabel("📊  Übersicht Einsatzstatistik")
        titel.setFont(QFont("Arial", 15, QFont.Weight.Bold))
        titel.setStyleSheet(f"color:{FIORI_TEXT}; padding:4px 0;")
        layout.addWidget(titel)

        # ── Filter: Jahr ───────────────────────────────────────────────────
        fff = QFrame()
        fff.setStyleSheet(
            "QFrame{background:#f8f9fa;border:1px solid #ddd;"
            "border-radius:4px;padding:4px;}"
        )
        fl = QHBoxLayout(fff)
        fl.setContentsMargins(8, 6, 8, 6)
        fl.setSpacing(10)
        fl.addWidget(QLabel("Jahr:"))
        self._combo_jahr = QComboBox()
        self._combo_jahr.setFixedWidth(90)
        self._combo_jahr.addItem("Alle", None)
        self._combo_jahr.currentIndexChanged.connect(self.refresh)
        fl.addWidget(self._combo_jahr)
        fl.addStretch()
        layout.addWidget(fff)

        # ── KPI-Kacheln ────────────────────────────────────────────────────
        kpi_frame = QFrame()
        kpi_layout = QHBoxLayout(kpi_frame)
        kpi_layout.setSpacing(12)
        kpi_layout.setContentsMargins(0, 0, 0, 0)

        self._kpi_gesamt   = self._kpi_card("📥  Gesamt",      "0", "#1565a8")
        self._kpi_ja       = self._kpi_card("✅  Angenommen", "0", "#2e7d32")
        self._kpi_nein     = self._kpi_card("❌  Abgelehnt",  "0", "#c62828")
        self._kpi_dauer    = self._kpi_card("⏱  Ø-Dauer",     "0 min", "#6a1b9a")

        for card in (self._kpi_gesamt, self._kpi_ja, self._kpi_nein, self._kpi_dauer):
            kpi_layout.addWidget(card, 1)
        layout.addWidget(kpi_frame)

        # ── Splitter: linke Tabelle + rechte Stichwort-Tabelle ─────────────
        from PySide6.QtWidgets import QSplitter
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setStyleSheet("QSplitter::handle{background:#ddd; width:2px;}")

        # Monatstabelle
        monat_widget = QWidget()
        ml = QVBoxLayout(monat_widget)
        ml.setContentsMargins(0, 0, 4, 0)
        ml.addWidget(QLabel("📅  Einsätze nach Monat"))
        self._monat_table = QTableWidget()
        self._monat_table.setColumnCount(4)
        self._monat_table.setHorizontalHeaderLabels(
            ["Monat", "Gesamt", "Angenommen", "Abgelehnt"]
        )
        mh = self._monat_table.horizontalHeader()
        mh.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        for c in range(1, 4):
            mh.setSectionResizeMode(c, QHeaderView.ResizeMode.ResizeToContents)
        self._monat_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._monat_table.setAlternatingRowColors(True)
        self._monat_table.verticalHeader().setVisible(False)
        self._monat_table.setStyleSheet("font-size:12px;")
        ml.addWidget(self._monat_table, 1)
        splitter.addWidget(monat_widget)

        # Stichwort-Tabelle
        stw_widget = QWidget()
        sl = QVBoxLayout(stw_widget)
        sl.setContentsMargins(4, 0, 0, 0)
        sl.addWidget(QLabel("🚨  Häufigste Einsatzstichwörter"))
        self._stw_table = QTableWidget()
        self._stw_table.setColumnCount(2)
        self._stw_table.setHorizontalHeaderLabels(["Einsatzstichwort", "Anzahl"])
        sh = self._stw_table.horizontalHeader()
        sh.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        sh.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self._stw_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._stw_table.setAlternatingRowColors(True)
        self._stw_table.verticalHeader().setVisible(False)
        self._stw_table.setStyleSheet("font-size:12px;")
        sl.addWidget(self._stw_table, 1)
        splitter.addWidget(stw_widget)

        splitter.setSizes([400, 400])
        layout.addWidget(splitter, 1)

        # ── Mitarbeiter-Tabelle ────────────────────────────────────────────
        ma_lbl = QLabel("👥  Einsätze nach Mitarbeiter")
        layout.addWidget(ma_lbl)
        self._ma_table = QTableWidget()
        self._ma_table.setColumnCount(3)
        self._ma_table.setHorizontalHeaderLabels(["Mitarbeiter", "Als MA 1", "Als MA 2"])
        mah = self._ma_table.horizontalHeader()
        mah.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        for c in range(1, 3):
            mah.setSectionResizeMode(c, QHeaderView.ResizeMode.ResizeToContents)
        self._ma_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._ma_table.setAlternatingRowColors(True)
        self._ma_table.verticalHeader().setVisible(False)
        self._ma_table.setMaximumHeight(160)
        self._ma_table.setStyleSheet("font-size:12px;")
        layout.addWidget(self._ma_table)

    @staticmethod
    def _kpi_card(label: str, wert: str, farbe: str) -> QFrame:
        card = QFrame()
        card.setMinimumHeight(80)
        card.setStyleSheet(
            f"QFrame{{background:{farbe};border-radius:8px;padding:8px;}}"
        )
        vl = QVBoxLayout(card)
        vl.setSpacing(2)
        lbl = QLabel(label)
        lbl.setStyleSheet("color:rgba(255,255,255,204);font-size:11px;background:transparent;")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        val = QLabel(wert)
        val.setObjectName("val")
        val.setFont(QFont("Arial", 22, QFont.Weight.Bold))
        val.setStyleSheet("color:white;background:transparent;")
        val.setAlignment(Qt.AlignmentFlag.AlignCenter)
        vl.addWidget(lbl)
        vl.addWidget(val)
        return card

    @staticmethod
    def _set_kpi(card: QFrame, wert: str):
        val = card.findChild(QLabel, "val")
        if val:
            val.setText(wert)

    def refresh(self):
        # Jahre in Combo befüllen
        current_j = self._combo_jahr.currentData()
        self._combo_jahr.blockSignals(True)
        self._combo_jahr.clear()
        self._combo_jahr.addItem("Alle", None)
        for j in verfuegbare_jahre_einsaetze():
            self._combo_jahr.addItem(str(j), j)
        for i in range(self._combo_jahr.count()):
            if self._combo_jahr.itemData(i) == current_j:
                self._combo_jahr.setCurrentIndex(i)
                break
        self._combo_jahr.blockSignals(False)

        jahr = self._combo_jahr.currentData()
        alle = lade_einsaetze(jahr=jahr)

        # ── KPIs ───────────────────────────────────────────────────────────
        n = len(alle)
        ang = sum(1 for e in alle if e.get("angenommen", 1))
        nein = n - ang
        dauern = [e["einsatzdauer"] for e in alle if e.get("einsatzdauer")]
        avg_d = round(sum(dauern) / len(dauern)) if dauern else 0

        self._set_kpi(self._kpi_gesamt, str(n))
        self._set_kpi(self._kpi_ja,     str(ang))
        self._set_kpi(self._kpi_nein,   str(nein))
        self._set_kpi(self._kpi_dauer,  f"{avg_d} min")

        # ── Monatstabelle ──────────────────────────────────────────────────
        _MONATE = ["","Januar","Februar","März","April","Mai","Juni",
                   "Juli","August","September","Oktober","November","Dezember"]
        monat_stats: dict[int, dict] = {}
        for e in alle:
            try:
                m = int(e["datum"].split(".")[1])
            except Exception:
                continue
            if m not in monat_stats:
                monat_stats[m] = {"gesamt": 0, "ang": 0, "nein": 0}
            monat_stats[m]["gesamt"] += 1
            if e.get("angenommen", 1):
                monat_stats[m]["ang"] += 1
            else:
                monat_stats[m]["nein"] += 1

        sorted_monate = sorted(monat_stats.keys())
        self._monat_table.setRowCount(len(sorted_monate))
        for row, m in enumerate(sorted_monate):
            s = monat_stats[m]
            self._monat_table.setItem(row, 0, QTableWidgetItem(_MONATE[m]))
            for col, key in enumerate(["gesamt", "ang", "nein"], start=1):
                it = QTableWidgetItem(str(s[key]))
                it.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                if col == 2 and s[key] > 0:
                    it.setBackground(QColor("#e8f5e9"))
                elif col == 3 and s[key] > 0:
                    it.setBackground(QColor("#ffe0e0"))
                self._monat_table.setItem(row, col, it)

        # ── Stichwort-Tabelle ──────────────────────────────────────────────
        from collections import Counter
        stw_counter = Counter(
            e["einsatzstichwort"] for e in alle
            if e.get("einsatzstichwort")
        )
        top_stw = stw_counter.most_common(20)
        self._stw_table.setRowCount(len(top_stw))
        for row, (stw, cnt) in enumerate(top_stw):
            self._stw_table.setItem(row, 0, QTableWidgetItem(stw))
            it = QTableWidgetItem(str(cnt))
            it.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self._stw_table.setItem(row, 1, it)

        # ── Mitarbeiter-Tabelle ────────────────────────────────────────────
        ma1_counter: dict[str, int] = {}
        ma2_counter: dict[str, int] = {}
        for e in alle:
            m1 = (e.get("drk_ma1") or "").strip()
            m2 = (e.get("drk_ma2") or "").strip()
            if m1:
                ma1_counter[m1] = ma1_counter.get(m1, 0) + 1
            if m2:
                ma2_counter[m2] = ma2_counter.get(m2, 0) + 1

        alle_ma = sorted(set(ma1_counter) | set(ma2_counter))
        self._ma_table.setRowCount(len(alle_ma))
        for row, ma in enumerate(alle_ma):
            self._ma_table.setItem(row, 0, QTableWidgetItem(ma))
            for col, counter in enumerate([ma1_counter, ma2_counter], start=1):
                it = QTableWidgetItem(str(counter.get(ma, 0)))
                it.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                self._ma_table.setItem(row, col, it)


# ──────────────────────────────────────────────────────────────────────────────
#  Haupt-Widget: Dienstliches
# ──────────────────────────────────────────────────────────────────────────────

class DienstlichesWidget(QWidget):
    """Haupt-Widget 'Dienstliches' mit Tabs für Einsätze und weitere Protokolle."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # ── Titelleiste ────────────────────────────────────────────────────
        header = QFrame()
        header.setFixedHeight(52)
        header.setStyleSheet(f"background:{FIORI_BLUE}; border:none;")
        hl = QHBoxLayout(header)
        hl.setContentsMargins(20, 0, 20, 0)
        lbl = QLabel("☕  Dienstliches")
        lbl.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        lbl.setStyleSheet("color:white; background:transparent;")
        hl.addWidget(lbl)
        hl.addStretch()

        btn_refresh = QPushButton("🔄")
        btn_refresh.setFixedSize(30, 30)
        btn_refresh.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_refresh.setToolTip("Ansicht aktualisieren")
        btn_refresh.setStyleSheet(
            "QPushButton{background:rgba(255,255,255,38);color:white;"
            "border:1px solid rgba(255,255,255,77);border-radius:4px;}"
            "QPushButton:hover{background:rgba(255,255,255,64);}"
        )
        btn_refresh.clicked.connect(self.refresh)
        hl.addWidget(btn_refresh)
        layout.addWidget(header)

        # ── Tab-Widget ─────────────────────────────────────────────────────
        self._tabs = QTabWidget()
        self._tabs.setDocumentMode(True)
        self._tabs.setStyleSheet("""
            QTabBar::tab {
                padding: 10px 20px;
                font-size: 13px;
                min-width: 120px;
                font-family: 'Segoe UI';
                color: #666;
                background: transparent;
                border-bottom: 3px solid transparent;
                margin-right: 4px;
            }
            QTabBar::tab:selected {
                color: #1565a8;
                font-weight: bold;
                border-bottom: 3px solid #1565a8;
            }
            QTabBar::tab:hover:!selected {
                color: #1565a8;
                border-bottom: 3px solid #ccddf5;
            }
        """)

        self._einsaetze_tab = _EinsaetzeTab()
        self._tabs.addTab(self._einsaetze_tab, "🚑  Einsätze")

        self._patienten_tab = _PatientenTab()
        self._tabs.addTab(self._patienten_tab, "🏥  Patienten DRK Station")

        self._uebersicht_tab = _UebersichtTab()
        self._tabs.addTab(self._uebersicht_tab, "📊  Übersicht")

        self._tabs.currentChanged.connect(self._on_tab_changed)
        layout.addWidget(self._tabs, 1)

    def _on_tab_changed(self, idx: int):
        if idx == 2:
            self._uebersicht_tab.refresh()

    def refresh(self):
        self._einsaetze_tab.refresh()
        self._patienten_tab.refresh()
        if self._tabs.currentIndex() == 2:
            self._uebersicht_tab.refresh()
