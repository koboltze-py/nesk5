"""
Vorkommnisse-Widget
Erfassung und Verwaltung von Vorkommnis-Berichten (PRM, Passagiere, Chronologie).
Export als Word-Dokument mit Firmenlogo und Anschrift.
"""
from __future__ import annotations

import os
import sys
from datetime import datetime, date
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from PySide6.QtCore import Qt, QDate, QTime, QPoint
from PySide6.QtGui import QFont, QColor, QAction
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QFormLayout, QGroupBox,
    QLabel, QLineEdit, QTextEdit, QPushButton, QComboBox,
    QDateEdit, QTimeEdit, QTableWidget, QTableWidgetItem,
    QHeaderView, QAbstractItemView, QScrollArea, QFrame,
    QSplitter, QSizePolicy, QMessageBox, QFileDialog, QTabWidget, QMenu,
    QCheckBox,
)

from config import BASE_DIR, FIORI_BLUE, FIORI_TEXT

# ── Vorkommnis-Typen ───────────────────────────────────────────────────────────
VORKOMMNIS_TYPEN = [
    "PRM-Betreuung",
    "Medizinischer Notfall",
    "Sicherheitsvorfall",
    "Verspätung/Offblock",
    "Fahrzeugschaden",
    "Kommunikationsfehler",
    "Sonstiges",
]

# ── PRM-Kategorien (IATA) ──────────────────────────────────────────────────────
PRM_KATEGORIEN = [
    "WCHS", "WCHR", "WCHC", "BLND", "DEAF", "DPNA",
    "UMNR", "STCR", "MEDA", "Sonstiges",
]

# ── Mitarbeiter-Rollen ─────────────────────────────────────────────────────────
MITARBEITER_ROLLEN = [
    "Disposition",
    "PRM-Begleitung",
    "Fahrer",
    "Servicepoint",
    "Leitung",
    "Sonstiges",
]

# ═══════════════════════════════════════════════════════════════════════════════
#  Stil-Helfer
# ═══════════════════════════════════════════════════════════════════════════════

_BTN_PRIM = f"""
    QPushButton {{
        background-color: {FIORI_BLUE};
        color: white;
        border: none;
        border-radius: 4px;
        padding: 6px 14px;
        font-size: 12px;
        font-weight: bold;
    }}
    QPushButton:hover {{
        background-color: #1a6fc4;
    }}
    QPushButton:disabled {{
        background-color: #aaa;
    }}
"""

_BTN_DANGER = """
    QPushButton {
        background-color: #e74c3c;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 5px 12px;
        font-size: 12px;
    }
    QPushButton:hover {
        background-color: #c0392b;
    }
"""

_BTN_SEC = """
    QPushButton {
        background-color: #4a6480;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 6px 14px;
        font-size: 12px;
    }
    QPushButton:hover {
        background-color: #5a7490;
    }
"""

_TABLE_STYLE = """
    QTableWidget {
        border: 1px solid #d0d8e4;
        border-radius: 4px;
        gridline-color: #e0e6ef;
        background: white;
        font-size: 12px;
    }
    QTableWidget::item:selected {
        background-color: #dbeafe;
        color: #1b3a5c;
    }
    QHeaderView::section {
        background-color: #1565a8;
        color: white;
        font-weight: bold;
        font-size: 12px;
        padding: 5px;
        border: none;
    }
"""

_GROUP_STYLE = """
    QGroupBox {
        font-weight: bold;
        font-size: 13px;
        color: #1b3a5c;
        border: 1px solid #c5d0de;
        border-radius: 6px;
        margin-top: 10px;
        padding-top: 8px;
    }
    QGroupBox::title {
        subcontrol-origin: margin;
        left: 10px;
        padding: 0 4px;
    }
"""

_LINE_EDIT_STYLE = """
    QLineEdit, QTextEdit, QComboBox, QDateEdit, QTimeEdit {
        border: 1px solid #c5d0de;
        border-radius: 4px;
        padding: 4px 8px;
        font-size: 12px;
        background: white;
    }
    QLineEdit:focus, QTextEdit:focus, QComboBox:focus,
    QDateEdit:focus, QTimeEdit:focus {
        border: 1.5px solid #1565a8;
    }
"""


# ═══════════════════════════════════════════════════════════════════════════════
#  Hilfsklasse: Editierbare Tabelle mit Hinzufügen/Löschen
# ═══════════════════════════════════════════════════════════════════════════════

class _EditTable(QWidget):
    """Generische editierbare Tabelle mit +/−-Buttons und Custom-Kontextmenü."""

    def __init__(
        self,
        headers: list[str],
        combo_columns: dict[int, list[str]] | None = None,
        column_widths: dict[int, int] | None = None,
        parent=None,
    ):
        super().__init__(parent)
        self._combo_columns = combo_columns or {}

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(6)

        # Tabelle
        self._table = QTableWidget(0, len(headers))
        self._table.setHorizontalHeaderLabels(headers)
        self._table.setStyleSheet(_TABLE_STYLE)
        self._table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self._table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked
                                    | QAbstractItemView.EditTrigger.SelectedClicked
                                    | QAbstractItemView.EditTrigger.AnyKeyPressed)
        self._table.verticalHeader().setVisible(False)
        self._table.setMinimumHeight(120)

        # Spaltenbreiten
        hdr = self._table.horizontalHeader()
        hdr.setStretchLastSection(True)
        for col in range(len(headers) - 1):
            hdr.setSectionResizeMode(col, QHeaderView.ResizeMode.Interactive)
            w = (column_widths or {}).get(col, 160)
            self._table.setColumnWidth(col, w)

        # Custom-Kontextmenü (ersetzt das OS-Standard-Menü)
        self._table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self._table.customContextMenuRequested.connect(self._show_context_menu)

        layout.addWidget(self._table)

        # Buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(6)
        self._btn_add = QPushButton("+ Hinzufügen")
        self._btn_add.setStyleSheet(_BTN_PRIM)
        self._btn_add.clicked.connect(self._add_row)
        self._btn_del = QPushButton("− Löschen")
        self._btn_del.setStyleSheet(_BTN_DANGER)
        self._btn_del.clicked.connect(self._del_row)
        btn_row.addWidget(self._btn_add)
        btn_row.addWidget(self._btn_del)
        btn_row.addStretch()
        layout.addLayout(btn_row)

    def _show_context_menu(self, pos: QPoint):
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background: white; border: 1px solid #c5d0de;
                font-size: 12px;
            }
            QMenu::item { padding: 6px 20px; }
            QMenu::item:selected { background: #dbeafe; color: #1b3a5c; }
            QMenu::separator { height: 1px; background: #e0e6ef; margin: 2px 8px; }
        """)
        act_add  = QAction("✚  Zeile hinzufügen", self)
        act_del  = QAction("✖  Ausgewählte Zeile löschen", self)
        act_up   = QAction("↑  Zeile nach oben", self)
        act_down = QAction("↓  Zeile nach unten", self)
        act_add.triggered.connect(self._add_row)
        act_del.triggered.connect(self._del_row)
        act_up.triggered.connect(self._move_up)
        act_down.triggered.connect(self._move_down)
        menu.addAction(act_add)
        menu.addAction(act_del)
        menu.addSeparator()
        menu.addAction(act_up)
        menu.addAction(act_down)
        menu.exec(self._table.viewport().mapToGlobal(pos))

    def _move_up(self):
        row = self._table.currentRow()
        if row <= 0:
            return
        self._swap_rows(row, row - 1)
        self._table.selectRow(row - 1)

    def _move_down(self):
        row = self._table.currentRow()
        if row < 0 or row >= self._table.rowCount() - 1:
            return
        self._swap_rows(row, row + 1)
        self._table.selectRow(row + 1)

    def _swap_rows(self, a: int, b: int):
        cols = self._table.columnCount()
        for col in range(cols):
            wa = self._table.cellWidget(a, col)
            wb = self._table.cellWidget(b, col)
            if wa and wb and isinstance(wa, QComboBox) and isinstance(wb, QComboBox):
                va, vb = wa.currentText(), wb.currentText()
                ia = wa.findText(vb); wa.setCurrentIndex(ia if ia >= 0 else 0)
                ib = wb.findText(va); wb.setCurrentIndex(ib if ib >= 0 else 0)
            else:
                ia_item = self._table.item(a, col)
                ib_item = self._table.item(b, col)
                ta = ia_item.text() if ia_item else ""
                tb = ib_item.text() if ib_item else ""
                self._table.setItem(a, col, QTableWidgetItem(tb))
                self._table.setItem(b, col, QTableWidgetItem(ta))

    def _add_row(self):
        row = self._table.rowCount()
        self._table.insertRow(row)
        for col in range(self._table.columnCount()):
            if col in self._combo_columns:
                combo = QComboBox()
                combo.addItems(self._combo_columns[col])
                self._table.setCellWidget(row, col, combo)
            else:
                self._table.setItem(row, col, QTableWidgetItem(""))
        self._table.scrollToBottom()
        self._table.selectRow(row)

    def _del_row(self):
        rows = sorted(
            {idx.row() for idx in self._table.selectedIndexes()},
            reverse=True,
        )
        for row in rows:
            self._table.removeRow(row)

    def get_data(self) -> list[list[str]]:
        result = []
        for row in range(self._table.rowCount()):
            cols = []
            for col in range(self._table.columnCount()):
                widget = self._table.cellWidget(row, col)
                if isinstance(widget, QComboBox):
                    cols.append(widget.currentText())
                else:
                    item = self._table.item(row, col)
                    cols.append(item.text().strip() if item else "")
            result.append(cols)
        return result

    def set_data(self, data: list[list[str]]):
        self._table.setRowCount(0)
        for row_data in data:
            row = self._table.rowCount()
            self._table.insertRow(row)
            for col, value in enumerate(row_data):
                if col in self._combo_columns:
                    combo = QComboBox()
                    combo.addItems(self._combo_columns[col])
                    idx = combo.findText(value)
                    if idx >= 0:
                        combo.setCurrentIndex(idx)
                    self._table.setCellWidget(row, col, combo)
                else:
                    self._table.setItem(row, col, QTableWidgetItem(value))


# ═══════════════════════════════════════════════════════════════════════════════
#  Haupt-Widget: VorkommnisseWidget
# ═══════════════════════════════════════════════════════════════════════════════

class VorkommnisseWidget(QWidget):

    def __init__(self, parent=None):
        super().__init__(parent)
        self._current_id: int | None = None   # None = neuer Datensatz
        self.setStyleSheet("background-color: #f5f6f7;")
        self._build_ui()
        self._aktualisiere_liste()

    # ── UI-Aufbau ──────────────────────────────────────────────────────────────

    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Header-Leiste
        header = QWidget()
        header.setFixedHeight(56)
        header.setStyleSheet(f"background-color: {FIORI_BLUE};")
        hdr_layout = QHBoxLayout(header)
        hdr_layout.setContentsMargins(20, 0, 20, 0)
        title_lbl = QLabel("⚠️  Vorkommnisse")
        title_lbl.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        title_lbl.setStyleSheet("color: white;")
        hdr_layout.addWidget(title_lbl)
        hdr_layout.addStretch()
        self._btn_neu = QPushButton("🔄  Neues Formular")
        self._btn_neu.setStyleSheet(_BTN_SEC)
        self._btn_neu.clicked.connect(self._neues_formular)
        self._btn_speichern = QPushButton("💾  Speichern")
        self._btn_speichern.setStyleSheet(_BTN_PRIM)
        self._btn_speichern.clicked.connect(self._speichern)
        self._btn_export = QPushButton("📄  Word exportieren")
        self._btn_export.setStyleSheet(_BTN_SEC)
        self._btn_export.clicked.connect(self._export_word)
        self._btn_ordner = QPushButton("📁  Berichte öffnen")
        self._btn_ordner.setStyleSheet(_BTN_SEC)
        self._btn_ordner.setToolTip("Ordner mit allen Vorkommnis-Berichten öffnen")
        self._btn_ordner.clicked.connect(self._ordner_oeffnen)
        hdr_layout.addWidget(self._btn_neu)
        hdr_layout.addWidget(self._btn_speichern)
        hdr_layout.addWidget(self._btn_export)
        hdr_layout.addWidget(self._btn_ordner)
        root.addWidget(header)

        # Tab-Widget
        self._tabs = QTabWidget()
        self._tabs.setStyleSheet("""
            QTabWidget::pane { border: none; background: #f5f6f7; }
            QTabBar::tab {
                background: #354a5e; color: #cdd5e0;
                padding: 8px 22px; font-size: 13px;
                border: none; margin-right: 2px;
            }
            QTabBar::tab:selected { background: #1565a8; color: white; font-weight: bold; }
            QTabBar::tab:hover    { background: #4a6480; color: white; }
        """)
        self._tabs.addTab(self._build_erfassung_tab(), "✏️  Erfassung")
        self._tabs.addTab(self._build_liste_tab(),     "📋  Gespeicherte Vorkommnisse")
        self._tabs.currentChanged.connect(self._on_tab_changed)
        root.addWidget(self._tabs)

    # ── Tab 1: Erfassung ───────────────────────────────────────────────────────

    def _build_erfassung_tab(self) -> QWidget:
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; background: #f5f6f7; }")
        content = QWidget()
        content.setStyleSheet("background: #f5f6f7;")
        c_layout = QVBoxLayout(content)
        c_layout.setContentsMargins(20, 16, 20, 20)
        c_layout.setSpacing(14)

        self._status_lbl = QLabel("🟢  Neuer Datensatz")
        self._status_lbl.setStyleSheet(
            "color: #27ae60; font-size: 12px; font-weight: bold;"
        )
        c_layout.addWidget(self._status_lbl)

        c_layout.addWidget(self._build_grunddaten())

        grp_pax = QGroupBox("1. Betroffene Passagiere")
        grp_pax.setStyleSheet(_GROUP_STYLE)
        pax_layout = QVBoxLayout(grp_pax)
        pax_layout.setContentsMargins(10, 14, 10, 10)
        self._pax_table = _EditTable(
            headers=["Passagier (Name, Vorname)", "Kategorie", "Anmerkung"],
            combo_columns={1: PRM_KATEGORIEN},
            column_widths={0: 220, 1: 110},
        )
        pax_layout.addWidget(self._pax_table)
        c_layout.addWidget(grp_pax)

        grp_personal = QGroupBox("2. Eingeteiltes Personal")
        grp_personal.setStyleSheet(_GROUP_STYLE)
        pers_layout = QVBoxLayout(grp_personal)
        pers_layout.setContentsMargins(10, 14, 10, 10)
        self._personal_table = _EditTable(
            headers=["Name (Mitarbeiter)", "Funktion/Rolle", "Anmerkung"],
            combo_columns={1: MITARBEITER_ROLLEN},
            column_widths={0: 200, 1: 150},
        )
        pers_layout.addWidget(self._personal_table)
        c_layout.addWidget(grp_personal)

        grp_chrono = QGroupBox("3. Chronologischer Ablauf")
        grp_chrono.setStyleSheet(_GROUP_STYLE)
        chrono_layout = QVBoxLayout(grp_chrono)
        chrono_layout.setContentsMargins(10, 14, 10, 10)
        self._chrono_table = _EditTable(
            headers=["Uhrzeit", "Ereignis"],
            column_widths={0: 110},
        )
        chrono_layout.addWidget(self._chrono_table)
        c_layout.addWidget(grp_chrono)

        grp_ursache = QGroupBox("4. Ursachenanalyse")
        grp_ursache.setStyleSheet(_GROUP_STYLE)
        ursache_layout = QVBoxLayout(grp_ursache)
        ursache_layout.setContentsMargins(10, 14, 10, 10)
        self._ursache_edit = QTextEdit()
        self._ursache_edit.setPlaceholderText(
            "Ursachen des Vorkommnisses stichpunktartig beschreiben ..."
        )
        self._ursache_edit.setMinimumHeight(100)
        self._ursache_edit.setStyleSheet(_LINE_EDIT_STYLE)
        ursache_layout.addWidget(self._ursache_edit)
        c_layout.addWidget(grp_ursache)

        grp_ergebnis = QGroupBox("5. Ergebnis")
        grp_ergebnis.setStyleSheet(_GROUP_STYLE)
        erg_layout = QVBoxLayout(grp_ergebnis)
        erg_layout.setContentsMargins(10, 14, 10, 10)
        self._ergebnis_edit = QTextEdit()
        self._ergebnis_edit.setPlaceholderText(
            "Ergebnis und Ausgang des Vorkommnisses ..."
        )
        self._ergebnis_edit.setMinimumHeight(80)
        self._ergebnis_edit.setStyleSheet(_LINE_EDIT_STYLE)
        erg_layout.addWidget(self._ergebnis_edit)
        c_layout.addWidget(grp_ergebnis)

        c_layout.addStretch()
        scroll.setWidget(content)
        return scroll

    _MONATE = [
        "Januar", "Februar", "März", "April", "Mai", "Juni",
        "Juli", "August", "September", "Oktober", "November", "Dezember",
    ]

    # ── Tab 2: Übersicht ───────────────────────────────────────────────────────

    def _build_liste_tab(self) -> QWidget:
        from datetime import date as _date
        from PySide6.QtWidgets import QSpinBox
        _heute = _date.today()

        w = QWidget()
        w.setStyleSheet("background: #f5f6f7;")
        layout = QVBoxLayout(w)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # ── Filter-Leiste ─────────────────────────────────────────────────────
        filter_row = QHBoxLayout()
        filter_row.setSpacing(8)

        lbl_monat = QLabel("Monat:")
        lbl_monat.setStyleSheet("font-weight: bold; color: #1b3a5c;")
        self._filter_monat = QComboBox()
        self._filter_monat.addItem("Alle Monate", 0)
        for i, name in enumerate(self._MONATE, start=1):
            self._filter_monat.addItem(name, i)
        self._filter_monat.setCurrentIndex(_heute.month)  # 1-12 → Index 1-12
        self._filter_monat.setFixedWidth(140)
        self._filter_monat.setStyleSheet(_LINE_EDIT_STYLE)

        lbl_jahr = QLabel("Jahr:")
        lbl_jahr.setStyleSheet("font-weight: bold; color: #1b3a5c;")
        self._filter_jahr = QSpinBox()
        self._filter_jahr.setRange(2020, 2040)
        self._filter_jahr.setValue(_heute.year)
        self._filter_jahr.setFixedWidth(80)
        self._filter_jahr.setStyleSheet(_LINE_EDIT_STYLE)

        btn_alle = QPushButton("🔄  Alle anzeigen")
        btn_alle.setStyleSheet(_BTN_SEC)
        btn_alle.setToolTip("Monatsfilter aufheben – alle Vorkommnisse anzeigen")
        btn_alle.clicked.connect(self._filter_alle_anzeigen)

        self._filter_zaehler = QLabel("")
        self._filter_zaehler.setStyleSheet("color: #555; font-size: 11px;")

        filter_row.addWidget(lbl_monat)
        filter_row.addWidget(self._filter_monat)
        filter_row.addSpacing(12)
        filter_row.addWidget(lbl_jahr)
        filter_row.addWidget(self._filter_jahr)
        filter_row.addSpacing(8)
        filter_row.addWidget(btn_alle)
        filter_row.addStretch()
        filter_row.addWidget(self._filter_zaehler)
        layout.addLayout(filter_row)

        self._filter_monat.currentIndexChanged.connect(self._aktualisiere_liste)
        self._filter_jahr.valueChanged.connect(self._aktualisiere_liste)

        # ── Toolbar ───────────────────────────────────────────────────────────
        toolbar = QHBoxLayout()
        btn_laden = QPushButton("📥  Öffnen / Bearbeiten")
        btn_laden.setStyleSheet(_BTN_PRIM)
        btn_laden.setToolTip("Ausgewähltes Vorkommnis in den Erfassungs-Tab laden")
        btn_laden.clicked.connect(self._laden_aus_liste)
        btn_loeschen = QPushButton("🗑  Löschen")
        btn_loeschen.setStyleSheet(_BTN_DANGER)
        btn_loeschen.setToolTip("Ausgewähltes Vorkommnis dauerhaft löschen")
        btn_loeschen.clicked.connect(self._loeschen_aus_liste)
        toolbar.addWidget(btn_laden)
        toolbar.addWidget(btn_loeschen)
        toolbar.addStretch()
        layout.addLayout(toolbar)

        self._liste_table = QTableWidget(0, 6)
        self._liste_table.setHorizontalHeaderLabels(
            ["ID", "Flug", "Typ", "Datum", "Ort", "Erstellt von"]
        )
        self._liste_table.setStyleSheet(_TABLE_STYLE)
        self._liste_table.setSelectionBehavior(
            QAbstractItemView.SelectionBehavior.SelectRows
        )
        self._liste_table.setEditTriggers(
            QAbstractItemView.EditTrigger.NoEditTriggers
        )
        self._liste_table.verticalHeader().setVisible(False)
        self._liste_table.horizontalHeader().setStretchLastSection(True)
        self._liste_table.setColumnWidth(0, 45)
        self._liste_table.setColumnWidth(1, 90)
        self._liste_table.setColumnWidth(2, 160)
        self._liste_table.setColumnWidth(3, 100)
        self._liste_table.setColumnWidth(4, 160)
        self._liste_table.doubleClicked.connect(self._laden_aus_liste)
        self._liste_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self._liste_table.customContextMenuRequested.connect(
            self._liste_context_menu
        )
        layout.addWidget(self._liste_table)
        return w

    def _liste_context_menu(self, pos: QPoint):
        menu = QMenu(self)
        menu.setStyleSheet(
            "QMenu { background: white; border: 1px solid #c5d0de; font-size: 12px; }"
            "QMenu::item { padding: 6px 20px; }"
            "QMenu::item:selected { background: #dbeafe; color: #1b3a5c; }"
        )
        act_open = QAction("📥  Öffnen / Bearbeiten", self)
        act_del  = QAction("🗑  Löschen", self)
        act_open.triggered.connect(self._laden_aus_liste)
        act_del.triggered.connect(self._loeschen_aus_liste)
        menu.addAction(act_open)
        menu.addAction(act_del)
        menu.exec(self._liste_table.viewport().mapToGlobal(pos))

    def _build_grunddaten(self) -> QGroupBox:
        grp = QGroupBox("Grunddaten des Vorkommnisses")
        grp.setStyleSheet(_GROUP_STYLE)
        layout = QFormLayout(grp)
        layout.setContentsMargins(14, 18, 14, 12)
        layout.setSpacing(10)
        layout.setLabelAlignment(Qt.AlignmentFlag.AlignRight)
        layout.setFieldGrowthPolicy(QFormLayout.FieldGrowthPolicy.ExpandingFieldsGrow)

        style = _LINE_EDIT_STYLE

        self._flug_edit = QLineEdit()
        self._flug_edit.setPlaceholderText("z. B. XQ983")
        self._flug_edit.setStyleSheet(style)
        layout.addRow("Flugnummer:", self._flug_edit)

        self._typ_combo = QComboBox()
        self._typ_combo.addItems(VORKOMMNIS_TYPEN)
        self._typ_combo.setStyleSheet(style)
        layout.addRow("Vorkommnis-Typ:", self._typ_combo)

        self._datum_edit = QDateEdit(QDate.currentDate())
        self._datum_edit.setDisplayFormat("dd.MM.yyyy")
        self._datum_edit.setCalendarPopup(True)
        self._datum_edit.setStyleSheet(style)
        layout.addRow("Datum:", self._datum_edit)

        self._ort_edit = QLineEdit()
        self._ort_edit.setPlaceholderText("z. B. Köln/Bonn (CGN)")
        self._ort_edit.setStyleSheet(style)
        layout.addRow("Ort:", self._ort_edit)

        self._offblock_plan_edit = QTimeEdit()
        self._offblock_plan_edit.setDisplayFormat("HH:mm")
        self._offblock_plan_edit.setStyleSheet(style)
        self._offblock_plan_edit.setEnabled(False)
        self._offblock_plan_chk = QCheckBox("angeben")
        self._offblock_plan_chk.setChecked(False)
        self._offblock_plan_chk.toggled.connect(self._offblock_plan_edit.setEnabled)
        _plan_row = QWidget()
        _plan_hl  = QHBoxLayout(_plan_row)
        _plan_hl.setContentsMargins(0, 0, 0, 0)
        _plan_hl.addWidget(self._offblock_plan_edit)
        _plan_hl.addWidget(self._offblock_plan_chk)
        layout.addRow("Geplanter Offblock:", _plan_row)

        self._offblock_ist_edit = QTimeEdit()
        self._offblock_ist_edit.setDisplayFormat("HH:mm")
        self._offblock_ist_edit.setStyleSheet(style)
        self._offblock_ist_edit.setEnabled(False)
        self._offblock_ist_chk = QCheckBox("angeben")
        self._offblock_ist_chk.setChecked(False)
        self._offblock_ist_chk.toggled.connect(self._offblock_ist_edit.setEnabled)
        _ist_row = QWidget()
        _ist_hl  = QHBoxLayout(_ist_row)
        _ist_hl.setContentsMargins(0, 0, 0, 0)
        _ist_hl.addWidget(self._offblock_ist_edit)
        _ist_hl.addWidget(self._offblock_ist_chk)
        layout.addRow("Tatsächlicher Offblock:", _ist_row)

        self._erstellt_von_edit = QLineEdit()
        self._erstellt_von_edit.setPlaceholderText("Name des Verfassers")
        self._erstellt_von_edit.setStyleSheet(style)
        layout.addRow("Erstellt von:", self._erstellt_von_edit)

        return grp

    # ── DB-Aktionen ────────────────────────────────────────────────────────────

    def _speichern(self):
        from functions.vorkommnisse_db import speichern, aktualisieren
        daten = self._sammle_daten()
        if not daten["flug"]:
            QMessageBox.warning(self, "Fehler", "Bitte Flugnummer eingeben.")
            return
        try:
            if self._current_id is None:
                self._current_id = speichern(daten)
                QMessageBox.information(
                    self, "Gespeichert",
                    f"Vorkommnis Flug {daten['flug']} wurde gespeichert."
                )
            else:
                aktualisieren(self._current_id, daten)
                QMessageBox.information(
                    self, "Aktualisiert",
                    f"Vorkommnis #{self._current_id} wurde aktualisiert."
                )
            self._aktualisiere_status()
            self._aktualisiere_liste()
        except Exception as exc:
            QMessageBox.critical(self, "Fehler beim Speichern", str(exc))

    def _neues_formular(self):
        self._current_id = None
        self._formular_leeren()
        self._aktualisiere_status()
        self._tabs.setCurrentIndex(0)

    def _laden_aus_liste(self):
        from functions.vorkommnisse_db import lade_ein
        row = self._liste_table.currentRow()
        if row < 0:
            QMessageBox.information(self, "Hinweis", "Bitte zuerst ein Vorkommnis auswählen.")
            return
        id_item = self._liste_table.item(row, 0)
        if not id_item:
            return
        vid = int(id_item.text())
        daten = lade_ein(vid)
        if daten is None:
            QMessageBox.warning(self, "Fehler", "Datensatz nicht gefunden.")
            return
        self._current_id = vid
        self._formular_befuellen(daten)
        self._aktualisiere_status()
        self._tabs.setCurrentIndex(0)

    def _loeschen_aus_liste(self):
        from functions.vorkommnisse_db import loeschen
        row = self._liste_table.currentRow()
        if row < 0:
            QMessageBox.information(self, "Hinweis", "Bitte zuerst ein Vorkommnis auswählen.")
            return
        id_item  = self._liste_table.item(row, 0)
        flg_item = self._liste_table.item(row, 1)
        if not id_item:
            return
        vid  = int(id_item.text())
        flug = flg_item.text() if flg_item else str(vid)
        antwort = QMessageBox.question(
            self, "Löschen bestätigen",
            f"Vorkommnis Flug »{flug}« wirklich löschen?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if antwort != QMessageBox.StandardButton.Yes:
            return
        loeschen(vid)
        if self._current_id == vid:
            self._current_id = None
            self._formular_leeren()
            self._aktualisiere_status()
        self._aktualisiere_liste()

    def _filter_alle_anzeigen(self):
        """Monatsfilter auf \"Alle\" setzen."""
        self._filter_monat.setCurrentIndex(0)   # "Alle Monate"

    def _aktualisiere_liste(self):
        try:
            from functions.vorkommnisse_db import lade_alle
            alle = lade_alle()
        except Exception:
            alle = []

        # Aktiven Filter auslesen
        sel_monat = self._filter_monat.currentData()   # 0 = alle, 1-12 = Monat
        sel_jahr  = self._filter_jahr.value()

        gefiltert = []
        for d in alle:
            if sel_monat == 0:
                gefiltert.append(d)
                continue
            # Datum im Format dd.MM.yyyy parsen
            try:
                from datetime import datetime as _dt
                dat = _dt.strptime(d["datum"], "%d.%m.%Y")
                if dat.month == sel_monat and dat.year == sel_jahr:
                    gefiltert.append(d)
            except ValueError:
                pass  # unbekanntes Datumsformat → überspringen

        # Zähler aktualisieren
        gesamt = len(alle)
        angezeigt = len(gefiltert)
        if sel_monat == 0:
            self._filter_zaehler.setText(f"{gesamt} Einträge")
        else:
            monat_name = self._MONATE[sel_monat - 1]
            self._filter_zaehler.setText(
                f"{angezeigt} von {gesamt} • {monat_name} {sel_jahr}"
            )

        self._liste_table.setRowCount(0)
        for d in gefiltert:
            r = self._liste_table.rowCount()
            self._liste_table.insertRow(r)
            for col, val in enumerate([
                str(d["id"]), d["flug"], d["typ"],
                d["datum"], d["ort"], d["erstellt_von"],
            ]):
                item = QTableWidgetItem(val)
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                if d["id"] == self._current_id:
                    item.setBackground(QColor("#dbeafe"))
                self._liste_table.setItem(r, col, item)

    def _aktualisiere_status(self):
        if self._current_id is None:
            self._status_lbl.setText("🟢  Neuer Datensatz")
            self._status_lbl.setStyleSheet(
                "color: #27ae60; font-size: 12px; font-weight: bold;"
            )
        else:
            self._status_lbl.setText(f"🟡  Bearbeitung: Datensatz #{self._current_id}")
            self._status_lbl.setStyleSheet(
                "color: #e67e22; font-size: 12px; font-weight: bold;"
            )

    def _on_tab_changed(self, idx: int):
        if idx == 1:
            self._aktualisiere_liste()

    # ── Formular-Helfer ────────────────────────────────────────────────────────

    def _formular_leeren(self):
        self._flug_edit.clear()
        self._typ_combo.setCurrentIndex(0)
        self._datum_edit.setDate(QDate.currentDate())
        self._ort_edit.clear()
        self._offblock_plan_chk.setChecked(False)
        self._offblock_plan_edit.setTime(QTime(0, 0))
        self._offblock_ist_chk.setChecked(False)
        self._offblock_ist_edit.setTime(QTime(0, 0))
        self._erstellt_von_edit.clear()
        self._pax_table.set_data([])
        self._personal_table.set_data([])
        self._chrono_table.set_data([])
        self._ursache_edit.clear()
        self._ergebnis_edit.clear()

    def _formular_befuellen(self, d: dict):
        self._flug_edit.setText(d.get("flug", ""))
        idx = self._typ_combo.findText(d.get("typ", ""))
        if idx >= 0:
            self._typ_combo.setCurrentIndex(idx)
        qdate = QDate.fromString(d.get("datum", ""), "dd.MM.yyyy")
        if qdate.isValid():
            self._datum_edit.setDate(qdate)
        self._ort_edit.setText(d.get("ort", ""))
        def _pt(s: str) -> QTime:
            t = QTime.fromString(s.split(" ")[0], "HH:mm")
            return t if t.isValid() else QTime(0, 0)
        _plan_raw = d.get("offblock_plan", "")
        _ist_raw  = d.get("offblock_ist",  "")
        _plan_has = bool(_plan_raw and _plan_raw not in ("", "00:00", "00:00 Uhr"))
        _ist_has  = bool(_ist_raw  and _ist_raw  not in ("", "00:00", "00:00 Uhr"))
        self._offblock_plan_chk.setChecked(_plan_has)
        self._offblock_plan_edit.setTime(_pt(_plan_raw) if _plan_has else QTime(0, 0))
        self._offblock_ist_chk.setChecked(_ist_has)
        self._offblock_ist_edit.setTime(_pt(_ist_raw) if _ist_has else QTime(0, 0))
        self._erstellt_von_edit.setText(d.get("erstellt_von", ""))
        self._pax_table.set_data(d.get("passagiere", []))
        self._personal_table.set_data(d.get("personal", []))
        self._chrono_table.set_data(d.get("chronologie", []))
        self._ursache_edit.setPlainText(d.get("ursache", ""))
        self._ergebnis_edit.setPlainText(d.get("ergebnis", ""))

    def _sammle_daten(self) -> dict:
        if self._offblock_plan_chk.isChecked():
            plan = self._offblock_plan_edit.time().toString("HH:mm")
            plan_min = (
                self._offblock_plan_edit.time().hour() * 60
                + self._offblock_plan_edit.time().minute()
            )
            offblock_plan_str = f"{plan} Uhr"
        else:
            plan_min = 0
            offblock_plan_str = ""

        if self._offblock_ist_chk.isChecked():
            ist = self._offblock_ist_edit.time().toString("HH:mm")
            ist_min = (
                self._offblock_ist_edit.time().hour() * 60
                + self._offblock_ist_edit.time().minute()
            )
            if self._offblock_plan_chk.isChecked():
                versp = ist_min - plan_min
                if versp < 0:
                    versp += 1440
                offblock_ist_str = f"{ist} Uhr"
                if versp > 0:
                    offblock_ist_str += f" (+{versp} Min.)"
            else:
                offblock_ist_str = f"{ist} Uhr"
        else:
            offblock_ist_str = ""
        return {
            "flug":          self._flug_edit.text().strip(),
            "typ":           self._typ_combo.currentText(),
            "datum":         self._datum_edit.date().toString("dd.MM.yyyy"),
            "ort":           self._ort_edit.text().strip(),
            "offblock_plan": offblock_plan_str,
            "offblock_ist":  offblock_ist_str,
            "erstellt_von":  self._erstellt_von_edit.text().strip(),
            "passagiere":    self._pax_table.get_data(),
            "personal":      self._personal_table.get_data(),
            "chronologie":   self._chrono_table.get_data(),
            "ursache":       self._ursache_edit.toPlainText().strip(),
            "ergebnis":      self._ergebnis_edit.toPlainText().strip(),
        }

    # ── Word-Export ────────────────────────────────────────────────────────────

    def _berichte_basis_dir(self) -> Path:
        """Basisordner für alle Vorkommnis-Berichte."""
        return Path(BASE_DIR) / "Daten" / "Vorkommnis Berichte"

    def _monats_dir(self, datum: str) -> Path:
        """Gibt den Monatsordner zurück und legt ihn bei Bedarf an.
        datum erwartet Format 'dd.MM.yyyy'."""
        try:
            d = datetime.strptime(datum, "%d.%m.%Y")
        except ValueError:
            d = datetime.today()
        # Ordnername z. B. "2026-04 April"
        monat_name = d.strftime("%B")  # lokalisierter Monatsname
        ordner = self._berichte_basis_dir() / f"{d.year}-{d.month:02d} {monat_name}"
        ordner.mkdir(parents=True, exist_ok=True)
        return ordner

    def _ordner_oeffnen(self):
        """Öffnet den Berichte-Basisordner im Windows Explorer."""
        basis = self._berichte_basis_dir()
        basis.mkdir(parents=True, exist_ok=True)
        os.startfile(str(basis))

    def _export_word(self):
        daten = self._sammle_daten()
        if not daten["flug"]:
            QMessageBox.warning(self, "Fehler", "Bitte Flugnummer eingeben.")
            return

        default_name = f"Vorkommnisbericht_{daten['flug'].replace('/', '-')}.docx"
        # Monatsordner basierend auf dem Datum des Vorkommnisses
        speicher_dir = self._monats_dir(daten["datum"])
        pfad, _ = QFileDialog.getSaveFileName(
            self, "Word-Dokument speichern",
            str(speicher_dir / default_name),
            "Word-Dokument (*.docx)",
        )
        if not pfad:
            return
        try:
            self._erstelle_word(pfad, daten)
            QMessageBox.information(
                self, "Exportiert",
                f"Vorkommnisbericht gespeichert:\n{pfad}",
            )
            os.startfile(pfad)
        except ImportError:
            QMessageBox.critical(
                self, "Fehler",
                "python-docx nicht installiert.\nBitte ausführen: pip install python-docx",
            )
        except Exception as exc:
            QMessageBox.critical(self, "Export-Fehler", str(exc))

    def _erstelle_word(self, pfad: str, d: dict):
        """Erstellt das Word-Dokument mit Kopf-/Fußzeile und allen Abschnitten."""
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Seitenränder ──────────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # ── Kopfzeile ─────────────────────────────────────────────────────────
        self._baue_kopfzeile(doc)

        # ── Titel ─────────────────────────────────────────────────────────────
        titel = doc.add_heading(
            f"Vorkommnis Bericht – {d['typ']} Flug {d['flug']}", level=1
        )
        titel.runs[0].font.color.rgb = RGBColor(0x15, 0x65, 0xa8)

        # ── Grunddaten-Tabelle ────────────────────────────────────────────────
        grund_daten = [
            ("Flug",   d["flug"]),
            ("Datum",  d["datum"]),
            ("Ort",    d["ort"]),
        ]
        if d.get("offblock_plan"):
            grund_daten.append(("Geplanter Offblock", d["offblock_plan"]))
        if d.get("offblock_ist"):
            grund_daten.append(("Tatsächlicher Offblock", d["offblock_ist"]))

        tbl_grund = doc.add_table(rows=len(grund_daten), cols=2)
        tbl_grund.style = "Table Grid"
        for i, (key, val) in enumerate(grund_daten):
            tbl_grund.rows[i].cells[0].text = key
            tbl_grund.rows[i].cells[1].text = val
            tbl_grund.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

        # ── 1. Passagiere ─────────────────────────────────────────────────────
        doc.add_heading("1. Betroffene Passagiere", level=2)
        if d["passagiere"]:
            tbl_pax = doc.add_table(rows=1, cols=3)
            tbl_pax.style = "Table Grid"
            for i, hdr in enumerate(["Passagier", "Kategorie", "Anmerkung"]):
                cell = tbl_pax.rows[0].cells[i]
                cell.text = hdr
                cell.paragraphs[0].runs[0].bold = True
            for row_data in d["passagiere"]:
                row = tbl_pax.add_row()
                for i, val in enumerate(row_data[:3]):
                    row.cells[i].text = val
        else:
            doc.add_paragraph("(keine Passagiere erfasst)")
        doc.add_paragraph()

        # ── 2. Personal ───────────────────────────────────────────────────────
        doc.add_heading("2. Eingeteiltes Personal", level=2)
        if d["personal"]:
            tbl_pers = doc.add_table(rows=1, cols=3)
            tbl_pers.style = "Table Grid"
            for i, hdr in enumerate(["Mitarbeiter", "Funktion/Rolle", "Anmerkung"]):
                cell = tbl_pers.rows[0].cells[i]
                cell.text = hdr
                cell.paragraphs[0].runs[0].bold = True
            for row_data in d["personal"]:
                row = tbl_pers.add_row()
                for i, val in enumerate(row_data[:3]):
                    row.cells[i].text = val
        else:
            doc.add_paragraph("(kein Personal erfasst)")
        doc.add_paragraph()

        # ── 3. Chronologischer Ablauf ─────────────────────────────────────────
        doc.add_heading("3. Chronologischer Ablauf", level=2)
        if d["chronologie"]:
            tbl_chr = doc.add_table(rows=1, cols=2)
            tbl_chr.style = "Table Grid"
            for i, hdr in enumerate(["Uhrzeit", "Ereignis"]):
                cell = tbl_chr.rows[0].cells[i]
                cell.text = hdr
                cell.paragraphs[0].runs[0].bold = True
            for row_data in d["chronologie"]:
                row = tbl_chr.add_row()
                for i, val in enumerate(row_data[:2]):
                    row.cells[i].text = val
        else:
            doc.add_paragraph("(kein Ablauf erfasst)")
        doc.add_paragraph()

        # ── 4. Ursachenanalyse ────────────────────────────────────────────────
        doc.add_heading("4. Ursachenanalyse", level=2)
        doc.add_paragraph(d["ursache"] or "(keine Angabe)")
        doc.add_paragraph()

        # ── 5. Ergebnis ───────────────────────────────────────────────────────
        doc.add_heading("5. Ergebnis", level=2)
        doc.add_paragraph(d["ergebnis"] or "(keine Angabe)")
        doc.add_paragraph()

        # ── Unterschrift ──────────────────────────────────────────────────────
        doc.add_paragraph()
        p_sign = doc.add_paragraph(f"Erstellt von: {d['erstellt_von'] or '________________'}")
        p_sign.add_run(f"        Datum: {d['datum']}")

        # ── Fußzeile ──────────────────────────────────────────────────────────
        self._baue_fusszeile(doc)

        doc.save(pfad)

    def _baue_kopfzeile(self, doc):
        """Firmenlogo + Anschrift in der Kopfzeile."""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for section in doc.sections:
            header = section.header
            # Erste Zeile: Firmenname fett
            p1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p1.clear()
            run = p1.add_run("Deutsches Rotes Kreuz – Erste-Hilfe-Station Flughafen Köln/Bonn")
            run.bold = True
            run.font.size = Pt(11)
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Zweite Zeile: Adresse
            p2 = header.add_paragraph()
            run2 = p2.add_run("Flughafen Köln/Bonn (CGN)  |  Kennedystraße  |  51147 Köln")
            run2.font.size = Pt(9)
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Logo versuchen einzubinden (optional, kein Fehler wenn nicht vorhanden)
            _base = Path(BASE_DIR) if not isinstance(BASE_DIR, Path) else BASE_DIR
            logo_pfad = _base / "json" / "logo.png"
            if not logo_pfad.exists():
                logo_pfad = _base / "gui" / "logo.png"
            if logo_pfad.exists():
                try:
                    p_logo = header.add_paragraph()
                    run_logo = p_logo.add_run()
                    run_logo.add_picture(str(logo_pfad), width=Cm(3))
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                except Exception:
                    pass  # Logo-Einbindung ist optional

    def _baue_fusszeile(self, doc):
        """Seitenzahl und Hinweis in der Fußzeile."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Text links
            run_links = p.add_run("Vertraulich – DRK Erste-Hilfe-Station CGN    |    Seite ")
            run_links.font.size = Pt(8)

            # Automatische Seitenzahl (XML-Feld)
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            instrText = OxmlElement("w:instrText")
            instrText.text = "PAGE"
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run_page = p.add_run()
            run_page._r.append(fldChar1)
            run_page._r.append(instrText)
            run_page._r.append(fldChar2)
            run_page.font.size = Pt(8)
