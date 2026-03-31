# -*- coding: utf-8 -*-
"""
DEMO – Dashboard-Panel (links) NEBEN normalem Stärkemeldungs-Export (rechts)
Links  : DRK-Kennzahlen-Panel (wie Screenshot: Datum, Betreuungen, Bulmor-Status)
Rechts : Normales Stärkemeldung-Format (aus StaerkemeldungExport / dienstplan.py)
Layout : A4 Hochformat – 2-spaltige Tabelle
Ausgabe: word neu 26.3/DEMO_Dashboard_links_Staerkemeldung_rechts.docx
"""

import os, sys
from pathlib import Path

_NESK3 = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_NESK3))

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Pfade ─────────────────────────────────────────────────────────────────────
LOGO = _NESK3 / "Daten" / "Email" / "Logo.jpg"
ZIEL = Path(__file__).parent  # word neu 26.3/

# ── Demo-Werte (aus Screenshot) ───────────────────────────────────────────────
DATUM      = "25.03.2026"
BET_JAHR   = 21_000
BET_VORT   = 300
EINZ_SL    = 7
BUL_ANZ    = 3      # 3 von 5 Bulmor aktiv → ORANGE
STATION    = "Erste-Hilfe-Station · Flughafen Köln/Bonn"
TEL        = "+49 2203 40-2323"
MAIL       = "erste-hilfe-station-flughafen@drk-koeln.de"
STATIONSLT     = "Peters"
SCHICHTL_TAG   = ("Peters",  "08:00–17:15")   # Schichtleiter Tagdienst
SCHICHTL_NACHT = ("Acar",    "18:00–06:00")   # Schichtleiter Nachtdienst
PAX_EINZEL     = 42_500   # Tages-PAX dieser Tag

# Demo-Personaldaten (aus Max30_Bul3_Orange.docx)
DISPO_TAG = [
    ("Peters",      "08:00–17:15"),
    ("Lehmann",     "08:00–16:00"),
]
DISPO_NACHT = [
    ("Acar",        "18:00–06:00"),
    ("Mantzas",     "18:00–06:20"),
]
BETREUER_TAG = [
    ("Athanasiou",  "06:00–18:00"), ("Badrieh",    "06:00–18:00"),
    ("Baluch",      "06:15–18:00"), ("Bauschke",   "06:00–17:00"),
    ("Gül",         "06:00–18:30"), ("Schneider",  "06:00–18:00"),
    ("Tamer",       "06:00–18:00"), ("Üzülmez",    "06:00–18:00"),
    ("El Mojahid",  "08:00–16:00"), ("Idic",       "08:00–18:00"),
    ("Doubli",      "09:00–19:00"), ("Heim",       "09:00–19:00"),
    ("Loukili",     "09:00–19:00"), ("Pieper",     "09:00–19:00"),
    ("Thiebes",     "09:00–19:00"), ("Delgado",    "09:30–19:00"),
    ("Cemal",       "10:00–18:00"),
]
BETREUER_NACHT = [
    ("Bakkal",      "18:00–06:00"), ("Bedl",       "18:30–06:00"),
    ("Bouladhane",  "18:00–06:00"), ("Campolo",    "18:00–06:00"),
    ("Dobrani",     "18:00–06:00"), ("Oh",         "18:00–06:00"),
    ("Taute",       "18:00–06:00"), ("Chugh",      "21:00–07:00"),
    ("Hein",        "21:00–07:00"), ("Irani",      "21:00–07:00"),
    ("Isa",         "21:00–07:00"), ("Tunahan",    "21:00–07:00"),
]
PAX = 42_500

# ── Farben ────────────────────────────────────────────────────────────────────
BG_DUNKEL = "1A3460"   # Dunkelblau (Dashboard-Hintergrund)
BG_MITTEL = "0F1F3C"   # Noch dunkler
HE        = "C8DAFF"   # Hellblau / Text auf dunkel
AZ        = "00C8FF"   # Akzent Cyan / Tagdienst
WEISS     = "FFFFFF"

# Bulmor-Statusfarben
def _bul_farben(n):
    if n <= 2: return "FF3333", "KRITISCH",      "3A0000", "CC2222"
    if n == 3: return "E07800", "EINGESCHRÄNKT", "3A2000", "B86000"
    return              "10A050", "VOLLSTÄNDIG",   "003A18", "0A8040"

FC_HEX, LBL_BUL, FC_BG, FC_BOX = _bul_farben(BUL_ANZ)


# ═══════════════════════════════════════════════════════════════════════════════
# Hilfs-Funktionen
# ═══════════════════════════════════════════════════════════════════════════════
def _rgb(hx: str) -> RGBColor:
    hx = hx.lstrip("#")
    return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

def _set_bg(cell, hx: str):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    s  = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hx.lstrip("#").upper())
    pr.append(s)

def _no_border(cell):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    b  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{side}"); e.set(qn("w:val"), "none"); b.append(e)
    pr.append(b)

def _para(cell, text, bold=False, size=9, fg="000000",
          align="left", sa=0, sb=0):
    p = cell.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center" else
                   WD_ALIGN_PARAGRAPH.RIGHT  if align == "right"  else
                   WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = _rgb(fg)
    return p

def _trenn(cell, hx, oben=False):
    p   = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr")
    e   = OxmlElement("w:top" if oben else "w:bottom")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "6")
    e.set(qn("w:space"), "1"); e.set(qn("w:color"), hx.upper())
    bdr.append(e); pPr.append(bdr)

def _hdr_box(cell, text, bg, fg="FFFFFF", w=Cm(8), sa=1, sb=1):
    t  = cell.add_table(rows=1, cols=1); t.style = "Table Grid"
    c  = t.cell(0, 0); _no_border(c); _set_bg(c, bg); c.width = w
    _para(c, text, bold=True, size=8, fg=fg, sa=sa, sb=sb)
    return t

def _zeitgruppen_para(cell, gruppen: dict, size=9.5):
    """Gibt Zeitgruppen im klassischen StaerkemeldungExport-Stil aus."""
    for zeit, namen in sorted(gruppen.items()):
        p   = cell.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        # Tab-Stop bei 4,5 cm
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left"); tab.set(qn("w:pos"), "2550")
        tabs.append(tab); pPr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "2550"); ind.set(qn("w:hanging"), "2550")
        pPr.append(ind)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)  # Zeilenabstand zwischen Zeitgruppen
        rz = p.add_run(f"{zeit}\t"); rz.font.size = Pt(size)
        rn = p.add_run(" / ".join(namen)); rn.font.size = Pt(size)


# ═══════════════════════════════════════════════════════════════════════════════
# Zeitgruppen aus Listen bauen
# ═══════════════════════════════════════════════════════════════════════════════
def _grup(lst):
    g = {}
    for name, zeit in lst:
        g.setdefault(zeit, []).append(name)
    return g


# ═══════════════════════════════════════════════════════════════════════════════
# Haupt-Funktion
# ═══════════════════════════════════════════════════════════════════════════════
def erstelle_demo():
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Cm(21.0); sec.page_height  = Cm(29.7)
        sec.top_margin    = Cm(0.5);  sec.bottom_margin = Cm(0.5)
        sec.left_margin   = Cm(0.5);  sec.right_margin  = Cm(0.5)

    # Kopfzeile (Word-nativer Header – erscheint auf ganzer Seite)
    hdr = doc.sections[0].header
    ht  = hdr.add_table(rows=1, cols=2, width=Inches(7.8)); ht.autofit = False
    # Logo links
    if LOGO.exists():
        lp = ht.rows[0].cells[0].paragraphs[0]
        lp.add_run().add_picture(str(LOGO), width=Inches(1.0))
    # Text rechts
    tp = ht.rows[0].cells[1].paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = tp.add_run("Deutsches Rotes Kreuz Kreisverband Köln e.V.\n")
    r1.font.size = Pt(12); r1.font.bold = True
    r2 = tp.add_run("Unfallhilfsstelle und Betreuungsstelle · Flughafen Köln/Bonn")
    r2.font.size = Pt(11)
    hdr.add_paragraph("_" * 90)

    # Fußzeile
    ftr   = doc.sections[0].footer
    fp    = ftr.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr    = fp.add_run(f"☎ {TEL}   |   ✉ {MAIL}   |   Stationsleitung: {STATIONSLT}")
    fr.font.size = Pt(11); fr.font.color.rgb = _rgb("777777")

    # ── Haupt-Tabelle: 2 Spalten ─────────────────────────────────────────────
    L_W = Cm(5.5)    # Links  – Dashboard-Panel (schmal)
    R_W = Cm(14.0)   # Rechts – Stärkemeldung klassisch (breit)

    main = doc.add_table(rows=1, cols=2); main.style = "Table Grid"
    lc = main.cell(0, 0); rc = main.cell(0, 1)
    lc.width = L_W;  rc.width = R_W
    _no_border(lc);  _no_border(rc)
    # kein farbiger Hintergrund – weißes Panel
    lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Zeile exakt auf Seitenhöhe fixieren → kein Seitenumbruch möglich
    # Exakt aus v8-Dokument gemessen: 12689 twips = 22.38 cm
    PAGE_H_CM = 22.38
    tr = main.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(PAGE_H_CM * 567)))  # 1 cm = 567 twips
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)

    # ══════════════════════════════════════════════════════════════════════════
    # LINKE SPALTE – Dashboard (wie Screenshot)
    # ══════════════════════════════════════════════════════════════════════════

    # Org-Info (kompakter wegen schmaler Spalte)
    _para(lc, "Deutsches Rotes Kreuz", bold=True, size=11, fg="1A3460",
          align="center", sb=2)
    _para(lc, "Kreisverband Köln e.V.", size=10, fg="1A3460", align="center")
    _para(lc, STATION, size=9, fg="666666", align="center")
    _para(lc, TEL, bold=True, size=10, fg="1A3460", align="center", sa=1)

    _trenn(lc, "1565a8")

    # Datum
    _para(lc, f"Datum:  {DATUM}", bold=True, size=12, fg="1A3460",
          align="center", sb=2, sa=2)

    _trenn(lc, "1565a8")

    # Kennzahlen  ✦ Label \n WERT
    kz = [
        ("✦ Betreuungen (2026)",    f"{BET_JAHR:,}".replace(",","."), "444444", "1565a8"),
        ("✦ Betreuungen (Vortag)",  str(BET_VORT),                   "444444", "1A3460"),
        ("✦ Einsätze SL",            str(EINZ_SL),                    "444444", "1565a8"),
        ("✦ PAX gestern",            f"{PAX_EINZEL:,}".replace(",","."), "444444", "0A7040"),
    ]
    for lbl, val, lbl_fg, val_fg in kz:
        p = lc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(0)
        r1 = p.add_run(f"{lbl}\n"); r1.font.size = Pt(9.5); r1.font.color.rgb = _rgb(lbl_fg)
        r2 = p.add_run(val)
        r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = _rgb(val_fg)

    # Schichtleiter Tag + Nacht im Dashboard
    _trenn(lc, "1565a8", oben=True)
    _para(lc, "SCHICHTLEITER", bold=True, size=9.5, fg="1565a8", align="center", sb=1)
    _para(lc, f"Tag:   {SCHICHTL_TAG[0]}",   size=10, fg="1A3460", align="center", sb=0)
    _para(lc, SCHICHTL_TAG[1],               size=9.5, fg="2A4A7F", align="center", sa=0)
    _para(lc, f"Nacht: {SCHICHTL_NACHT[0]}", size=10, fg="1A3460", align="center", sb=0)
    _para(lc, SCHICHTL_NACHT[1],                size=9.5, fg="2A4A7F", align="center", sa=1)

    _trenn(lc, "1565a8", oben=True)

    # BULMOR – FAHRZEUGSTATUS Header
    p_bul_hdr = lc.add_paragraph()
    p_bul_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bul_hdr.paragraph_format.space_before = Pt(2)
    p_bul_hdr.paragraph_format.space_after  = Pt(1)
    rb = p_bul_hdr.add_run("BULMOR – FAHRZEUGSTATUS")
    rb.bold = True; rb.font.size = Pt(10); rb.font.color.rgb = _rgb("1A3460")

    # B1–B5 als Tabelle
    btbl = lc.add_table(rows=2, cols=5); btbl.style = "Table Grid"
    for i in range(5):
        aktiv  = (i + 1) <= BUL_ANZ
        bg_top = FC_HEX   if aktiv else "CC3333"
        bg_bot = "E8F0FB" if aktiv else "FFE0E0"
        ct = btbl.cell(0, i); cb = btbl.cell(1, i)
        _no_border(ct); _no_border(cb)
        _set_bg(ct, bg_top); _set_bg(cb, bg_bot)
        # Obere Zeile: Kreis + B-Nummer
        pt = ct.paragraphs[0]; pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.paragraph_format.space_before = Pt(1); pt.paragraph_format.space_after = Pt(0)
        rt = pt.add_run(f"●\nB{i+1}")
        rt.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = _rgb(WEISS)
        # Untere Zeile: Status-Text
        pb = cb.paragraphs[0]; pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pb.paragraph_format.space_before = Pt(0); pb.paragraph_format.space_after = Pt(1)
        status = "Dienst" if aktiv else "Aus"
        rb2 = pb.add_run(status)
        rb2.font.size = Pt(10)
        rb2.font.color.rgb = _rgb(FC_HEX if aktiv else "CC3333")

    # Gesamt-Box – kein farbiger Hintergrund
    gbox = lc.add_table(rows=1, cols=1); gbox.style = "Table Grid"
    gc   = gbox.cell(0, 0); _no_border(gc); gc.width = L_W
    pg   = gc.paragraphs[0]; pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pg.paragraph_format.space_before = Pt(2); pg.paragraph_format.space_after = Pt(2)
    rg   = pg.add_run(f"Gesamt: {BUL_ANZ}/5  ◇  {LBL_BUL}")
    rg.bold = True; rg.font.size = Pt(10.5); rg.font.color.rgb = _rgb("1A3460")

    # ══════════════════════════════════════════════════════════════════════════
    # RECHTE SPALTE – klassische Stärkemeldung (StaerkemeldungExport-Stil)
    # ══════════════════════════════════════════════════════════════════════════

    # Sub-Header (kein farbiger Block)
    p_hdr = rc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(1)
    p_hdr.paragraph_format.space_after  = Pt(1)
    rh = p_hdr.add_run(f"STÄRKEMELDUNG  ·  {DATUM}")
    rh.bold = True; rh.font.size = Pt(13); rh.font.color.rgb = _rgb("1A3460")
    # Leerabsatz nach Header (sa=1, sb=1) – wie in v8
    p_sep0 = rc.add_paragraph()
    p_sep0.paragraph_format.space_before = Pt(1); p_sep0.paragraph_format.space_after = Pt(1)

    # Zeitraum-Zeile
    p_zr = rc.add_paragraph()
    p_zr.paragraph_format.space_before = Pt(0)
    p_zr.paragraph_format.space_after  = Pt(1)
    rz   = p_zr.add_run(f"Zeitraum:\t{DATUM} bis {DATUM}")
    rz.font.size = Pt(12); rz.font.bold = True
    # Leerabsatz nach Zeitraum (sa=1, sb=0) – wie in v8
    p_sep1 = rc.add_paragraph()
    p_sep1.paragraph_format.space_before = Pt(0); p_sep1.paragraph_format.space_after = Pt(1)

    # ── Schichtleiter ─────────────────────────────────────────────────────────
    ph_sl = rc.add_paragraph()
    ph_sl.paragraph_format.space_before = Pt(1)
    ph_sl.paragraph_format.space_after  = Pt(0)
    rh_sl = ph_sl.add_run("Schichtleiter"); rh_sl.font.bold = True; rh_sl.font.size = Pt(12)
    # Leerabsatz nach Schichtleiter-Heading (sa=0, sb=1) – wie in v8
    p_sep2 = rc.add_paragraph()
    p_sep2.paragraph_format.space_before = Pt(1); p_sep2.paragraph_format.space_after = Pt(0)

    for prefix, (name, zeit) in [("Tag:  ", SCHICHTL_TAG), ("Nacht:", SCHICHTL_NACHT)]:
        p_sl = rc.add_paragraph()
        pPr  = p_sl._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left"); tab.set(qn("w:pos"), "2550")
        tabs.append(tab); pPr.append(tabs)
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "2550"); ind.set(qn("w:hanging"), "2550")
        pPr.append(ind)
        p_sl.paragraph_format.space_before = Pt(0); p_sl.paragraph_format.space_after = Pt(0)
        r_pf = p_sl.add_run(f"{prefix}\t"); r_pf.font.size = Pt(12)
        r_nm = p_sl.add_run(f"{name}  ({zeit})"); r_nm.font.size = Pt(12); r_nm.bold = True

    rc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Disposition ───────────────────────────────────────────────────────────
    ph = rc.add_paragraph()
    ph.paragraph_format.space_before = Pt(1)
    ph.paragraph_format.space_after  = Pt(0)
    rh = ph.add_run("Disposition"); rh.font.bold = True; rh.font.size = Pt(12)
    # Leerabsatz nach Disposition-Heading (sa=0, sb=1) – wie in v8
    p_sep3 = rc.add_paragraph()
    p_sep3.paragraph_format.space_before = Pt(1); p_sep3.paragraph_format.space_after = Pt(0)

    all_dispo = DISPO_TAG + DISPO_NACHT
    _zeitgruppen_para(rc, _grup(all_dispo), size=11.5)
    rc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Behindertenbetreuer ───────────────────────────────────────────────────
    ph2 = rc.add_paragraph()
    ph2.paragraph_format.space_before = Pt(1)
    ph2.paragraph_format.space_after  = Pt(0)
    rh2 = ph2.add_run("Behindertenbetreuer"); rh2.font.bold = True; rh2.font.size = Pt(12)
    # Leerabsatz nach Behindertenbetreuer-Heading (sa=0, sb=1) – wie in v8
    p_sep4 = rc.add_paragraph()
    p_sep4.paragraph_format.space_before = Pt(1); p_sep4.paragraph_format.space_after = Pt(0)

    all_bet = BETREUER_TAG + BETREUER_NACHT
    _zeitgruppen_para(rc, _grup(all_bet), size=11.5)
    rc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── PAX-Zahl (jetzt hier, kein separater Block mehr am Ende) ──────
    ph_pax = rc.add_paragraph()
    ph_pax.paragraph_format.space_before = Pt(1)
    ph_pax.paragraph_format.space_after  = Pt(0)
    rh_pax = ph_pax.add_run("PAX gestern"); rh_pax.font.bold = True; rh_pax.font.size = Pt(12)
    # Leerabsatz nach PAX-Heading (sa=0, sb=1) – wie in v8
    p_sep5 = rc.add_paragraph()
    p_sep5.paragraph_format.space_before = Pt(1); p_sep5.paragraph_format.space_after = Pt(0)
    p_paxz = rc.add_paragraph()
    pPr_p  = p_paxz._p.get_or_add_pPr()
    tabs_p = OxmlElement("w:tabs"); tab_p = OxmlElement("w:tab")
    tab_p.set(qn("w:val"), "left"); tab_p.set(qn("w:pos"), "2550")
    tabs_p.append(tab_p); pPr_p.append(tabs_p)
    ind_p  = OxmlElement("w:ind"); ind_p.set(qn("w:left"), "2550"); ind_p.set(qn("w:hanging"), "2550")
    pPr_p.append(ind_p)
    p_paxz.paragraph_format.space_before = Pt(0); p_paxz.paragraph_format.space_after = Pt(0)
    rp1 = p_paxz.add_run(f"{DATUM}\t"); rp1.font.size = Pt(12)
    rp2 = p_paxz.add_run(f"{PAX_EINZEL:,}".replace(",", ".") + " Passagiere")
    rp2.font.size = Pt(12); rp2.bold = True

    # ── Speichern ─────────────────────────────────────────────────────────────
    out = ZIEL / f"DEMO_Dashboard_v9_{DATUM.replace('.','')}.docx"
    doc.save(str(out))
    print(f"[OK] Gespeichert: {out}")
    return str(out)


if __name__ == "__main__":
    erstelle_demo()
