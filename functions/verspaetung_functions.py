"""
Verspätungs-Funktionen
Erstellen des ausgefüllten Word-Dokuments aus der FO-Vorlage.
"""
import os
import shutil
from pathlib import Path
from datetime import datetime

from config import BASE_DIR as _BASE_DIR
BASE_DIR = Path(_BASE_DIR)

VORLAGE_PFAD = (
    BASE_DIR / "Daten" / "Spät" / "FO_CGN_27_Unpünktlicher Dienstantritt.docx"
)
PROTOKOLL_DIR = BASE_DIR / "Daten" / "Spät" / "Protokoll"

# Automatisch voreingestellte Dienstbeginn-Zeiten je Dienstart
_DIENST_BEGINN = {
    "T":   "06:00",
    "T10": "09:00",
    "N":   "18:00",
    "N10": "21:00",
}


def dienstbeginn_fuer(dienst: str) -> str:
    """Gibt die übliche Dienstbeginn-Zeit für eine Dienstart zurück."""
    return _DIENST_BEGINN.get(dienst, "06:00")


def berechne_verspaetung_min(dienstbeginn: str, dienstantritt: str) -> int:
    """
    Berechnet die Verspätung in Minuten.
    Positiver Wert = zu spät; Negativer = zu früh.
    """
    try:
        h1, m1 = map(int, dienstbeginn.split(":"))
        h2, m2 = map(int, dienstantritt.split(":"))
        return (h2 * 60 + m2) - (h1 * 60 + m1)
    except Exception:
        return 0


def _fill_cell(cell, text: str):
    """
    Setzt den Text der ersten Paragraph eines Word-Tabellenfeldes.
    Entfernt ALLE Inline-Inhalte (auch Nicht-Run-Elemente wie eingebettete
    Kästchen-Zeichen) und schreibt einen neuen Run (Arial 12pt) mit dem Text.
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt
    para = cell.paragraphs[0]
    p_elem = para._p
    # Alle Kinder außer den Absatz-Eigenschaften (w:pPr) entfernen
    for child in list(p_elem):
        if child.tag != qn("w:pPr"):
            p_elem.remove(child)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)


def erstelle_verspaetungs_dokument(daten: dict) -> str:
    """
    Füllt die Word-Vorlage mit den übergebenen Daten aus,
    speichert ein neues Dokument in PROTOKOLL_DIR
    und gibt den vollständigen Dateipfad zurück.

    Erwartete Schlüssel in ``daten``:
        mitarbeiter, datum, dienst, dienstbeginn, dienstantritt,
        begruendung, aufgenommen_von
    """
    from docx import Document

    PROTOKOLL_DIR.mkdir(parents=True, exist_ok=True)

    if not VORLAGE_PFAD.exists():
        raise FileNotFoundError(
            f"Vorlage nicht gefunden:\n{VORLAGE_PFAD}\n\n"
            "Bitte stelle sicher, dass die Datei "
            "'FO_CGN_27_Unpünktlicher Dienstantritt.docx' "
            "im Ordner Daten/Spät liegt."
        )

    # Dateiname
    ma_name  = daten.get("mitarbeiter", "Unbekannt").replace(" ", "_")
    datum    = daten.get("datum", "").replace(".", "")
    ts       = datetime.now().strftime("%H%M%S")
    dateiname = f"Verspaetung_{ma_name}_{datum}_{ts}.docx"
    ziel_pfad = PROTOKOLL_DIR / dateiname

    shutil.copy2(str(VORLAGE_PFAD), str(ziel_pfad))

    doc = Document(str(ziel_pfad))
    t0 = doc.tables[0]
    t1 = doc.tables[1]

    # ── Tabelle 0 ──────────────────────────────────────────────────────────────
    # Datum
    _fill_cell(t0.rows[0].cells[1], daten.get("datum", ""))

    # Mitarbeiter
    _fill_cell(t0.rows[1].cells[1], daten.get("mitarbeiter", ""))

    # Dienst-Checkboxen
    dienst = daten.get("dienst", "T")
    _fill_cell(t0.rows[2].cells[1], "☑ T"   if dienst == "T"   else " T")
    _fill_cell(t0.rows[2].cells[3], "☑ N"   if dienst == "N"   else " N")
    _fill_cell(t0.rows[3].cells[1], "☑ T10" if dienst == "T10" else " T10")
    _fill_cell(t0.rows[3].cells[3], "☑ N10" if dienst == "N10" else " N10")

    # Dienstbeginn (Stunde / Minute)
    beginn_parts = daten.get("dienstbeginn", "06:00").split(":")
    _fill_cell(t0.rows[4].cells[1], beginn_parts[0] if len(beginn_parts) > 0 else "06")
    _fill_cell(t0.rows[4].cells[2], beginn_parts[1] if len(beginn_parts) > 1 else "00")

    # Dienstantritt (Stunde / Minute)
    antritt_parts = daten.get("dienstantritt", "").split(":")
    _fill_cell(t0.rows[5].cells[1], antritt_parts[0] if len(antritt_parts) > 0 else "")
    _fill_cell(t0.rows[5].cells[2], antritt_parts[1] if len(antritt_parts) > 1 else "")

    # Begründung
    _fill_cell(t0.rows[6].cells[1], daten.get("begruendung", ""))

    # ── Tabelle 1 ──────────────────────────────────────────────────────────────
    # Den Namen in den leeren Paragraph direkt VOR der Tabelle (im Dokument-Body)
    # schreiben – das ist die Position, die im Formular über dem Strich erscheint.
    from docx.oxml.ns import qn
    from docx.shared import Pt
    body = doc.element.body
    t1_elem = doc.tables[1]._tbl
    body_children = list(body)
    t1_idx = body_children.index(t1_elem)
    # Paragraph direkt vor der Tabelle befüllen
    para_vor_tabelle = body_children[t1_idx - 1] if t1_idx > 0 else None
    if para_vor_tabelle is not None and para_vor_tabelle.tag.endswith("}p"):
        from docx.oxml.ns import qn as _qn
        # Vorhandene Runs löschen
        for child in list(para_vor_tabelle):
            if child.tag != _qn("w:pPr"):
                para_vor_tabelle.remove(child)
        # Neuen Run mit Arial 12pt einfügen
        from docx.oxml import OxmlElement
        new_r = OxmlElement("w:r")
        new_rpr = OxmlElement("w:rPr")
        new_fonts = OxmlElement("w:rFonts")
        new_fonts.set(qn("w:ascii"), "Arial")
        new_fonts.set(qn("w:hAnsi"), "Arial")
        new_rpr.append(new_fonts)
        new_sz = OxmlElement("w:sz")
        new_sz.set(qn("w:val"), "24")   # 24 half-points = 12pt
        new_rpr.append(new_sz)
        new_r.append(new_rpr)
        new_t = OxmlElement("w:t")
        new_t.text = daten.get("aufgenommen_von", "")
        new_r.append(new_t)
        para_vor_tabelle.append(new_r)

    doc.save(str(ziel_pfad))
    try:
        from functions.dokument_archiv import kopiere_ins_archiv
        kopiere_ins_archiv(str(ziel_pfad), "verspaetung")
    except Exception:
        pass
    return str(ziel_pfad)


def oeffne_dokument(pfad: str):
    """Öffnet ein Dokument mit dem Standard-Programm (Windows)."""
    if os.path.isfile(pfad):
        os.startfile(pfad)
    else:
        raise FileNotFoundError(f"Datei nicht gefunden: {pfad}")
