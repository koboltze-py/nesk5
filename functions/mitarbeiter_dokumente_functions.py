"""
Mitarbeiter-Dokumente – Hilfsfunktionen
Erstellen von Word-Dokumenten mit der DRK-Kopf-/Fußzeile aus der Vorlage
"""
import os
import shutil
from datetime import datetime
from pathlib import Path
from functions.stellungnahmen_db import eintrag_speichern as _db_eintrag_speichern
try:
    from functions.dienstanweisungen_db import eintrag_speichern as _da_db_speichern
except Exception:
    _da_db_speichern = None  # DB nicht verfügbar – kein Fehler
from config import BASE_DIR

# Pfad zur Vorlage mit Kopf-/Fußzeile
VORLAGE_PFAD = os.path.join(
    BASE_DIR, "Daten", "Mitarbeiter Vorlagen",
    "Kopf und Fußzeile",
    "Stärkemeldung 31.01.2026 bis 01.02.2026.docx",
)

# Ordner für erstellte Dokumente (kategorisiert)
DOKUMENTE_BASIS = os.path.join(BASE_DIR, "Daten", "Mitarbeiterdokumente")

# Externer Stellungnahmen-Ordner (OneDrive-Ablage)
STELLUNGNAHMEN_EXTERN_PFAD = os.path.join(
    os.path.dirname(os.path.dirname(BASE_DIR)),
    "97_Stellungnahmen",
)

KATEGORIEN = [
    "Stellungnahmen",
    "Bescheinigungen und Anträge",
    "Dienstanweisungen",
    "PSA",
    "Verspätung",
    "Schulungen",
]


def sicherungsordner() -> str:
    """Gibt den Basispfad für Mitarbeiterdokumente zurück und legt ihn an."""
    os.makedirs(DOKUMENTE_BASIS, exist_ok=True)
    for kat in KATEGORIEN:
        os.makedirs(os.path.join(DOKUMENTE_BASIS, kat), exist_ok=True)
    return DOKUMENTE_BASIS


def lade_dokumente_nach_kategorie() -> dict[str, list[dict]]:
    """
    Gibt alle Dokumente je Kategorie zurück.
    Rückgabe: { "Stellungnahmen": [{"name": ..., "pfad": ..., "geaendert": ...}, ...], ... }
    """
    sicherungsordner()
    ergebnis: dict[str, list[dict]] = {}
    for kat in KATEGORIEN:
        ordner = os.path.join(DOKUMENTE_BASIS, kat)
        dateien = []
        if os.path.isdir(ordner):
            for fname in sorted(os.listdir(ordner)):
                if fname.lower().endswith((".docx", ".doc", ".pdf", ".txt")):
                    pfad = os.path.join(ordner, fname)
                    mtime = datetime.fromtimestamp(os.path.getmtime(pfad))
                    dateien.append({
                        "name": fname,
                        "pfad": pfad,
                        "geaendert": mtime.strftime("%d.%m.%Y %H:%M"),
                    })
        ergebnis[kat] = dateien
    return ergebnis


def erstelle_dokument_aus_vorlage(
    kategorie: str,
    titel: str,
    mitarbeiter: str,
    datum: str,
    inhalt: str,
    dateiname: str = "",
) -> str:
    """
    Erstellt ein neues Word-Dokument mit der DRK-Kopf-/Fußzeile aus der Vorlage.
    Gibt den Pfad der erstellten Datei zurück.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen.")

    sicherungsordner()

    # Dateiname generieren
    if not dateiname:
        safe_titel = "".join(c for c in titel if c.isalnum() or c in " _-").strip()[:40]
        safe_ma = "".join(c for c in mitarbeiter if c.isalnum() or c in " _-").strip()[:20]
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        dateiname = f"{safe_titel}_{safe_ma}_{stamp}.docx"

    ziel_pfad = os.path.join(DOKUMENTE_BASIS, kategorie, dateiname)

    # Vorlage öffnen (enthält Kopf-/Fußzeile)
    if os.path.isfile(VORLAGE_PFAD):
        doc = Document(VORLAGE_PFAD)
        # Bestehende Absätze im Body entfernen
        for para in doc.paragraphs[:]:
            p = para._element
            p.getparent().remove(p)
        for table in doc.tables[:]:
            t = table._element
            t.getparent().remove(t)
    else:
        # Fallback: leeres Dokument
        doc = Document()

    # Inhalt einfügen
    # Titel
    titel_para = doc.add_paragraph()
    run = titel_para.add_run(titel)
    run.bold = True
    run.font.size = Pt(16)
    titel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Leerzeile

    # Meta-Informationen
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Mitarbeiter: ").bold = True
    meta_para.add_run(mitarbeiter)
    meta_para.add_run(f"   Datum: ").bold = True
    meta_para.add_run(datum)

    doc.add_paragraph()  # Leerzeile

    # Hauptinhalt
    for zeile in inhalt.split("\n"):
        doc.add_paragraph(zeile)

    # Datum + Unterschrift
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f"Köln/Bonn Flughafen, {datum}")
    doc.add_paragraph()
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Unterschrift")

    doc.save(ziel_pfad)
    try:
        from functions.dokument_archiv import kopiere_ins_archiv
        _bereich = {
            "Stellungnahmen": "stellungnahmen",
            "Dienstanweisungen": "dienstanweisungen",
            "Verspätung": "verspaetung",
        }.get(kategorie, "mitarbeiterdokumente")
        kopiere_ins_archiv(ziel_pfad, _bereich)
    except Exception:
        pass
    return ziel_pfad


def oeffne_datei(pfad: str) -> None:
    """Öffnet eine Datei mit der zugehörigen Windows-Standardanwendung."""
    import subprocess
    subprocess.Popen(["start", "", pfad], shell=True)


def loesche_dokument(pfad: str) -> bool:
    """Löscht eine Datei, gibt True zurück wenn erfolgreich."""
    try:
        if os.path.isfile(pfad):
            os.remove(pfad)
            return True
    except Exception:
        pass
    return False


def umbenennen_dokument(alter_pfad: str, neuer_name: str) -> str:
    """Benennt ein Dokument um. Gibt den neuen Pfad zurück."""
    ordner = os.path.dirname(alter_pfad)
    neuer_pfad = os.path.join(ordner, neuer_name)
    os.rename(alter_pfad, neuer_pfad)
    return neuer_pfad


# ══════════════════════════════════════════════════════════════════════════════
#  Stellungnahme – strukturiertes Word-Dokument
# ══════════════════════════════════════════════════════════════════════════════

def erstelle_stellungnahme(daten: dict) -> tuple[str, str]:
    """
    Erstellt eine strukturierte Stellungnahme als Word-Dokument.
    Speichert in BEIDEN Ordnern:
      - intern: Daten/Mitarbeiterdokumente/Stellungnahmen/
      - extern: 97_Stellungnahmen/

    daten-Keys:
        mitarbeiter, datum (dd.MM.yyyy), verfasst_am (dd.MM.yyyy)
        art: "flug" | "beschwerde" | "nicht_mitgeflogen"
        --- bei art == "flug" ---
        flugnummer, verspaetung (bool)
        onblock, offblock  (HH:MM, nur wenn verspaetung)
        richtung: "inbound" | "outbound" | "beides"
        ankunft_lfz, auftragsende  (HH:MM, bei inbound/beides)
        paxannahme_zeit            (HH:MM, bei outbound/beides)
        paxannahme_ort             (str,   bei outbound/beides)
        sachverhalt                (Freitext)
        --- bei art == "beschwerde" ---
        beschwerde_text (Freitext)
        sachverhalt
        --- bei art == "nicht_mitgeflogen" ---
        flugnummer, sachverhalt

    Gibt zurück: (intern_pfad, extern_pfad)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise RuntimeError("python-docx ist nicht installiert.")

    sicherungsordner()
    os.makedirs(STELLUNGNAHMEN_EXTERN_PFAD, exist_ok=True)

    # ── Dateiname ──────────────────────────────────────────────────────────────
    art = daten.get("art", "flug")
    ma_safe = "".join(c for c in daten.get("mitarbeiter", "Unbekannt")
                      if c.isalnum() or c in " _-").strip().replace(" ", "_")[:20]
    fn_prefix = daten.get("flugnummer", "").replace(" ", "") or "oF"
    dateiname = f"SN_{ma_safe}_{fn_prefix}.docx"

    intern_pfad = os.path.join(DOKUMENTE_BASIS, "Stellungnahmen", dateiname)
    extern_pfad = os.path.join(STELLUNGNAHMEN_EXTERN_PFAD, dateiname)

    # ── Vorlage öffnen ─────────────────────────────────────────────────────────
    if os.path.isfile(VORLAGE_PFAD):
        doc = Document(VORLAGE_PFAD)
        for para in doc.paragraphs[:]:
            para._element.getparent().remove(para._element)
        for table in doc.tables[:]:
            table._element.getparent().remove(table._element)
    else:
        doc = Document()

    def _h(text: str, size: int = 13, bold: bool = True, center: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _row(label: str, value: str):
        p = doc.add_paragraph()
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(value or "–")
        r_val.font.size = Pt(11)
        return p

    def _trennlinie():
        p = doc.add_paragraph("─" * 60)
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    def _abschnitt(titel: str):
        p = doc.add_paragraph()
        run = p.add_run(f"▌ {titel}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xA8)

    def _freitext(text: str):
        for zeile in (text or "").split("\n"):
            p = doc.add_paragraph(zeile)
            p.runs[0].font.size = Pt(11) if p.runs else None

    # ── Titel ──────────────────────────────────────────────────────────────────
    _h("S T E L L U N G N A H M E", size=18, center=True)
    doc.add_paragraph()
    _row("Mitarbeiter", daten.get("mitarbeiter", ""))
    _row("Datum des Vorfalls", daten.get("datum", ""))
    _row("Verfasst am", daten.get("verfasst_am", datetime.now().strftime("%d.%m.%Y")))
    _trennlinie()

    # ── Art der Stellungnahme ──────────────────────────────────────────────────
    art_label = {
        "flug": "Flug-bezogener Vorfall",
        "beschwerde": "Passagierbeschwerde",
        "nicht_mitgeflogen": "Passagier nicht mitgeflogen (kein PRM-Dienst)",
    }.get(art, art)
    _row("Art der Stellungnahme", art_label)
    doc.add_paragraph()

    # ── FLUG ───────────────────────────────────────────────────────────────────
    if art == "flug":
        _abschnitt("Flugdaten")
        _row("Flugnummer", daten.get("flugnummer", ""))

        verspaetung = daten.get("verspaetung", False)
        _row("Verspätung", "Ja" if verspaetung else "Nein")
        if verspaetung:
            _row("  Onblock-Zeit (tatsächlich)", daten.get("onblock", ""))
            _row("  Offblock-Zeit (geplant)", daten.get("offblock", ""))

        richtung = daten.get("richtung", "")
        _row("Flugrichtung", {
            "inbound": "Inbound (Ankunft)",
            "outbound": "Outbound (Abflug)",
            "beides": "Inbound + Outbound",
        }.get(richtung, richtung))
        doc.add_paragraph()

        if richtung in ("inbound", "beides"):
            _abschnitt("Inbound-Einsatz")
            _row("Ankunft Luftfahrzeug", daten.get("ankunft_lfz", ""))
            _row("Auftragsende", daten.get("auftragsende", ""))
            doc.add_paragraph()

        if richtung in ("outbound", "beides"):
            _abschnitt("Outbound-Einsatz")
            _row("Paxannahme-Zeit", daten.get("paxannahme_zeit", ""))
            _row("Ort der Paxannahme", daten.get("paxannahme_ort", ""))
            doc.add_paragraph()

        _abschnitt("Sachverhalt")
        _freitext(daten.get("sachverhalt", ""))

    # ── BESCHWERDE ─────────────────────────────────────────────────────────────
    elif art == "beschwerde":
        _abschnitt("Flugzeiten")
        _row("Onblock-Zeit (tatsächlich)", daten.get("onblock", ""))
        _row("Offblock-Zeit (geplant)", daten.get("offblock", ""))
        doc.add_paragraph()
        _abschnitt("Sachverhalt / Beschwerde")
        _freitext(daten.get("sachverhalt", ""))
        if daten.get("beschwerde_text"):
            doc.add_paragraph()
            _abschnitt("Beschreibung der Beschwerde")
            _freitext(daten.get("beschwerde_text", ""))

    # ── NICHT MITGEFLOGEN ─────────────────────────────────────────────────────
    elif art == "nicht_mitgeflogen":
        _abschnitt("Flugdaten")
        _row("Flugnummer", daten.get("flugnummer", ""))
        doc.add_paragraph()
        _abschnitt("Sachverhalt")
        _freitext(daten.get("sachverhalt", ""))

    # ── Unterschrift ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    _trennlinie()
    p = doc.add_paragraph(
        f"Köln/Bonn Flughafen, {daten.get('verfasst_am', datetime.now().strftime('%d.%m.%Y'))}"
    )
    p.runs[0].font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph("_______________________________")
    p_unt = doc.add_paragraph(daten.get("mitarbeiter", "Unterschrift"))
    p_unt.runs[0].bold = True
    p_unt.runs[0].font.size = Pt(11)

    # ── Speichern ──────────────────────────────────────────────────────────────
    doc.save(intern_pfad)
    try:
        from functions.dokument_archiv import kopiere_ins_archiv
        kopiere_ins_archiv(intern_pfad, "stellungnahmen")
    except Exception:
        pass
    shutil.copy2(intern_pfad, extern_pfad)
    # ── Datenbank-Eintrag ───────────────────────────────────────────────────
    try:
        _db_eintrag_speichern(daten, intern_pfad, extern_pfad)
    except Exception:
        pass  # DB-Fehler soll kein Dokument-Erstellen blockieren
    return intern_pfad, extern_pfad


# ══════════════════════════════════════════════════════════════════════════════
#  Dienstanweisung – Freitext-Erstellung mit Formatwahl
# ══════════════════════════════════════════════════════════════════════════════

def dienstanweisung_text_passt(
    text: str,
    ausrichtung: str = "hoch",
    schriftgroesse: int = 11,
) -> tuple[bool, str, int, int]:
    """
    Prüft anhand des eingegebenen Textes (Freitext), ob er auf eine A4-Seite passt.
    Gibt zurück: (passt: bool, hinweis: str, gesamt_zeilen: int, zeilen_pro_seite: int)
    """
    import math
    # A4 in Punkten: 595 × 842 pt
    A4_W, A4_H = 595.0, 842.0
    MARGIN_TB = 72.0   # 2.54 cm oben/unten
    MARGIN_LR = 90.0   # 3.17 cm links/rechts

    if ausrichtung == "quer":
        usable_w = A4_H - 2 * MARGIN_LR
        usable_h = A4_W - 2 * MARGIN_TB
    else:
        usable_w = A4_W - 2 * MARGIN_LR
        usable_h = A4_H - 2 * MARGIN_TB

    char_w = schriftgroesse * 0.5
    chars_per_line = max(1, int(usable_w / char_w))
    line_h = schriftgroesse * 1.5
    lines_per_page = max(1, int(usable_h / line_h))

    # Titelzeile + Leerzeile
    EXTRA_ZEILEN = 2
    gesamt_zeilen = EXTRA_ZEILEN
    for zeile in (text or "").split("\n"):
        if not zeile:
            gesamt_zeilen += 1
        else:
            gesamt_zeilen += math.ceil(len(zeile) / chars_per_line)

    passt = gesamt_zeilen <= lines_per_page
    if passt:
        hinweis = (
            f"✅  Text passt auf eine A4-Seite  "
            f"({gesamt_zeilen} / {lines_per_page} Zeilen)"
        )
    else:
        seiten = math.ceil(gesamt_zeilen / lines_per_page)
        hinweis = (
            f"⚠️  Text benötigt ca. {seiten} Seiten  "
            f"({gesamt_zeilen} Zeilen, max. {lines_per_page} pro Seite)"
        )
    return passt, hinweis, gesamt_zeilen, lines_per_page


def erstelle_dienstanweisung_freitext(
    titel: str,
    inhalt: str,
    ausrichtung: str = "hoch",
    schriftgroesse: int = 11,
) -> str:
    """
    Erstellt eine neue Dienstanweisung als Word-Dokument aus freiem Text.
    Verwendet die DRK-Vorlage (Kopf-/Fußzeile), setzt Ausrichtung und Schriftgröße.
    Speichert unter Daten/Mitarbeiterdokumente/Dienstanweisungen/ und gibt den Pfad zurück.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx ist nicht installiert.")

    sicherungsordner()

    safe_titel = "".join(c for c in (titel or "Dienstanweisung")
                         if c.isalnum() or c in " _-").strip()[:40]
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    dateiname = f"DA_{safe_titel}_{stamp}.docx"
    ziel_pfad = os.path.join(DOKUMENTE_BASIS, "Dienstanweisungen", dateiname)

    # Leeres Dokument – keine Kopf-/Fußzeile aus Vorlage
    doc = Document()

    # Seitenausrichtung
    section = doc.sections[0]
    if ausrichtung == "quer":
        section.orientation = WD_ORIENT.LANDSCAPE
        new_w = max(section.page_width, section.page_height)
        new_h = min(section.page_width, section.page_height)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        new_w = min(section.page_width, section.page_height)
        new_h = max(section.page_width, section.page_height)
    section.page_width  = new_w
    section.page_height = new_h

    # Titel
    p_titel = doc.add_paragraph()
    run_titel = p_titel.add_run(titel or "Dienstanweisung")
    run_titel.bold = True
    run_titel.font.size = Pt(schriftgroesse + 4)
    p_titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Inhalt – jede Zeile als eigener Absatz
    for zeile in (inhalt or "").split("\n"):
        p = doc.add_paragraph(zeile)
        for run in p.runs:
            run.font.size = Pt(schriftgroesse)

    doc.save(ziel_pfad)
    try:
        from functions.dokument_archiv import kopiere_ins_archiv
        kopiere_ins_archiv(ziel_pfad, "dienstanweisungen")
    except Exception:
        pass
    return ziel_pfad
