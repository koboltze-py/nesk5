"""
Stärkemeldung Export nach Word
Adaptiert aus Nesk2/Function/staerkemeldung_export.py für Nesk3
DRK-Vorlage mit Kopfzeile (Logo), Fußzeile, Datumsbereich, Zeitgruppen
"""
import os
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Logo-Pfad relativ zum Projekt-Stammverzeichnis
_BASE_DIR  = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO_PFAD  = os.path.join(_BASE_DIR, "Daten", "Email", "Logo.jpg")


class StaerkemeldungExport:
    """
    Exportiert geparste Dienstplan-Daten als Stärkemeldung (Word-Dokument).
    Struktur: DRK-Kopfzeile → Zeitraum → Disposition → Behindertenbetreuer → PAX → Fußzeile
    """

    def __init__(self, dienstplan_data: dict, ausgabe_pfad: str,
                 von_datum: datetime, bis_datum: datetime, pax_zahl: int = 0,
                 ausgeschlossene_vollnamen=None):
        """
        Args:
            dienstplan_data:          Rückgabe von DienstplanParser.parse()
            ausgabe_pfad:             Pfad zur Ziel-.docx-Datei
            von_datum:                Startdatum
            bis_datum:                Enddatum
            pax_zahl:                 Passagierzahl
            ausgeschlossene_vollnamen: set/list von Vollnamen (lowercase) die
                                      nicht im Dokument erscheinen sollen
        """
        self.data        = dienstplan_data
        self.ausgabe_pfad = ausgabe_pfad
        self.von_datum   = von_datum
        self.bis_datum   = bis_datum
        self.pax_zahl    = pax_zahl
        self.ausgeschlossene = set(n.lower().strip() for n in (ausgeschlossene_vollnamen or []))

    # ------------------------------------------------------------------
    # Öffentliche Methode
    # ------------------------------------------------------------------

    def export(self) -> tuple[str, list[str]]:
        """
        Erstellt Word-Dokument.

        Returns:
            (dateipfad: str, warnungen: list[str])
            warnungen ist eine leere Liste wenn alles in Ordnung ist.
        """
        warnungen = []

        # Leere Kranke aus aktiven Listen herausfiltern
        kranke_set     = set(id(m) for m in self.data.get('kranke', []))
        dispo_aktiv    = [m for m in self.data.get('dispo', [])
                         if id(m) not in kranke_set
                         and m.get('vollname', '').lower() not in self.ausgeschlossene]
        betreuer_aktiv = [m for m in self.data.get('betreuer', [])
                         if id(m) not in kranke_set
                         and m.get('vollname', '').lower() not in self.ausgeschlossene]

        # Nach Startzeit sortieren (None-Werte ans Ende)
        dispo_aktiv    = sorted(dispo_aktiv,    key=lambda x: x.get('start_zeit') or 'ZZZZ')
        betreuer_aktiv = sorted(betreuer_aktiv, key=lambda x: x.get('start_zeit') or 'ZZZZ')

        # --- Dokument aufbauen ---
        doc = Document()

        self._add_header(doc)
        self._add_footer(doc)

        # Datumsbereich
        von_str  = self.von_datum.strftime("%d.%m.%Y")
        bis_str  = self.bis_datum.strftime("%d.%m.%Y")
        datum_p  = doc.add_paragraph()
        run      = datum_p.add_run(f"Zeitraum:\t{von_str} bis {bis_str}")
        run.font.size = Pt(12)
        run.font.bold = True

        doc.add_paragraph()  # Leerzeile

        # Disposition
        if dispo_aktiv:
            h = doc.add_paragraph("Disposition")
            h.runs[0].font.bold = True
            self._add_dienst_gruppe(doc, dispo_aktiv, ist_dispo=True)
            doc.add_paragraph()

        # Behindertenbetreuer
        if betreuer_aktiv:
            h = doc.add_paragraph("Behindertenbetreuer")
            h.runs[0].font.bold = True
            self._add_dienst_gruppe(doc, betreuer_aktiv, ist_dispo=False)
            doc.add_paragraph()

        # PAX-Zahl
        pax_p = doc.add_paragraph(f"- {self.pax_zahl} -")
        pax_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pax_p.runs[0].font.size = Pt(12)
        pax_p.runs[0].font.bold = True

        # Zuerst in temporäre Datei speichern, dann umbenennen.
        # Falls die Zieldatei in Word geöffnet ist (WinError 32),
        # wird ein Ausweich-Dateiname mit Zeitstempel verwendet.
        tmp_pfad = self.ausgabe_pfad + '.nesk3tmp'
        pfad_final = self.ausgabe_pfad
        try:
            doc.save(tmp_pfad)
            try:
                if os.path.exists(pfad_final):
                    os.replace(tmp_pfad, pfad_final)
                else:
                    os.rename(tmp_pfad, pfad_final)
            except (PermissionError, OSError):
                # Zieldatei gesperrt (z.B. in Word offen) → Ausweich-Name
                base, ext = os.path.splitext(pfad_final)
                ts = datetime.now().strftime('%H%M%S')
                pfad_final = f'{base}_{ts}{ext}'
                os.rename(tmp_pfad, pfad_final)
                warnungen.append(
                    f'Die Datei "{os.path.basename(self.ausgabe_pfad)}" war noch in Word geöffnet.\n'
                    f'Das Dokument wurde stattdessen gespeichert als:\n{os.path.basename(pfad_final)}'
                )
        finally:
            if os.path.exists(tmp_pfad):
                try:
                    os.remove(tmp_pfad)
                except OSError:
                    pass

        try:
            from functions.dokument_archiv import kopiere_ins_archiv
            kopiere_ins_archiv(pfad_final, "staerkemeldung")
        except Exception:
            pass
        return pfad_final, warnungen

    # ------------------------------------------------------------------
    # Private Methoden
    # ------------------------------------------------------------------

    def _add_header(self, doc):
        """DRK-Kopfzeile: Logo links, Organisationstext rechts, Trennlinie."""
        section = doc.sections[0]
        header  = section.header

        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.autofit = False

        # Logo (links)
        logo_path = Path(LOGO_PFAD)
        if logo_path.exists():
            logo_cell = header_table.rows[0].cells[0]
            logo_para = logo_cell.paragraphs[0]
            run = logo_para.add_run()
            run.add_picture(str(logo_path), width=Inches(1.2))
        else:
            header_table.rows[0].cells[0].paragraphs[0].add_run("[Logo nicht gefunden]")

        # Text (rechts)
        text_cell = header_table.rows[0].cells[1]
        text_para = text_cell.paragraphs[0]
        text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        line1 = text_para.add_run("Deutsches Rotes Kreuz Kreisverband Köln e.V.\n")
        line1.font.size = Pt(10)
        line1.font.bold = True

        line2 = text_para.add_run(
            "- Unfallhilfsstelle und Betreuungsstelle für Behinderte am Flughafen Köln/Bonn"
        )
        line2.font.size = Pt(9)

        header.add_paragraph("_" * 80)

    def _add_footer(self, doc):
        """DRK-Fußzeile mit Kontaktdaten."""
        section    = doc.sections[0]
        footer     = section.footer
        footer_p   = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_p.add_run(
            "Telefon: +49 220340 – 2323\t\t"
            "email: flughafen@drk-koeln.de\t\t"
            "Stationsleitung: Lars Peters"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_dienst_gruppe(self, doc, mitarbeiter_liste: list, ist_dispo: bool = False):
        """
        Fügt Mitarbeiter gruppiert nach Uhrzeit ein.
        Format:   06:00 bis 14:00   [TAB]   Müller / Schmidt / Meier
        Tabulator-Stop bei 4,5 cm (2550 Twips), Hanging-Indent.
        """
        # Zeiten-Gruppen aufbauen
        zeiten_gruppen: dict[str, list[str]] = {}

        for person in mitarbeiter_liste:
            start = (person.get('start_zeit') or '')[:5]
            end   = (person.get('end_zeit')   or '')[:5]

            # Dispo: Minuten abschneiden – außer wenn der Nutzer die Zeit manuell angepasst hat
            if ist_dispo and not person.get('manuell_geaendert'):
                if start and ':' in start:
                    start = f"{int(start.split(':')[0]):02d}:00"
                if end and ':' in end:
                    end = f"{int(end.split(':')[0]):02d}:00"

            # Spezialfall NF 16:00–04:00
            if start == '16:00' and end in ('04:00', '04'):
                zeit_key = '16:00 bis 04:00'
            else:
                zeit_key = f"{start} bis {end}"

            zeiten_gruppen.setdefault(zeit_key, []).append(person.get('anzeigename', ''))

        # Gruppen als Paragraph mit Tabulator-Stop ausgeben
        for zeit, namen in sorted(zeiten_gruppen.items()):
            namen_text = " / ".join(namen)
            para       = doc.add_paragraph()

            # Tabulator-Stop bei 4,5 cm + Hanging-Indent
            pPr = para._p.get_or_add_pPr()

            tabs = OxmlElement('w:tabs')
            tab  = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'left')
            tab.set(qn('w:pos'), '2550')
            tabs.append(tab)
            pPr.append(tabs)

            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'),    '2550')
            ind.set(qn('w:hanging'), '2550')
            pPr.append(ind)

            zeit_run = para.add_run(f"{zeit}\t")
            zeit_run.font.size = Pt(10)

            namen_run = para.add_run(namen_text)
            namen_run.font.size = Pt(10)
