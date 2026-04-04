# -*- coding: utf-8 -*-
"""
Staerkemeldung mit Dashboard-Panel (neues Format)
Linke Spalte : Blaues Dashboard-Panel (Datum, PAX-Kennzahlen, Schichtleiter, Bulmor)
Rechte Spalte: Klassische Staerkemeldung (Zeitgruppen wie StaerkemeldungExport)
Layout        : A4 Hochformat, 2-spaltige Tabelle, 0.8 cm Raender
"""

from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEL   = "+49 2203 40-2323"
MAIL  = "erste-hilfe-station-flughafen@drk-koeln.de"

BG_DUNKEL = "1A3460"
HE        = "C8DAFF"
AZ        = "00C8FF"
WEISS     = "FFFFFF"

_LOGO = (
    Path(__file__).resolve().parents[1]
    / "Daten" / "Email" / "Logo.jpg"
)

_DISPO_TAG_TYPEN   = {"DT", "DT3", "D"}
_DISPO_NACHT_TYPEN = {"DN", "DN3"}


def _rgb(hx):
    hx = hx.lstrip("#")
    return RGBColor(int(hx[0:2],16), int(hx[2:4],16), int(hx[4:6],16))

def _set_bg(cell, hx):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"), hx.lstrip("#").upper()); pr.append(s)

def _no_border(cell):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    b=OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        e=OxmlElement(f"w:{side}"); e.set(qn("w:val"),"none"); b.append(e)
    pr.append(b)

def _para(cell, text, bold=False, size=9, fg="000000", align="left", sa=0, sb=0):
    p=cell.add_paragraph()
    p.alignment=(WD_ALIGN_PARAGRAPH.CENTER if align=="center" else
                 WD_ALIGN_PARAGRAPH.RIGHT  if align=="right"  else
                 WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after=Pt(sa); p.paragraph_format.space_before=Pt(sb)
    r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(size); r.font.color.rgb=_rgb(fg)
    r.font.name="Aptos"
    return p

def _trenn(cell, hx, oben=False):
    p=cell.add_paragraph()
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(1)
    pPr=p._p.get_or_add_pPr(); bdr=OxmlElement("w:pBdr")
    e=OxmlElement("w:top" if oben else "w:bottom")
    e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"6")
    e.set(qn("w:space"),"1"); e.set(qn("w:color"),hx.upper())
    bdr.append(e); pPr.append(bdr)

def _zeit_key(p, ist_dispo=False):
    start=(p.get("start_zeit") or "")[:5]
    end  =(p.get("end_zeit")   or "")[:5]
    if ist_dispo and not p.get("manuell_geaendert"):
        if start and ":" in start: start=f"{int(start.split(':')[0]):02d}:00"
        if end   and ":" in end:   end  =f"{int(end.split(':')[0]):02d}:00"
    return f"{start}-{end}" if (start and end) else "?-?"

def _zeitgruppen(personen, ist_dispo=False):
    g={}
    for p in personen:
        key=_zeit_key(p,ist_dispo)
        name=p.get("anzeigename") or p.get("vollname") or "?"
        g.setdefault(key,[]).append(name)
    return g

def _zeitgruppen_para(cell, gruppen, size=10):
    for zeit,namen in sorted(gruppen.items()):
        p=cell.add_paragraph(); pPr=p._p.get_or_add_pPr()
        tabs=OxmlElement("w:tabs"); tab=OxmlElement("w:tab")
        tab.set(qn("w:val"),"left"); tab.set(qn("w:pos"),"2550")
        tabs.append(tab); pPr.append(tabs)
        ind=OxmlElement("w:ind"); ind.set(qn("w:left"),"2550"); ind.set(qn("w:hanging"),"2550")
        pPr.append(ind)
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(4)
        rz=p.add_run(f"{zeit}\t"); rz.font.size=Pt(size); rz.font.name="Aptos"
        rn=p.add_run(" / ".join(namen)); rn.font.size=Pt(size); rn.font.name="Aptos"

def _bul_farben(n):
    if n<=2: return "FF3333","KRITISCH"
    if n==3: return "E07800","EINGESCHRAENKT"
    return "10A050","VOLLSTÄNDIG"


class StaerkemeldungDashboardExport:
    BULMOR_GESAMT=5

    def __init__(self, dienstplan_data, ausgabe_pfad, von_datum, bis_datum,
                 pax_zahl=0, ausgeschlossene_vollnamen=None,
                 bulmor_aktiv=5, einsaetze_zahl=0,
                 sl_tag_name='', sl_nacht_name='', stationsleitung=''):
        self._data          =dienstplan_data
        self._pfad          =ausgabe_pfad
        self._von           =von_datum
        self._bis           =bis_datum
        self._pax           =pax_zahl
        self._einsaetze     =einsaetze_zahl
        self._sl_tag_name   =sl_tag_name.strip()
        self._sl_nacht_name =sl_nacht_name.strip()
        self._ausschl       ={n.lower().strip() for n in (ausgeschlossene_vollnamen or set())}
        self._bul_aktiv     =bulmor_aktiv
        self._stationslt    =stationsleitung
        # Schichtleiter-Namen (Nachname allein oder "Vorname Nachname") als Ausschluss-Set
        sl_namen = {n.lower().strip() for n in [sl_tag_name, sl_nacht_name] if n.strip()}
        kranke_ids=set(id(m) for m in dienstplan_data.get("kranke",[]))
        def _ist_sl(m):
            vn = m.get("vollname","").lower().strip()
            nachname = vn.split(",")[0].strip() if "," in vn else vn.split()[-1].lower() if vn else ""
            return vn in sl_namen or nachname in sl_namen
        self._dispo=sorted(
            [m for m in dienstplan_data.get("dispo",[])
             if id(m) not in kranke_ids
             and m.get("vollname","").lower() not in self._ausschl
             and not _ist_sl(m)],
            key=lambda x: x.get("start_zeit") or "ZZZZ")
        self._betreuer=sorted(
            [m for m in dienstplan_data.get("betreuer",[])
             if id(m) not in kranke_ids and m.get("vollname","").lower() not in self._ausschl],
            key=lambda x: x.get("start_zeit") or "ZZZZ")

    def export(self):
        warnungen=[]
        doc=self._init_doc()
        self._add_header(doc)
        self._add_footer(doc)
        self._fill_body(doc,warnungen)
        # Alle <w:p>-Elemente direkt im Body entfernen (nur Tabelle + sectPr sollen bleiben)
        body=doc._element.body
        for p_elem in list(body.iterchildren()):
            tag=p_elem.tag.split("}")[-1] if "}" in p_elem.tag else p_elem.tag
            if tag=="p":
                body.remove(p_elem)

        # Trailing-Paragraph mit echter 1pt-Zeilenhöhe bauen (OOXML-Pflicht)
        # Er muss VOR dem <w:sectPr> eingefügt werden, sonst landet er auf Seite 2.
        trail=OxmlElement("w:p")
        trail_pPr=OxmlElement("w:pPr")
        # snapToGrid=0: verhindert dass Word den Paragraph auf linePitch=360 rastet
        snap=OxmlElement("w:snapToGrid"); snap.set(qn("w:val"),"0")
        trail_pPr.append(snap)
        # Kein Abstand vor/nach, exakt 1pt Zeilenhöhe
        trail_sp=OxmlElement("w:spacing")
        trail_sp.set(qn("w:before"),"0"); trail_sp.set(qn("w:after"),"0")
        trail_sp.set(qn("w:line"),"20"); trail_sp.set(qn("w:lineRule"),"exact")
        trail_pPr.append(trail_sp)
        trail_rPr=OxmlElement("w:rPr")
        trail_sz=OxmlElement("w:sz"); trail_sz.set(qn("w:val"),"2")
        trail_rPr.append(trail_sz); trail_pPr.append(trail_rPr)
        trail.append(trail_pPr)

        # VOR sectPr einfügen (nicht danach)
        sectPr=body.find(qn("w:sectPr"))
        if sectPr is not None:
            body.insert(list(body).index(sectPr), trail)
        else:
            body.append(trail)
        doc.save(self._pfad)
        return self._pfad, warnungen

    def _init_doc(self):
        doc=Document()
        for sec in doc.sections:
            sec.page_width=Cm(21.0);  sec.page_height=Cm(29.7)
            sec.top_margin=Cm(1.8);   sec.bottom_margin=Cm(0.8)
            sec.left_margin=Cm(0.5);  sec.right_margin=Cm(0.5)
            sec.header_distance=Cm(0.4); sec.footer_distance=Cm(0.4)
        # docGrid auf noGrid setzen:
        # Ohne expliziten type-Wert interpretiert Word den Grid als 'lines',
        # was alle Inhalte auf linePitch=360 twips rastet und trotz
        # hRule=exact die Seite auf 2 Seiten anwachsen lassen kann.
        sectPr=doc.sections[0]._sectPr
        docGrid=sectPr.find(qn('w:docGrid'))
        if docGrid is None:
            docGrid=OxmlElement('w:docGrid')
            sectPr.append(docGrid)
        docGrid.set(qn('w:type'),'noGrid')
        return doc

    def _add_header(self,doc):
        hdr=doc.sections[0].header
        ht=hdr.add_table(rows=1,cols=2,width=Inches(7.8)); ht.autofit=False
        if _LOGO.exists():
            lp=ht.rows[0].cells[0].paragraphs[0]
            lp.add_run().add_picture(str(_LOGO),width=Inches(1.0))
        tp=ht.rows[0].cells[1].paragraphs[0]; tp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        r1=tp.add_run("Deutsches Rotes Kreuz Kreisverband Koeln e.V.\n")
        r1.font.size=Pt(9); r1.font.bold=True; r1.font.name="Aptos"
        r2=tp.add_run("Unfallhilfsstelle und Betreuungsstelle - Flughafen Koeln/Bonn")
        r2.font.size=Pt(8); r2.font.name="Aptos"
        hdr.add_paragraph("_"*90)

    def _add_footer(self,doc):
        ftr=doc.sections[0].footer; fp=ftr.paragraphs[0]
        fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        sl=f"   |   Stationsleitung: {self._stationslt}" if self._stationslt else ""
        fr=fp.add_run(f"Tel: {TEL}   |   {MAIL}{sl}")
        fr.font.size=Pt(8); fr.font.color.rgb=_rgb("777777"); fr.font.name="Aptos"

    def _fill_body(self,doc,warnungen):
        L_W=Cm(5.5); R_W=Cm(14.0)
        main=doc.add_table(rows=1,cols=2); main.style="Table Grid"
        lc=main.cell(0,0); rc=main.cell(0,1)
        lc.width=L_W; rc.width=R_W
        _no_border(lc); _no_border(rc)
        lc.vertical_alignment=WD_ALIGN_VERTICAL.TOP
        rc.vertical_alignment=WD_ALIGN_VERTICAL.TOP
        # Seitenhoehe: h=16838 top=1020 bot=454 → Brutto-Koerper=15364 twips.
        # ABER: Logo 452x339px @ 1.0" Breite = 1080 twips hoch.
        #       Header-Puffer = top - hdr_distance = 1020 - 227 = 793 twips.
        #       Logo-Überlauf = 1080 - 793 = 287 twips → Word schiebt Body nach unten.
        #       Effektive Koerperhoehe ≈ 15364 - 287 = ~15077 (Word-intern etwas weniger).
        # Empirisch: manuell auf 14417 verkleinert = 1 Seite. Daher sicherer Wert:
        TABLE_H_TWIPS = 14350
        tr=main.rows[0]._tr; trPr=tr.get_or_add_trPr()
        trH=OxmlElement("w:trHeight")
        trH.set(qn("w:val"),str(TABLE_H_TWIPS)); trH.set(qn("w:hRule"),"exact")
        trPr.append(trH)
        self._build_links(lc)
        self._build_rechts(rc)

    def _find_sl(self,nacht):
        manuell=self._sl_nacht_name if nacht else self._sl_tag_name
        typen=_DISPO_NACHT_TYPEN if nacht else _DISPO_TAG_TYPEN
        for p in self._dispo:
            dk=(p.get("dienst_kategorie") or "").upper()
            if dk in typen:
                start=(p.get("start_zeit") or "")[:5]
                end  =(p.get("end_zeit")   or "")[:5]
                zeit =f"{start}-{end}" if (start and end) else "?-?"
                # Name nur ausgeben wenn manuell eingegeben, sonst leer lassen
                return manuell, zeit
        # Kein Dispo-Eintrag gefunden
        if manuell:
            return manuell, "?"
        # Kein Eintrag und kein Name → Zeile trotzdem zeigen, aber leer
        return "", ""

    def _build_links(self,lc):
        _para(lc,"Deutsches Rotes Kreuz",bold=True,size=10.5,fg="000000",align="center",sb=2)
        _para(lc,"Kreisverband Koeln e.V.",size=9,fg="000000",align="center")
        _para(lc,"Erste-Hilfe-Station - Flughafen Koeln/Bonn",size=8,fg="000000",align="center")
        _para(lc,TEL,bold=True,size=9,fg="000000",align="center",sa=1)
        _trenn(lc,AZ)
        try:
            from database.pax_db import lade_jahres_pax,lade_tages_pax,lade_jahres_einsaetze,lade_tages_einsaetze
            from datetime import timedelta
            pax_jahr=lade_jahres_pax(self._von.year)
            pax_gestern=lade_tages_pax((self._von-timedelta(days=1)).strftime("%Y-%m-%d"))
            einsaetze_gestern=lade_tages_einsaetze((self._von-timedelta(days=1)).strftime("%Y-%m-%d"))
        except Exception:
            pax_jahr=0; pax_gestern=0; einsaetze_gestern=0
        for lbl,val in [
            ("+ PAX aktuelles Jahr", f"{pax_jahr:,}".replace(",",".")),
            ("+ PAX Vortag",         f"{pax_gestern:,}".replace(",",".") if pax_gestern else "-"),
            ("+ PAX heute",          f"{self._pax:,}".replace(",",".")),
        ]:
            p=lc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(0)
            r1=p.add_run(f"{lbl}\n"); r1.font.size=Pt(8.5); r1.font.color.rgb=_rgb("000000")
            r1.bold=True; r1.font.name="Aptos"
            r2=p.add_run(val); r2.bold=False; r2.font.size=Pt(13); r2.font.color.rgb=_rgb("000000")
            r2.font.name="Aptos"
        _trenn(lc,AZ)
        for lbl,val in [
            ("SL-Einsaetze Vortag", str(einsaetze_gestern) if einsaetze_gestern else "-"),
            ("SL-Einsaetze heute",  str(self._einsaetze)),
        ]:
            p=lc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(0)
            r1=p.add_run(f"{lbl}\n"); r1.font.size=Pt(8.5); r1.font.color.rgb=_rgb("000000")
            r1.bold=True; r1.font.name="Aptos"
            r2=p.add_run(val); r2.bold=False; r2.font.size=Pt(13); r2.font.color.rgb=_rgb("000000")
            r2.font.name="Aptos"
        _trenn(lc,AZ,oben=True)
        fc_hex,lbl_bul=_bul_farben(self._bul_aktiv)
        _para(lc,"BULMOR - FAHRZEUGSTATUS",bold=True,size=8.5,fg="000000",align="center",sb=1)
        bul_tbl=lc.add_table(rows=1,cols=self.BULMOR_GESAMT); bul_tbl.style="Table Grid"
        for i in range(self.BULMOR_GESAMT):
            c=bul_tbl.cell(0,i); _set_bg(c,fc_hex if i<self._bul_aktiv else "444444")
            p=c.add_paragraph(f"B{i+1}"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
            p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=_rgb(WEISS)
            p.runs[0].bold=True; p.runs[0].font.name="Aptos"
        ges=lc.add_paragraph(); ges.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ges.paragraph_format.space_before=Pt(1); ges.paragraph_format.space_after=Pt(0)
        r1=ges.add_run(f"{self._bul_aktiv}/{self.BULMOR_GESAMT}  ")
        r1.bold=True; r1.font.size=Pt(13); r1.font.color.rgb=_rgb(fc_hex); r1.font.name="Aptos"
        r2=ges.add_run(lbl_bul); r2.font.size=Pt(10); r2.font.color.rgb=_rgb(fc_hex); r2.font.name="Aptos"

    def _build_rechts(self,rc):
        datum_str=self._von.strftime("%d.%m.%Y"); bis_str=self._bis.strftime("%d.%m.%Y")
        sl_tag=self._find_sl(False); sl_nacht=self._find_sl(True)
        ph=rc.add_paragraph(); ph.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before=Pt(2); ph.paragraph_format.space_after=Pt(4)
        rh=ph.add_run("STÄRKEMELDUNG"); rh.font.bold=True; rh.font.size=Pt(13)
        rh.font.color.rgb=_rgb("000000"); rh.font.name="Aptos"
        pz=rc.add_paragraph()
        pz.paragraph_format.space_before=Pt(0); pz.paragraph_format.space_after=Pt(6)
        pPr_pz=pz._p.get_or_add_pPr()
        tabs_pz=OxmlElement("w:tabs"); tab_pz=OxmlElement("w:tab")
        tab_pz.set(qn("w:val"),"left"); tab_pz.set(qn("w:pos"),"2550")
        tabs_pz.append(tab_pz); pPr_pz.append(tabs_pz)
        ind_pz=OxmlElement("w:ind"); ind_pz.set(qn("w:left"),"2550"); ind_pz.set(qn("w:hanging"),"2550")
        pPr_pz.append(ind_pz)
        zeitraum=datum_str if datum_str==bis_str else f"{datum_str} bis {bis_str}"
        rz_lbl=pz.add_run("Zeitraum:\t"); rz_lbl.font.size=Pt(11); rz_lbl.font.bold=True; rz_lbl.font.name="Aptos"
        rz_val=pz.add_run(zeitraum); rz_val.font.size=Pt(11); rz_val.font.bold=False; rz_val.font.name="Aptos"
        if True:  # Schichtleiter-Block immer anzeigen
            ph_sl=rc.add_paragraph()
            ph_sl.paragraph_format.space_before=Pt(2); ph_sl.paragraph_format.space_after=Pt(1)
            rh_sl=ph_sl.add_run("Schichtleiter"); rh_sl.font.bold=True; rh_sl.font.size=Pt(11)
            rh_sl.font.name="Aptos"
            for sl in [sl_tag, sl_nacht]:
                if sl is None: continue
                p_sl=rc.add_paragraph(); pPr=p_sl._p.get_or_add_pPr()
                tabs=OxmlElement("w:tabs"); tab=OxmlElement("w:tab")
                tab.set(qn("w:val"),"left"); tab.set(qn("w:pos"),"2550")
                tabs.append(tab); pPr.append(tabs)
                ind=OxmlElement("w:ind"); ind.set(qn("w:left"),"2550"); ind.set(qn("w:hanging"),"2550")
                pPr.append(ind)
                p_sl.paragraph_format.space_before=Pt(0); p_sl.paragraph_format.space_after=Pt(0)
                r_pf=p_sl.add_run(f"{sl[1]}\t"); r_pf.font.size=Pt(10.5); r_pf.font.name="Aptos"
                r_nm=p_sl.add_run(sl[0]); r_nm.font.size=Pt(10.5); r_nm.bold=False
                r_nm.font.name="Aptos"
            rc.add_paragraph().paragraph_format.space_after=Pt(2)
        ph_d=rc.add_paragraph(); ph_d.paragraph_format.space_before=Pt(2); ph_d.paragraph_format.space_after=Pt(1)
        rh_d=ph_d.add_run("Disposition"); rh_d.font.bold=True; rh_d.font.size=Pt(11)
        rh_d.font.name="Aptos"
        _zeitgruppen_para(rc,_zeitgruppen(self._dispo,ist_dispo=True))
        rc.add_paragraph().paragraph_format.space_after=Pt(2)
        ph_b=rc.add_paragraph(); ph_b.paragraph_format.space_before=Pt(2); ph_b.paragraph_format.space_after=Pt(1)
        rh_b=ph_b.add_run("Behindertenbetreuer"); rh_b.font.bold=True; rh_b.font.size=Pt(11)
        rh_b.font.name="Aptos"
        _zeitgruppen_para(rc,_zeitgruppen(self._betreuer))