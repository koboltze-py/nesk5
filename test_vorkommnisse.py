"""
Tests für die Vorkommnisse-Funktion.

Testet:
  1. DB-Schicht  – speichern, laden, aktualisieren, löschen, lade_alle
  2. Word-Export – 5 verschiedene Szenarien, Ausgabe nach
                   C:\\Users\\DRKairport\\OneDrive - Deutsches Rotes Kreuz - ...
                   \\Desktop\\bei\\vor\\

Ausführen:
    python3.13 test_vorkommnisse.py
"""
from __future__ import annotations

import os
import sys
import tempfile
import unittest
from pathlib import Path

# Projektordner in den Pfad aufnehmen
BASE_TEST_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(BASE_TEST_DIR))

# Word-Ausgabeordner
WORD_OUT = Path(
    r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
    r"\Desktop\bei\vor"
)

# ── Testdaten ──────────────────────────────────────────────────────────────────

TESTFALL_EINFACH = {
    "flug":          "XQ983",
    "typ":           "PRM-Betreuung",
    "datum":         "20.04.2026",
    "ort":           "Köln/Bonn (CGN)",
    "offblock_plan": "10:30 Uhr",
    "offblock_ist":  "10:45 Uhr (+15 Min.)",
    "erstellt_von":  "M. Mustermann",
    "passagiere":    [
        ["Fischer, Anna", "WCHS", "Rollstuhl bereitgestellt"],
        ["Müller, Klaus", "BLND", "Begleitung bis Gate"],
    ],
    "personal":      [
        ["Schmidt, P.", "Betreuer", "Hauptverantwortlich"],
        ["Weber, T.",   "Fahrer",   "Bulmor 3"],
    ],
    "chronologie":   [
        ["10:15", "Passagierin am Check-in abgeholt"],
        ["10:25", "Rollstuhl bereitgestellt, Transfer zu Gate B12"],
        ["10:40", "Übergabe an Boarding-Personal"],
    ],
    "ursache":       "Rollstuhl wurde nicht rechtzeitig am Check-in bereitgestellt.",
    "ergebnis":      "Flug konnte pünktlich abgefertigt werden. Passagierin zufrieden.",
}

TESTFALL_MEDIZINISCH = {
    "flug":          "EW2241",
    "typ":           "Medizinischer Notfall",
    "datum":         "19.04.2026",
    "ort":           "Gate C15, Köln/Bonn (CGN)",
    "offblock_plan": "14:00 Uhr",
    "offblock_ist":  "14:00 Uhr",
    "erstellt_von":  "K. Heinz",
    "passagiere":    [
        ["Berger, Hans", "DPNA", "Herzrhythmusstörung"],
    ],
    "personal":      [
        ["Dr. Lange",    "Notarzt",    ""],
        ["Richter, S.",  "Sanitäter",  "Erstversorgung"],
        ["Becker, M.",   "Betreuer",   "Begleitung"],
    ],
    "chronologie":   [
        ["13:45", "Notruf eingehend – Passagier bewusstlos am Gate"],
        ["13:46", "Ersthelfer vor Ort, AED angewendet"],
        ["13:50", "Notarzt eingetroffen"],
        ["14:10", "Patient stabilisiert, Übergabe an RTW"],
        ["14:20", "Flugzeug verzögert wegen medizinischem Vorfall"],
    ],
    "ursache":       (
        "Passagier hatte vorbekannte Herzerkrankung.\n"
        "Unzureichende Vorausinformation von der Airline."
    ),
    "ergebnis":      "Patient wurde ins Krankenhaus eingeliefert. Flug 45 Min. verspätet.",
}

TESTFALL_SICHERHEIT = {
    "flug":          "FR8819",
    "typ":           "Sicherheitsvorfall",
    "datum":         "18.04.2026",
    "ort":           "Rollfeld / Vorfeldbereich",
    "offblock_plan": "06:00 Uhr",
    "offblock_ist":  "06:00 Uhr",
    "erstellt_von":  "A. Koch",
    "passagiere":    [],
    "personal":      [
        ["Koch, A.",    "Sicherheitsbeauftragter", "Erstmeldung"],
        ["Neumann, F.", "Vorfeldbetreuer",          ""],
    ],
    "chronologie":   [
        ["05:50", "Unbekannte Person ohne Ausweis im Vorfeldbereich gesichtet"],
        ["05:52", "Sicherheitspersonal alarmiert"],
        ["05:58", "Person gestellt und Ausweis geprüft – Fehlalarm"],
    ],
    "ursache":       "Mitarbeiter hatte den Ausweis im Fahrzeug vergessen.",
    "ergebnis":      "Kein sicherheitsrelevanter Vorfall. Mitarbeiter verwarnt.",
}

TESTFALL_VERSPAETUNG = {
    "flug":          "QY401",
    "typ":           "Verspätung/Offblock",
    "datum":         "17.04.2026",
    "ort":           "Gate A7, Köln/Bonn (CGN)",
    "offblock_plan": "07:15 Uhr",
    "offblock_ist":  "08:30 Uhr (+75 Min.)",
    "erstellt_von":  "B. Schulz",
    "passagiere":    [
        ["Tanaka, Yuki",     "WCHS", "Benötigt Rollstuhlhilfe"],
        ["Tanaka, Hiroshi",  "WCHS", "Begleitung"],
        ["Hoffmann, R.",     "WCHC", "Liegendtransport"],
    ],
    "personal":      [
        ["Schulz, B.",    "Dienstleiter",  "Koordination"],
        ["Peters, L.",    "Betreuer",      ""],
        ["Wagner, C.",    "Fahrer",        "Bulmor 1"],
        ["Braun, M.",     "Betreuer",      ""],
    ],
    "chronologie":   [
        ["07:00", "PRM-Passagiere am Check-in abgeholt"],
        ["07:15", "Warten am Gate – Offblock-Verzögerung durch Technik"],
        ["07:45", "Information: Flugzeug noch nicht bereit"],
        ["08:10", "Boarding freigegeben"],
        ["08:30", "Offblock nach 75 Min. Verspätung"],
    ],
    "ursache":       "Technischer Defekt am Flugzeug (Hydraulik). Reparatur vor Ort.",
    "ergebnis":      "Alle PRM-Passagiere pünktlich eingeschifft. Keine Beschwerden.",
}

TESTFALL_MINIMAL = {
    "flug":          "AB1234",
    "typ":           "Sonstiges",
    "datum":         "20.04.2026",
    "ort":           "Köln/Bonn (CGN)",
    "offblock_plan": "00:00 Uhr",
    "offblock_ist":  "00:00 Uhr",
    "erstellt_von":  "Test Nutzer",
    "passagiere":    [],
    "personal":      [],
    "chronologie":   [],
    "ursache":       "",
    "ergebnis":      "",
}

ALLE_TESTFAELLE = [
    ("PRM_Betreuung_XQ983",      TESTFALL_EINFACH),
    ("Medizinischer_Notfall_EW2241", TESTFALL_MEDIZINISCH),
    ("Sicherheitsvorfall_FR8819", TESTFALL_SICHERHEIT),
    ("Verspaetung_QY401",         TESTFALL_VERSPAETUNG),
    ("Minimal_AB1234",            TESTFALL_MINIMAL),
]


# ══════════════════════════════════════════════════════════════════════════════
#  1. DB-Tests (kein Qt nötig)
# ══════════════════════════════════════════════════════════════════════════════

class TestVorkommnissDB(unittest.TestCase):
    """Testet die CRUD-Funktionen in functions/vorkommnisse_db.py."""

    def setUp(self):
        """Temporäre Test-DB anlegen, VORKOMMNISSE_DB_PATH überschreiben."""
        self._tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        self._tmp.close()
        import config as _cfg
        self._orig_path = _cfg.VORKOMMNISSE_DB_PATH
        _cfg.VORKOMMNISSE_DB_PATH = self._tmp.name
        # Modul neu laden damit _DB_PFAD aktualisiert wird
        import importlib
        import functions.vorkommnisse_db as _m
        importlib.reload(_m)

    def tearDown(self):
        import config as _cfg
        _cfg.VORKOMMNISSE_DB_PATH = self._orig_path
        import importlib, functions.vorkommnisse_db as _m
        importlib.reload(_m)
        try:
            os.unlink(self._tmp.name)
        except Exception:
            pass

    def _db(self):
        import functions.vorkommnisse_db as m
        return m

    # ── speichern ──────────────────────────────────────────────────────────────

    def test_speichern_gibt_id_zurueck(self):
        m = self._db()
        vid = m.speichern(TESTFALL_EINFACH)
        self.assertIsInstance(vid, int)
        self.assertGreater(vid, 0)

    def test_speichern_mehrere_ergibt_unterschiedliche_ids(self):
        m = self._db()
        id1 = m.speichern(TESTFALL_EINFACH)
        id2 = m.speichern(TESTFALL_MEDIZINISCH)
        self.assertNotEqual(id1, id2)

    # ── lade_ein ───────────────────────────────────────────────────────────────

    def test_lade_ein_grunddaten(self):
        m = self._db()
        vid = m.speichern(TESTFALL_EINFACH)
        d = m.lade_ein(vid)
        self.assertIsNotNone(d)
        self.assertEqual(d["flug"],  "XQ983")
        self.assertEqual(d["typ"],   "PRM-Betreuung")
        self.assertEqual(d["datum"], "20.04.2026")
        self.assertEqual(d["ort"],   "Köln/Bonn (CGN)")
        self.assertEqual(d["erstellt_von"], "M. Mustermann")

    def test_lade_ein_json_felder_dekodiert(self):
        m = self._db()
        vid = m.speichern(TESTFALL_EINFACH)
        d = m.lade_ein(vid)
        self.assertIsInstance(d["passagiere"],  list)
        self.assertIsInstance(d["personal"],    list)
        self.assertIsInstance(d["chronologie"], list)
        self.assertEqual(len(d["passagiere"]),  2)
        self.assertEqual(len(d["personal"]),    2)
        self.assertEqual(len(d["chronologie"]), 3)

    def test_lade_ein_nicht_vorhanden_gibt_none(self):
        m = self._db()
        d = m.lade_ein(9999)
        self.assertIsNone(d)

    def test_lade_ein_leere_json_felder(self):
        m = self._db()
        vid = m.speichern(TESTFALL_MINIMAL)
        d = m.lade_ein(vid)
        self.assertEqual(d["passagiere"],  [])
        self.assertEqual(d["personal"],    [])
        self.assertEqual(d["chronologie"], [])

    # ── lade_alle ──────────────────────────────────────────────────────────────

    def test_lade_alle_leer(self):
        m = self._db()
        alle = m.lade_alle()
        self.assertEqual(alle, [])

    def test_lade_alle_gibt_alle_zurueck(self):
        m = self._db()
        m.speichern(TESTFALL_EINFACH)
        m.speichern(TESTFALL_MEDIZINISCH)
        m.speichern(TESTFALL_SICHERHEIT)
        alle = m.lade_alle()
        self.assertEqual(len(alle), 3)

    def test_lade_alle_neuestes_zuerst(self):
        m = self._db()
        id1 = m.speichern(TESTFALL_EINFACH)
        id2 = m.speichern(TESTFALL_MEDIZINISCH)
        alle = m.lade_alle()
        # Neuester (höchste ID) soll an erster Stelle stehen
        self.assertEqual(alle[0]["id"], id2)
        self.assertEqual(alle[1]["id"], id1)

    # ── aktualisieren ──────────────────────────────────────────────────────────

    def test_aktualisieren_aendert_felder(self):
        m = self._db()
        vid = m.speichern(TESTFALL_EINFACH)
        geaendert = {**TESTFALL_EINFACH, "flug": "XX999", "ort": "München (MUC)"}
        m.aktualisieren(vid, geaendert)
        d = m.lade_ein(vid)
        self.assertEqual(d["flug"], "XX999")
        self.assertEqual(d["ort"],  "München (MUC)")

    def test_aktualisieren_json_felder(self):
        m = self._db()
        vid = m.speichern(TESTFALL_EINFACH)
        neue_pax = [["Neuer, Name", "WCHR", ""]]
        geaendert = {**TESTFALL_EINFACH, "passagiere": neue_pax}
        m.aktualisieren(vid, geaendert)
        d = m.lade_ein(vid)
        self.assertEqual(len(d["passagiere"]), 1)
        self.assertEqual(d["passagiere"][0][0], "Neuer, Name")

    # ── loeschen ───────────────────────────────────────────────────────────────

    def test_loeschen_entfernt_datensatz(self):
        m = self._db()
        vid = m.speichern(TESTFALL_EINFACH)
        m.loeschen(vid)
        d = m.lade_ein(vid)
        self.assertIsNone(d)

    def test_loeschen_andere_bleiben_erhalten(self):
        m = self._db()
        id1 = m.speichern(TESTFALL_EINFACH)
        id2 = m.speichern(TESTFALL_MEDIZINISCH)
        m.loeschen(id1)
        alle = m.lade_alle()
        self.assertEqual(len(alle), 1)
        self.assertEqual(alle[0]["id"], id2)

    def test_loeschen_nicht_vorhanden_kein_fehler(self):
        m = self._db()
        # Soll keine Exception werfen
        m.loeschen(9999)

    # ── Vollständiger Roundtrip ────────────────────────────────────────────────

    def test_roundtrip_alle_felder(self):
        m = self._db()
        vid = m.speichern(TESTFALL_VERSPAETUNG)
        d = m.lade_ein(vid)
        self.assertEqual(d["flug"],          TESTFALL_VERSPAETUNG["flug"])
        self.assertEqual(d["typ"],           TESTFALL_VERSPAETUNG["typ"])
        self.assertEqual(d["datum"],         TESTFALL_VERSPAETUNG["datum"])
        self.assertEqual(d["ort"],           TESTFALL_VERSPAETUNG["ort"])
        self.assertEqual(d["offblock_plan"], TESTFALL_VERSPAETUNG["offblock_plan"])
        self.assertEqual(d["offblock_ist"],  TESTFALL_VERSPAETUNG["offblock_ist"])
        self.assertEqual(d["erstellt_von"],  TESTFALL_VERSPAETUNG["erstellt_von"])
        self.assertEqual(d["ursache"],       TESTFALL_VERSPAETUNG["ursache"])
        self.assertEqual(d["ergebnis"],      TESTFALL_VERSPAETUNG["ergebnis"])
        self.assertEqual(len(d["passagiere"]),  3)
        self.assertEqual(len(d["personal"]),    4)
        self.assertEqual(len(d["chronologie"]), 5)


# ══════════════════════════════════════════════════════════════════════════════
#  2. Word-Export-Tests
# ══════════════════════════════════════════════════════════════════════════════

class TestVorkommnisWordExport(unittest.TestCase):
    """Erstellt echte Word-Dateien für alle Testfälle."""

    @classmethod
    def setUpClass(cls):
        """Qt-App einmalig starten (nur für Widget-Instanz nötig)."""
        try:
            from PySide6.QtWidgets import QApplication
            cls._app = QApplication.instance() or QApplication(sys.argv[:1])
        except Exception as e:
            raise unittest.SkipTest(f"Qt nicht verfügbar: {e}")

        try:
            from gui.vorkommnisse import VorkommnisseWidget
            cls._widget = VorkommnisseWidget()
        except Exception as e:
            raise unittest.SkipTest(f"VorkommnisseWidget nicht ladbar: {e}")

        WORD_OUT.mkdir(parents=True, exist_ok=True)

    def _erstelle_word(self, daten: dict, dateiname: str) -> Path:
        pfad = WORD_OUT / dateiname
        self._widget._erstelle_word(str(pfad), daten)
        return pfad

    # ── Einzelne Testfälle ────────────────────────────────────────────────────

    def test_word_prm_betreuung(self):
        pfad = self._erstelle_word(
            TESTFALL_EINFACH, "Test_01_PRM_Betreuung_XQ983.docx"
        )
        self.assertTrue(pfad.exists(), f"Datei nicht erstellt: {pfad}")
        self.assertGreater(pfad.stat().st_size, 1000)
        print(f"  ✓ Word erstellt: {pfad.name}  ({pfad.stat().st_size // 1024} KB)")

    def test_word_medizinischer_notfall(self):
        pfad = self._erstelle_word(
            TESTFALL_MEDIZINISCH, "Test_02_Medizinischer_Notfall_EW2241.docx"
        )
        self.assertTrue(pfad.exists())
        self.assertGreater(pfad.stat().st_size, 1000)
        print(f"  ✓ Word erstellt: {pfad.name}  ({pfad.stat().st_size // 1024} KB)")

    def test_word_sicherheitsvorfall(self):
        pfad = self._erstelle_word(
            TESTFALL_SICHERHEIT, "Test_03_Sicherheitsvorfall_FR8819.docx"
        )
        self.assertTrue(pfad.exists())
        self.assertGreater(pfad.stat().st_size, 1000)
        print(f"  ✓ Word erstellt: {pfad.name}  ({pfad.stat().st_size // 1024} KB)")

    def test_word_verspaetung(self):
        pfad = self._erstelle_word(
            TESTFALL_VERSPAETUNG, "Test_04_Verspaetung_QY401.docx"
        )
        self.assertTrue(pfad.exists())
        self.assertGreater(pfad.stat().st_size, 1000)
        print(f"  ✓ Word erstellt: {pfad.name}  ({pfad.stat().st_size // 1024} KB)")

    def test_word_minimal(self):
        pfad = self._erstelle_word(
            TESTFALL_MINIMAL, "Test_05_Minimal_AB1234.docx"
        )
        self.assertTrue(pfad.exists())
        self.assertGreater(pfad.stat().st_size, 500)
        print(f"  ✓ Word erstellt: {pfad.name}  ({pfad.stat().st_size // 1024} KB)")

    # ── Inhaltsprüfung ────────────────────────────────────────────────────────

    def test_word_inhalt_flugnummer(self):
        """Prüft ob die Flugnummer im Dokument vorhanden ist."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx nicht installiert")
        pfad = self._erstelle_word(TESTFALL_EINFACH, "_tmp_inhalt_test.docx")
        doc = Document(str(pfad))
        volltext = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("XQ983", volltext)
        pfad.unlink(missing_ok=True)

    def test_word_inhalt_passagier(self):
        """Prüft ob ein Passagier-Name im Dokument vorkommt."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx nicht installiert")
        pfad = self._erstelle_word(TESTFALL_EINFACH, "_tmp_pax_test.docx")
        doc = Document(str(pfad))
        # Tabellenzellen durchsuchen
        tabellen_text = ""
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    tabellen_text += cell.text + " "
        self.assertIn("Fischer", tabellen_text)
        pfad.unlink(missing_ok=True)

    def test_word_inhalt_chronologie(self):
        """Prüft ob Chronologie-Einträge vorhanden sind."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx nicht installiert")
        pfad = self._erstelle_word(
            TESTFALL_MEDIZINISCH, "_tmp_chrono_test.docx"
        )
        doc = Document(str(pfad))
        tabellen_text = ""
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    tabellen_text += cell.text + " "
        self.assertIn("Notarzt eingetroffen", tabellen_text)
        pfad.unlink(missing_ok=True)

    def test_word_leer_ohne_passagiere(self):
        """Minimales Dokument (keine Passagiere) darf keinen Fehler werfen."""
        pfad = self._erstelle_word(TESTFALL_MINIMAL, "_tmp_leer_test.docx")
        self.assertTrue(pfad.exists())
        pfad.unlink(missing_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
#  Einstiegspunkt
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 65)
    print("  Nesk3 – Vorkommnisse Tests")
    print(f"  Word-Ausgabe: {WORD_OUT}")
    print("=" * 65)
    loader  = unittest.TestLoader()
    suite   = unittest.TestSuite()
    # DB-Tests zuerst
    suite.addTests(loader.loadTestsFromTestCase(TestVorkommnissDB))
    # Word-Tests danach
    suite.addTests(loader.loadTestsFromTestCase(TestVorkommnisWordExport))
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    print()
    if result.wasSuccessful():
        print(f"✅  Alle {result.testsRun} Tests bestanden.")
        print(f"   Word-Dateien gespeichert in:\n   {WORD_OUT}")
    else:
        print(f"❌  {len(result.failures)} Fehler, {len(result.errors)} Exceptions")
    sys.exit(0 if result.wasSuccessful() else 1)
